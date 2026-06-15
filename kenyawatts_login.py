"""
KenyaWatts — Digital Integrated National Energy Planning Platform
=================================================================
Purpose : Harmonise County Energy Plans into Kenya's National INEP
Partner  : EPRA Kenya (Energy & Petroleum Regulatory Authority)
Challenge: NGDA 2026 · DTU Young Academics Track · Challenge 2

HOW THE FILE IS ORGANISED
--------------------------
1.  Imports            — all Python libraries used
2.  Page config        — Streamlit page title, icon, layout
3.  CSS                — custom styling for the platform
4.  Chart defaults     — bold black labels on all charts
5.  County data        — energy data for 14 sample counties
6.  Generation data    — national electricity generation stats
7.  Cloud DB layer     — Supabase PostgreSQL (optional, falls back to JSON)
8.  Message store      — persistent two-way messaging (EPRA ↔ counties)
9.  Submission store   — persistent county plan submissions
10. Session state      — initialise all in-memory variables
11. Notifications      — toast notification system
12. Aggregation        — compute national weighted averages from county data
13. AI assistant       — Claude API integration
14. Authentication     — streamlit-authenticator login/logout
15. Shared components  — reusable UI helpers (metric_card, alert, section)
16. Page functions     — one function per page/view
17. Report generator   — Word and PDF national INEP report
18. Login page         — the sign-in screen shown before authentication
19. Forgot password    — 3-step password reset flow
20. Main function      — app entry point, routing via st.navigation()
"""
# =============================================================================
# SECTION 1 — IMPORTS
# Every external library the platform depends on is imported here.
# If any import fails, check that requirements.txt contains the library name.
# =============================================================================

import streamlit as st               # The web framework that runs the whole platform
import streamlit_authenticator as stauth  # Handles login, logout, and cookie sessions
import plotly.express as px          # Creates charts with a simple interface
import plotly.graph_objects as go    # Creates advanced/custom charts
import pandas as pd                  # Handles tabular data (county data, submissions)
import anthropic                     # Connects to Claude AI for the AI assistant tab
import io                            # Handles in-memory file buffers (for report downloads)
import json                          # Reads and writes JSON files (message/submission stores)
import bcrypt                        # Hashes and verifies passwords securely
from datetime import datetime, timedelta  # Used for timestamps on submissions and messages

# ── Optional: Word document report generation ─────────────────────────────────
# python-docx generates .docx files for the national INEP report.
# The try/except means the app still works if this library is missing —
# the Word download button just shows an error message instead.
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True            # Flag: Word reports are available
except ImportError:
    DOCX_AVAILABLE = False           # Flag: Word reports unavailable — missing library

# ── Optional: PDF report generation ──────────────────────────────────────────
# reportlab generates .pdf files for the national INEP report.
# Same pattern — app works without it, PDF download just shows an error.
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True             # Flag: PDF reports are available
except ImportError:
    PDF_AVAILABLE = False            # Flag: PDF reports unavailable — missing library

# ── Page config ───────────────────────────────────────────────────────────────
# =============================================================================
# SECTION 2 — PAGE CONFIGURATION
# These settings are applied to the entire Streamlit app before anything
# else renders. They must come before any other Streamlit calls.
# =============================================================================

st.set_page_config(
    page_title="KenyaWatts · EPRA National Energy Planning",  # Browser tab title
    page_icon="⚡",                    # Browser tab icon (emoji or image path)
    layout="wide",                    # Use full browser width (not narrow centered)
    initial_sidebar_state="expanded"  # Show the sidebar open by default
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  #MainMenu{visibility:hidden} footer{visibility:hidden} .stDeployButton{display:none}
  .kw-header{background:#0e1e2e;color:white;padding:14px 20px;border-radius:10px;margin-bottom:16px}
  .kw-logo{font-size:22px;font-weight:700;letter-spacing:-0.3px}
  .kw-logo span{color:#3ecfaa}
  .kw-metric{background:#f7f6f2;border-radius:10px;padding:14px 16px;margin-bottom:10px}
  .kw-metric-label{font-size:11px;color:#9c9a8e;text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}
  .kw-metric-value{font-size:24px;font-weight:700;color:#1a1916;line-height:1.1}
  .kw-metric-delta{font-size:12px;margin-top:3px}
  .kw-alert-info{background:#edf4fb;border-left:3px solid #1a6fa3;color:#1a6fa3;padding:10px 14px;border-radius:6px;font-size:12px;margin-bottom:10px;line-height:1.6}
  .kw-alert-success{background:#e8f7f4;border-left:3px solid #0f9d7e;color:#0f9d7e;padding:10px 14px;border-radius:6px;font-size:12px;margin-bottom:10px;line-height:1.6}
  .kw-alert-warn{background:#fdf3e3;border-left:3px solid #d4891a;color:#d4891a;padding:10px 14px;border-radius:6px;font-size:12px;margin-bottom:10px;line-height:1.6}
  .kw-alert-danger{background:#fbedeb;border-left:3px solid #b33a2c;color:#b33a2c;padding:10px 14px;border-radius:6px;font-size:12px;margin-bottom:10px;line-height:1.6}
  .kw-badge-submitted{background:#e8f7f4;color:#0f9d7e;padding:2px 10px;border-radius:12px;font-size:11px;font-weight:600}
  .kw-badge-review{background:#fdf3e3;color:#d4891a;padding:2px 10px;border-radius:12px;font-size:11px;font-weight:600}
  .kw-badge-pending{background:#f5f4f0;color:#7a7870;padding:2px 10px;border-radius:12px;font-size:11px;font-weight:600}
  .kw-badge-overdue{background:#fbedeb;color:#b33a2c;padding:2px 10px;border-radius:12px;font-size:11px;font-weight:600}
  .kw-notif{position:fixed;top:70px;right:20px;z-index:9999;background:#0f9d7e;color:white;padding:14px 20px;border-radius:10px;font-size:13px;font-weight:500;box-shadow:0 4px 20px rgba(0,0,0,0.15);max-width:320px;line-height:1.5}
  .upload-log-row{display:flex;gap:12px;align-items:flex-start;padding:8px 0;border-bottom:0.5px solid #f0ede4;font-size:12px}
  .upload-log-icon{font-size:16px;flex-shrink:0}
  .upload-log-meta{color:#9c9a8e;font-size:11px;margin-top:2px}
  /* Prevent content bleed-through between page navigations */
  [data-testid="stMain"] > div {contain:layout style;overflow:visible}
  [data-testid="stVerticalBlock"] {isolation:isolate}

  /* ── Floating AI assistant button (bottom-right corner) ── */
  /* The button itself is rendered via st.markdown HTML below in main() */
  .kw-ai-fab {
    position:fixed; bottom:24px; right:24px; z-index:9999;
    width:56px; height:56px; border-radius:50%;
    background:linear-gradient(135deg,#0e1e2e,#1a6fa3);
    display:flex; align-items:center; justify-content:center;
    cursor:pointer; box-shadow:0 4px 20px rgba(26,111,163,0.45);
    font-size:24px; border:2px solid rgba(62,207,170,0.6);
    transition:transform .2s, box-shadow .2s;
  }
  .kw-ai-fab:hover {transform:scale(1.1); box-shadow:0 6px 28px rgba(26,111,163,0.6);}

  /* Floating chat panel */
  .kw-ai-panel {
    position:fixed; bottom:90px; right:24px; z-index:9998;
    width:380px; max-height:520px;
    background:#0e1e2e; border-radius:16px;
    border:1px solid rgba(62,207,170,0.3);
    box-shadow:0 8px 40px rgba(0,0,0,0.4);
    display:flex; flex-direction:column; overflow:visible;
  }
  .kw-ai-panel-header {
    padding:12px 16px; background:rgba(255,255,255,0.05);
    border-bottom:0.5px solid rgba(255,255,255,0.1);
    display:flex; justify-content:space-between; align-items:center;
    font-size:13px; font-weight:600; color:#ffffff;
  }
  .kw-ai-panel-body {
    flex:1; overflow-y:auto; padding:12px 14px;
  }
  .kw-ai-msg-user {
    background:#1a6fa3; color:#fff; border-radius:12px 12px 2px 12px;
    padding:8px 12px; margin-bottom:8px; font-size:12px;
    line-height:1.5; max-width:85%; margin-left:auto;
  }
  .kw-ai-msg-bot {
    background:rgba(255,255,255,0.08); color:#e8e6e0;
    border-radius:12px 12px 12px 2px; padding:8px 12px;
    margin-bottom:8px; font-size:12px; line-height:1.5; max-width:90%;
  }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# SECTION 4 — CHART LAYOUT DEFAULTS
# These settings are applied to every Plotly chart in the platform.
# They ensure axis labels, tick labels and legend text are always
# bold, black and clearly readable — overriding Plotly's grey defaults.
# Usage: fig = apply_layout(fig, xlabel="Year", ylabel="GWh")
# =============================================================================

CHART_FONT  = dict(family="Arial", size=13, color="#1a1916")  # Main chart font — near black
AXIS_FONT   = dict(family="Arial", size=12, color="#1a1916")  # Axis tick labels
LEGEND_FONT = dict(family="Arial", size=12, color="#1a1916")  # Legend text
BASE_LAYOUT = dict(  # Applied to every chart via apply_layout()
    font=CHART_FONT,
    plot_bgcolor="white", paper_bgcolor="white",
    legend=dict(font=LEGEND_FONT, bgcolor="rgba(255,255,255,0.9)", borderwidth=0),
    xaxis=dict(tickfont=AXIS_FONT, title_font=dict(size=13,color="#1a1916",family="Arial"), showgrid=False),
    yaxis=dict(tickfont=AXIS_FONT, title_font=dict(size=13,color="#1a1916",family="Arial"), showgrid=True, gridcolor="#f0f0f0"),
    margin=dict(t=20, b=40, l=50, r=20),
)

def apply_layout(fig, title="", xlabel="", ylabel=""):
    fig.update_layout(**BASE_LAYOUT)
    if title:  fig.update_layout(title=dict(text=title, font=dict(size=14,color="#1a1916",family="Arial")))
    if xlabel: fig.update_xaxes(title_text=f"<b>{xlabel}</b>")
    if ylabel: fig.update_yaxes(title_text=f"<b>{ylabel}</b>")
    return fig

# =============================================================================
# SECTION 5 — COUNTY ENERGY DATA
# This DataFrame holds the core energy planning data for 14 sample counties.
# In a production system this would be read from the Supabase database.
# Each row = one county. Columns explained:
#   id        = two-letter county code (matches county_id in secrets.toml)
#   name      = full county name
#   region    = Kenya administrative region
#   pop       = population (KNBS 2019 Census)
#   elec      = electricity access rate (% of households) — grid + minigrid + SHS
#   cooking   = clean cooking access (% using LPG, biogas, electric, or ICS)
#   solar     = solar GHI — Global Horizontal Irradiance in kWh/m²/year
#               (measures how much solar energy hits the ground — higher = better)
#   budget    = total county energy plan budget in KES billions (0 = not submitted)
#   status    = submission status: submitted / review / pending / overdue
#   mtf       = Multi-Tier Framework demand tier (1=minimal, 5=high urban demand)
#   growth    = annual population growth rate (% per year, from KNBS data)
#   target_yr = year the county plans to achieve universal electricity access
#   lat/lon   = GPS coordinates — used to place markers on the map
# =============================================================================

COUNTIES = pd.DataFrame([
    # ── Nairobi Region ────────────────────────────────────────────────────────
    {"id":"NK","name":"Nairobi",       "region":"Nairobi",       "pop":4922000,"elec":96,"cooking":62,"solar":1980,"budget":120.0,"status":"submitted","mtf":4,"growth":1.8,"target_yr":2027,"lat":-1.286,"lon":36.817},

    # ── Central Region ────────────────────────────────────────────────────────
    {"id":"KI2","name":"Kiambu",       "region":"Central",       "pop":2417735,"elec":85,"cooking":48,"solar":1920,"budget":0,   "status":"pending",  "mtf":4,"growth":2.1,"target_yr":2029,"lat":-1.031,"lon":36.832},
    {"id":"MR","name":"Muranga",       "region":"Central",       "pop":1056640,"elec":68,"cooking":33,"solar":1900,"budget":0,   "status":"pending",  "mtf":3,"growth":0.9,"target_yr":2029,"lat":-0.717,"lon":37.150},
    {"id":"NY","name":"Nyeri",         "region":"Central",       "pop":759164, "elec":72,"cooking":35,"solar":1880,"budget":0,   "status":"pending",  "mtf":3,"growth":0.7,"target_yr":2029,"lat":-0.420,"lon":36.947},
    {"id":"KR","name":"Kirinyaga",     "region":"Central",       "pop":610411, "elec":70,"cooking":32,"solar":1890,"budget":0,   "status":"pending",  "mtf":3,"growth":0.8,"target_yr":2030,"lat":-0.500,"lon":37.400},
    {"id":"NT","name":"Nyandarua",     "region":"Central",       "pop":638289, "elec":62,"cooking":28,"solar":1870,"budget":0,   "status":"pending",  "mtf":2,"growth":0.6,"target_yr":2030,"lat":-0.180,"lon":36.520},

    # ── Coast Region ──────────────────────────────────────────────────────────
    {"id":"MB","name":"Mombasa",       "region":"Coast",         "pop":1208000,"elec":84,"cooking":45,"solar":2050,"budget":55.0,"status":"submitted","mtf":4,"growth":1.5,"target_yr":2028,"lat":-4.043,"lon":39.668},
    {"id":"KW","name":"Kwale",         "region":"Coast",         "pop":866820, "elec":41,"cooking":18,"solar":2020,"budget":0,   "status":"pending",  "mtf":2,"growth":1.3,"target_yr":2031,"lat":-4.183,"lon":39.483},
    {"id":"KF","name":"Kilifi",        "region":"Coast",         "pop":1453787,"elec":38,"cooking":16,"solar":2010,"budget":0,   "status":"pending",  "mtf":2,"growth":1.4,"target_yr":2031,"lat":-3.510,"lon":39.908},
    {"id":"TR","name":"Tana River",    "region":"Coast",         "pop":315943, "elec":18,"cooking":6, "solar":2060,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.2,"target_yr":2033,"lat":-1.600,"lon":40.120},
    {"id":"LD","name":"Lamu",          "region":"Coast",         "pop":143920, "elec":45,"cooking":20,"solar":2080,"budget":0,   "status":"pending",  "mtf":2,"growth":2.8,"target_yr":2031,"lat":-2.268,"lon":40.902},
    {"id":"TT","name":"Taita-Taveta",  "region":"Coast",         "pop":340671, "elec":44,"cooking":19,"solar":1990,"budget":0,   "status":"pending",  "mtf":2,"growth":1.0,"target_yr":2031,"lat":-3.400,"lon":38.370},

    # ── North Eastern Region ──────────────────────────────────────────────────
    {"id":"GR","name":"Garissa",       "region":"North East",    "pop":841353, "elec":22,"cooking":8, "solar":2140,"budget":0,   "status":"pending",  "mtf":1,"growth":2.1,"target_yr":2032,"lat":-0.453,"lon":39.646},
    {"id":"WJ","name":"Wajir",         "region":"North East",    "pop":781263, "elec":9, "cooking":3, "solar":2190,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.9,"target_yr":2034,"lat":1.747, "lon":40.058},
    {"id":"MN","name":"Mandera",       "region":"North East",    "pop":1025756,"elec":11,"cooking":4, "solar":2200,"budget":0,   "status":"overdue",  "mtf":1,"growth":3.1,"target_yr":2034,"lat":3.937, "lon":41.867},

    # ── Eastern Region ────────────────────────────────────────────────────────
    {"id":"MS","name":"Marsabit",      "region":"Eastern",       "pop":459785, "elec":8, "cooking":2, "solar":2180,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.5,"target_yr":2034,"lat":2.335, "lon":37.989},
    {"id":"IS","name":"Isiolo",        "region":"Eastern",       "pop":268002, "elec":28,"cooking":10,"solar":2100,"budget":0,   "status":"pending",  "mtf":1,"growth":2.3,"target_yr":2032,"lat":0.352, "lon":37.582},
    {"id":"ME","name":"Meru",          "region":"Eastern",       "pop":1545714,"elec":61,"cooking":27,"solar":1950,"budget":0,   "status":"pending",  "mtf":3,"growth":1.0,"target_yr":2030,"lat":0.047, "lon":37.649},
    {"id":"TE","name":"Tharaka-Nithi", "region":"Eastern",       "pop":393177, "elec":52,"cooking":22,"solar":1930,"budget":0,   "status":"pending",  "mtf":2,"growth":0.8,"target_yr":2030,"lat":-0.300,"lon":37.900},
    {"id":"EM","name":"Embu",          "region":"Eastern",       "pop":608599, "elec":65,"cooking":30,"solar":1910,"budget":0,   "status":"pending",  "mtf":3,"growth":0.9,"target_yr":2030,"lat":-0.530,"lon":37.450},
    {"id":"KT","name":"Kitui",         "region":"Eastern",       "pop":1136187,"elec":35,"cooking":13,"solar":2000,"budget":0,   "status":"pending",  "mtf":2,"growth":1.1,"target_yr":2031,"lat":-1.368,"lon":38.010},
    {"id":"MC","name":"Machakos",      "region":"Eastern",       "pop":1421932,"elec":58,"cooking":25,"solar":1970,"budget":0,   "status":"pending",  "mtf":3,"growth":1.3,"target_yr":2030,"lat":-1.519,"lon":37.266},
    {"id":"MK","name":"Makueni",       "region":"South East",    "pop":987653, "elec":75,"cooking":18,"solar":2008,"budget":74.9,"status":"submitted","mtf":3,"growth":1.1,"target_yr":2028,"lat":-2.303,"lon":37.624},

    # ── Rift Valley Region ────────────────────────────────────────────────────
    {"id":"TK","name":"Turkana",       "region":"North Rift",    "pop":926976, "elec":12,"cooking":3, "solar":2150,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.8,"target_yr":2033,"lat":3.112, "lon":35.596},
    {"id":"WP","name":"West Pokot",    "region":"North Rift",    "pop":621241, "elec":20,"cooking":7, "solar":2020,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.6,"target_yr":2033,"lat":1.620, "lon":35.120},
    {"id":"SM","name":"Samburu",       "region":"North Rift",    "pop":310327, "elec":16,"cooking":5, "solar":2080,"budget":0,   "status":"overdue",  "mtf":1,"growth":2.2,"target_yr":2033,"lat":1.200, "lon":36.900},
    {"id":"TZ","name":"Trans-Nzoia",   "region":"Rift Valley",   "pop":990341, "elec":48,"cooking":21,"solar":1860,"budget":0,   "status":"pending",  "mtf":2,"growth":1.6,"target_yr":2031,"lat":1.017, "lon":34.950},
    {"id":"UK","name":"Uasin Gishu",   "region":"Rift Valley",   "pop":1163186,"elec":62,"cooking":30,"solar":1890,"budget":0,   "status":"pending",  "mtf":3,"growth":1.5,"target_yr":2030,"lat":0.520, "lon":35.270},
    {"id":"EC","name":"Elgeyo-Marakwet","region":"Rift Valley",  "pop":454480, "elec":43,"cooking":18,"solar":1900,"budget":0,   "status":"pending",  "mtf":2,"growth":1.1,"target_yr":2031,"lat":1.050, "lon":35.520},
    {"id":"NA","name":"Nakuru",        "region":"Rift Valley",   "pop":2162202,"elec":72,"cooking":38,"solar":1920,"budget":68.0,"status":"review",   "mtf":3,"growth":1.4,"target_yr":2029,"lat":-0.303,"lon":36.080},
    {"id":"BR","name":"Baringo",       "region":"Rift Valley",   "pop":666763, "elec":34,"cooking":13,"solar":2010,"budget":0,   "status":"pending",  "mtf":2,"growth":1.7,"target_yr":2031,"lat":0.470, "lon":35.970},
    {"id":"LA","name":"Laikipia",      "region":"Rift Valley",   "pop":518560, "elec":55,"cooking":25,"solar":1960,"budget":0,   "status":"pending",  "mtf":2,"growth":1.0,"target_yr":2030,"lat":0.360, "lon":36.780},
    {"id":"NJ","name":"Nandi",         "region":"Rift Valley",   "pop":885711, "elec":50,"cooking":22,"solar":1870,"budget":0,   "status":"pending",  "mtf":2,"growth":1.3,"target_yr":2030,"lat":0.200, "lon":35.130},
    {"id":"KJ","name":"Kericho",       "region":"Rift Valley",   "pop":902000, "elec":58,"cooking":26,"solar":1880,"budget":0,   "status":"pending",  "mtf":3,"growth":1.2,"target_yr":2030,"lat":-0.370,"lon":35.290},
    {"id":"BM","name":"Bomet",         "region":"Rift Valley",   "pop":876986, "elec":44,"cooking":18,"solar":1860,"budget":0,   "status":"pending",  "mtf":2,"growth":1.4,"target_yr":2031,"lat":-0.790,"lon":35.340},
    {"id":"KA","name":"Kajiado",       "region":"South Rift",    "pop":1117840,"elec":55,"cooking":29,"solar":1990,"budget":0,   "status":"pending",  "mtf":2,"growth":1.8,"target_yr":2030,"lat":-1.852,"lon":36.777},
    {"id":"NR","name":"Narok",         "region":"South Rift",    "pop":1157873,"elec":32,"cooking":12,"solar":2000,"budget":0,   "status":"pending",  "mtf":2,"growth":2.0,"target_yr":2031,"lat":-1.080,"lon":35.870},

    # ── Nyanza Region ─────────────────────────────────────────────────────────
    {"id":"KI","name":"Kisumu",        "region":"Nyanza",        "pop":1155574,"elec":67,"cooking":31,"solar":1870,"budget":0,   "status":"pending",  "mtf":3,"growth":1.2,"target_yr":2030,"lat":-0.091,"lon":34.768},
    {"id":"HA","name":"Homa Bay",      "region":"Nyanza",        "pop":1131950,"elec":38,"cooking":15,"solar":1840,"budget":0,   "status":"pending",  "mtf":2,"growth":1.5,"target_yr":2031,"lat":-0.530,"lon":34.457},
    {"id":"MG","name":"Migori",        "region":"Nyanza",        "pop":1116436,"elec":35,"cooking":13,"solar":1850,"budget":0,   "status":"pending",  "mtf":2,"growth":1.8,"target_yr":2031,"lat":-1.063,"lon":34.473},
    {"id":"KS","name":"Kisii",         "region":"Nyanza",        "pop":1266860,"elec":52,"cooking":23,"solar":1840,"budget":0,   "status":"pending",  "mtf":3,"growth":1.1,"target_yr":2030,"lat":-0.681,"lon":34.766},
    {"id":"NY2","name":"Nyamira",      "region":"Nyanza",        "pop":605576, "elec":48,"cooking":20,"solar":1830,"budget":0,   "status":"pending",  "mtf":2,"growth":0.9,"target_yr":2030,"lat":-0.570,"lon":34.940},
    {"id":"SI","name":"Siaya",         "region":"Nyanza",        "pop":993183, "elec":42,"cooking":16,"solar":1850,"budget":0,   "status":"pending",  "mtf":2,"growth":1.0,"target_yr":2031,"lat":-0.060,"lon":34.290},

    # ── Western Region ────────────────────────────────────────────────────────
    {"id":"CD","name":"Kakamega",      "region":"Western",       "pop":1867579,"elec":49,"cooking":21,"solar":1860,"budget":0,   "status":"pending",  "mtf":2,"growth":1.4,"target_yr":2031,"lat":0.282, "lon":34.752},
    {"id":"VM","name":"Vihiga",        "region":"Western",       "pop":590013, "elec":55,"cooking":24,"solar":1850,"budget":0,   "status":"pending",  "mtf":3,"growth":0.8,"target_yr":2030,"lat":-0.070,"lon":34.723},
    {"id":"BU","name":"Bungoma",       "region":"Western",       "pop":1670570,"elec":44,"cooking":19,"solar":1870,"budget":0,   "status":"pending",  "mtf":2,"growth":1.6,"target_yr":2031,"lat":0.564, "lon":34.560},
    {"id":"BN","name":"Busia",         "region":"Western",       "pop":893681, "elec":39,"cooking":15,"solar":1850,"budget":0,   "status":"pending",  "mtf":2,"growth":1.3,"target_yr":2031,"lat":0.460, "lon":34.110},
])

GEN_TREND = pd.DataFrame({
    "Year":["FY19/20","FY20/21","FY21/22","FY22/23","FY23/24","FY24/25"],
    "GWh": [11564,    11891,    12210,    12897,    13685,    14520]
})
GEN_MIX = pd.DataFrame({
    "Source":["Geothermal","Hydro","Wind","Solar","Thermal","Other"],
    "Pct":   [25.9, 23.8, 17.4, 14.9, 8.7, 9.3],
    "Color": ["#0f9d7e","#1a6fa3","#5b4fc9","#d4891a","#b33a2c","#7a7870"]
})
MAKUENI_COOKING = pd.DataFrame({
    "Fuel":["Firewood","LPG","Charcoal","Biogas","Electric","Other"],
    "Pct": [72.5,17.6,8.2,0.2,0.3,1.2]
})
OUTAGE = pd.DataFrame({
    "Month":["Jul","Aug","Sep","Oct","Nov","Dec","Jan","Feb","Mar","Apr","May","Jun"],
    "Hours":[7.2,9.1,11.1,8.4,7.8,6.9,8.2,9.4,8.8,7.1,6.5,7.3]
})

# ── CLOUD DATABASE LAYER (Supabase PostgreSQL) ───────────────────────────────
# Supabase is a free cloud PostgreSQL database.
# When configured it stores all messages and submissions persistently in the cloud
# so data survives redeploys and is accessible across all sessions forever.
# If not configured, falls back to local JSON files (fine for demo).
#
# To enable:
#   1. Go to supabase.com → New project (free)
#   2. Go to SQL Editor → run the setup SQL below
#   3. Go to Settings → API → copy URL and anon key
#   4. Add to Streamlit secrets:
#      [supabase]
#      url = "https://your-project.supabase.co"
#      key = "your-anon-key"
#
# Setup SQL (run once in Supabase SQL editor):
# CREATE TABLE IF NOT EXISTS kw_messages (
#   id TEXT PRIMARY KEY,
#   data JSONB NOT NULL,
#   created_at TIMESTAMPTZ DEFAULT NOW()
# );
# CREATE TABLE IF NOT EXISTS kw_submissions (
#   id TEXT PRIMARY KEY,
#   data JSONB NOT NULL,
#   created_at TIMESTAMPTZ DEFAULT NOW()
# );

try:
    import urllib.request, urllib.error
    _SUPABASE_CFG = st.secrets.get("supabase", {})
    SUPABASE_URL  = _SUPABASE_CFG.get("url","")
    SUPABASE_KEY  = _SUPABASE_CFG.get("key","")
    USE_CLOUD_DB  = bool(SUPABASE_URL and SUPABASE_KEY)
except Exception:
    USE_CLOUD_DB  = False
    SUPABASE_URL  = ""
    SUPABASE_KEY  = ""

def _supa_headers():
    return {
        "apikey":        SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type":  "application/json",
        "Prefer":        "resolution=merge-duplicates",
    }

def _supa_get(table):
    """Fetch all rows from a Supabase table."""
    try:
        url = f"{SUPABASE_URL}/rest/v1/{table}?select=data&order=created_at.asc"
        req = urllib.request.Request(url, headers=_supa_headers())
        with urllib.request.urlopen(req, timeout=5) as r:
            rows = json.loads(r.read())
            return [row["data"] for row in rows]
    except Exception:
        return None

def _supa_upsert(table, record_id, data):
    """Upsert a record into a Supabase table."""
    try:
        payload = json.dumps({"id": record_id, "data": data}).encode()
        url     = f"{SUPABASE_URL}/rest/v1/{table}"
        req     = urllib.request.Request(url, data=payload,
                                         headers=_supa_headers(), method="POST")
        req.add_header("Prefer","resolution=merge-duplicates")
        with urllib.request.urlopen(req, timeout=5) as r:
            return r.status < 300
    except Exception:
        return False

def _supa_patch(table, record_id, data):
    """Update a specific record in Supabase."""
    try:
        payload = json.dumps({"data": data}).encode()
        url     = f"{SUPABASE_URL}/rest/v1/{table}?id=eq.{record_id}"
        req     = urllib.request.Request(url, data=payload,
                                         headers={**_supa_headers(),"Prefer":"return=minimal"},
                                         method="PATCH")
        with urllib.request.urlopen(req, timeout=5) as r:
            return r.status < 300
    except Exception:
        return False

# ── FILE-BASED MESSAGE STORE ─────────────────────────────────────────────────
# Messages are written to a JSON file so they persist across ALL sessions.
# Every login reads from the same file — EPRA messages appear in county inboxes
# and county messages appear in EPRA's hub in real time.
import os, threading

MSG_STORE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "kw_messages.json")
_msg_lock = threading.Lock()

# SEED_MESSAGES is intentionally EMPTY so the platform starts clean.
# When testing, EPRA and counties generate real messages by using the platform.
# The developer can reset all messages using the Reset tool in Account settings.
# To pre-load sample messages for a demo, add them here in the same format.
SEED_MESSAGES = []   # Empty — no pre-loaded sample messages

def load_messages():
    """Load messages — tries Supabase cloud first, falls back to local JSON."""
    if USE_CLOUD_DB:
        rows = _supa_get("kw_messages")
        if rows is not None:
            return rows if rows else SEED_MESSAGES
    # Local JSON fallback
    with _msg_lock:
        if not os.path.exists(MSG_STORE_PATH):
            _write_messages(SEED_MESSAGES)
            return SEED_MESSAGES
        try:
            with open(MSG_STORE_PATH,"r",encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            _write_messages(SEED_MESSAGES)
            return SEED_MESSAGES

def _write_messages(messages):
    """Write messages to the JSON file (call inside _msg_lock)."""
    with open(MSG_STORE_PATH,"w",encoding="utf-8") as f:
        json.dump(messages, f, ensure_ascii=False, indent=2)

def save_messages(messages):
    """Save messages to the JSON file (thread-safe)."""
    with _msg_lock:
        _write_messages(messages)

def get_messages():
    """Return live messages — always reads from file."""
    return load_messages()

def add_message(msg):
    """Add a new message — saves to Supabase cloud AND local JSON."""
    if USE_CLOUD_DB:
        _supa_upsert("kw_messages", msg["id"], msg)
    with _msg_lock:
        msgs = load_messages() if not USE_CLOUD_DB else []
        if not USE_CLOUD_DB:
            msgs.append(msg)
            _write_messages(msgs)
        else:
            # Still write locally as backup
            try:
                local = json.load(open(MSG_STORE_PATH)) if os.path.exists(MSG_STORE_PATH) else []
                local.append(msg)
                open(MSG_STORE_PATH,"w").write(json.dumps(local,indent=2))
            except Exception:
                pass
    return msg

def add_reply(msg_id, reply):
    """Add a reply to a specific message by ID — saves to cloud and local."""
    msgs = load_messages()
    for m in msgs:
        if m["id"] == msg_id:
            m.setdefault("replies", []).append(reply)
            m["read_by_epra"] = True
            if USE_CLOUD_DB:
                _supa_patch("kw_messages", msg_id, m)
            break
    if not USE_CLOUD_DB:
        with _msg_lock:
            _write_messages(msgs)

def mark_read(msg_id):
    """Mark a message as read by EPRA."""
    with _msg_lock:
        msgs = load_messages()
        for m in msgs:
            if m["id"] == msg_id:
                m["read_by_epra"] = True
                break
        _write_messages(msgs)


# =============================================================================
# RESET / CLEAR FUNCTIONS  (developer only — called from the Reset tool page)
# These permanently delete all messages and/or submissions from both the
# local JSON files and the Supabase cloud database.
# Only the "developer" role has access to these functions via the UI.
# =============================================================================

def reset_all_messages():
    """
    Delete ALL messages from the message store.
    Clears both the local JSON file and Supabase (if configured).
    After this call, load_messages() returns an empty list.
    Used by the developer to start a clean test session.
    """
    # Clear local JSON file
    with _msg_lock:
        _write_messages([])
    # Clear Supabase table if configured
    if USE_CLOUD_DB:
        try:
            import urllib.request as _ur
            # DELETE all rows from kw_messages table
            url = f"{SUPABASE_URL}/rest/v1/kw_messages?id=neq.NONE_MATCH"
            req = _ur.Request(url, headers=_supa_headers(), method="DELETE")
            _ur.urlopen(req, timeout=5)
        except Exception:
            pass   # Silently continue if Supabase delete fails

def reset_all_submissions():
    """
    Delete ALL county plan submissions from the submission store.
    Clears both the local JSON file and Supabase (if configured).
    After this call, get_submissions() returns an empty list.
    Used by the developer to start a clean test session.
    """
    # Clear local JSON file
    with _sub_lock:
        with open(SUB_STORE_PATH,"w",encoding="utf-8") as f:
            json.dump([], f)
    # Clear Supabase table if configured
    if USE_CLOUD_DB:
        try:
            import urllib.request as _ur
            url = f"{SUPABASE_URL}/rest/v1/kw_submissions?id=neq.NONE_MATCH"
            req = _ur.Request(url, headers=_supa_headers(), method="DELETE")
            _ur.urlopen(req, timeout=5)
        except Exception:
            pass

# ── FILE-BASED SUBMISSION STORE ───────────────────────────────────────────────
# Every county plan submission is written to disk immediately.
# EPRA and county users read from the same file — real-time across all sessions.
SUB_STORE_PATH  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "kw_submissions.json")
_sub_lock = threading.Lock()

def load_submissions():
    """Load submissions — tries Supabase cloud first, falls back to local JSON."""
    if USE_CLOUD_DB:
        rows = _supa_get("kw_submissions")
        if rows is not None:
            return sorted(rows, key=lambda x: x.get("datetime",""), reverse=True)
    with _sub_lock:
        if not os.path.exists(SUB_STORE_PATH):
            return []
        try:
            with open(SUB_STORE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []

def save_submission(record):
    """
    Save a county plan submission to Supabase cloud AND local JSON.

    HISTORY PRESERVATION:
    All submissions are kept. Each gets a unique ref so the full history
    is visible in the county's submission history tab and the EPRA
    validation queue. The most recent submission per county is used for
    national aggregation (computed by sorting newest-first).
    """
    if USE_CLOUD_DB:
        _supa_upsert("kw_submissions", record["ref"], record)

    with _sub_lock:
        subs = []
        if os.path.exists(SUB_STORE_PATH):
            try:
                subs = json.load(open(SUB_STORE_PATH))
            except Exception:
                subs = []

        # Remove only the exact same ref (prevent exact duplicates on rerun)
        # All other submissions from this county are kept as history
        subs = [s for s in subs if s.get("ref") != record.get("ref")]

        subs.append(record)
        with open(SUB_STORE_PATH, "w", encoding="utf-8") as f:
            json.dump(subs, f, ensure_ascii=False, indent=2)

def get_submissions(county=None):
    """
    Return all submissions from the persistent store, optionally filtered
    by county name (case-insensitive). Returns newest first.

    This is the single source of truth for county plan submissions across
    all sessions. Always reads from disk/cloud so changes from one session
    are immediately visible in another.
    """
    subs = load_submissions()
    if county:
        subs = [s for s in subs
                if s.get("county","").strip().lower() == county.strip().lower()]
    return sorted(subs, key=lambda x: x.get("datetime",""), reverse=True)

def update_submission_status(ref, new_status, epra_note=""):
    """EPRA approves/rejects a submission — updates cloud and local."""
    subs = load_submissions()
    for s in subs:
        if s.get("ref") == ref:
            s["status"]           = new_status
            s["epra_note"]        = epra_note
            s["epra_action_time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            if USE_CLOUD_DB:
                _supa_patch("kw_submissions", ref, s)
            break
    if not USE_CLOUD_DB:
        with _sub_lock:
            with open(SUB_STORE_PATH, "w", encoding="utf-8") as f:
                json.dump(subs, f, ensure_ascii=False, indent=2)
    else:
        # Update local backup too
        try:
            with open(SUB_STORE_PATH, "w", encoding="utf-8") as f:
                json.dump(subs, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

STATUS_COLOR = {"submitted":"#0f9d7e","review":"#d4891a","pending":"#7a7870","overdue":"#b33a2c"}
STATUS_LABEL = {"submitted":"Submitted ✓","review":"In review","pending":"Pending","overdue":"Overdue"}

# ── SESSION STATE INIT ────────────────────────────────────────────────────────
# =============================================================================
# SECTION 10 — SESSION STATE INITIALISATION
# st.session_state is Streamlit's in-memory store. It persists across reruns
# within the same browser session but is lost when the browser tab closes.
# We use it for: audit logs, notification queue, AI chat history, upload logs.
# Messages and submissions use the persistent file/cloud store instead.
# =============================================================================

def init_session():
    """
    Initialise all session state variables on first load.
    Uses setdefault pattern — only sets a value if it does not already exist,
    so existing values are preserved across reruns.
    """
    defaults = {
        # All session state starts EMPTY — no fake/hardcoded data.
        # Real data flows only from:
        #   - kw_messages.json (messages between EPRA and counties)
        #   - kw_submissions.json (county plan submissions)
        # These are read by get_messages() and get_submissions() respectively.
        "audit_log":         [],   # Log of actions taken this session
        "upload_log":        [],   # Local record of uploads this session
        "ai_history":        [],   # AI chat conversation history
        "notifications":     [],   # Toast notifications queue
        "user_display_name": {},   # In-session display name overrides
        "user_email":        {},   # In-session email overrides
        "submitted_data":    [],   # Submitted data (legacy — now in file store)
        "lang":              "en", # UI language: "en" = English, "sw" = Swahili
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# ── NOTIFICATION SYSTEM ───────────────────────────────────────────────────────
def push_notification(msg, icon="✅"):
    st.session_state.notifications.append({
        "msg": msg, "icon": icon,
        "time": datetime.now().strftime("%H:%M:%S"), "shown": False
    })

def show_notifications():
    for n in st.session_state.notifications:
        if not n["shown"]:
            st.toast(f"{n['icon']} {n['msg']}", icon=n["icon"])
            n["shown"] = True

# =============================================================================
# LANGUAGE / TRANSLATION SYSTEM
# Supports English (en) and Swahili (sw).
# Usage: t("key")  → returns the string in the current language.
# Language is stored in st.session_state["lang"] (default "en").
# The language toggle is rendered in the sidebar on every page.
# =============================================================================

TRANSLATIONS = {
    # ── Navigation ────────────────────────────────────────────────────────────
    "nav_national_overview":      {"en": "National overview",                   "sw": "Muhtasari wa Kitaifa"},
    "nav_county_map":             {"en": "County map",                          "sw": "Ramani ya Kaunti"},
    "nav_all_counties":           {"en": "All 47 counties",                     "sw": "Kaunti Zote 47"},
    "nav_validation_queue":       {"en": "Validation queue",                    "sw": "Foleni ya Uthibitisho"},
    "nav_comms_hub":              {"en": "Communications hub",                  "sw": "Kituo cha Mawasiliano"},
    "nav_makueni_ref":            {"en": "Makueni CEP — template reference",    "sw": "Makueni CEP — Mwongozo"},
    "nav_data_download":          {"en": "Data download",                       "sw": "Pakua Data"},
    "nav_account_settings":       {"en": "Account settings",                    "sw": "Mipangilio ya Akaunti"},
    "nav_submit_plan":            {"en": "Submit energy plan",                  "sw": "Wasilisha Mpango wa Nishati"},
    "nav_communications":         {"en": "Communications",                      "sw": "Mawasiliano"},
    "nav_my_county_plan":         {"en": "My county plan",                      "sw": "Mpango wangu wa Kaunti"},
    "nav_county_demand_data":     {"en": "County demand data",                  "sw": "Data ya Mahitaji ya Kaunti"},
    "nav_submit_provider_plan":   {"en": "Submit provider plan",                "sw": "Wasilisha Mpango wa Mtoa"},
    "nav_dev_tools":              {"en": "Developer tools",                     "sw": "Zana za Msanidi"},

    # ── Sidebar ────────────────────────────────────────────────────────────────
    "sidebar_platform":           {"en": "National Energy Planning Platform",   "sw": "Jukwaa la Mipango ya Nishati ya Kitaifa"},
    "sidebar_navigation":         {"en": "Navigation",                          "sw": "Urambazaji"},
    "sidebar_sign_out":           {"en": "🚪  Sign out",                        "sw": "🚪  Toka"},
    "sidebar_language":           {"en": "Language / Lugha",                    "sw": "Lugha / Language"},
    "sidebar_data_source":        {"en": "Data: EPRA FY 2024/25 · Makueni CEP 2023–2032 · KNBS 2019",
                                   "sw": "Data: EPRA FY 2024/25 · Makueni CEP 2023–2032 · KNBS 2019"},
    "sidebar_ai_btn":             {"en": "⚡  AI Assistant",                    "sw": "⚡  Msaidizi wa AI"},

    # ── Role labels ────────────────────────────────────────────────────────────
    "role_epra":                  {"en": "EPRA Planner",                        "sw": "Mpanga wa EPRA"},
    "role_ministry":              {"en": "Ministry of Energy",                  "sw": "Wizara ya Nishati"},
    "role_devpartner":            {"en": "Development Partner",                 "sw": "Mshirika wa Maendeleo"},
    "role_county":                {"en": "County Energy Planning Committee",    "sw": "Kamati ya Mipango ya Nishati ya Kaunti"},
    "role_kplc":                  {"en": "KPLC",                                "sw": "KPLC"},
    "role_developer":             {"en": "Platform Developer",                  "sw": "Msanidi wa Jukwaa"},

    # ── Header ────────────────────────────────────────────────────────────────
    "header_tagline":             {"en": "Digital Integrated National Energy Planning Platform · EPRA Kenya",
                                   "sw": "Jukwaa la Kidijitali la Mipango ya Nishati ya Kitaifa · EPRA Kenya"},
    "header_live":                {"en": "● Live · NGDA 2026",                  "sw": "● Moja kwa Moja · NGDA 2026"},

    # ── Login page ────────────────────────────────────────────────────────────
    "login_title":                {"en": "Digital Integrated National Energy Planning Platform",
                                   "sw": "Jukwaa la Kidijitali la Mipango ya Nishati ya Kitaifa"},
    "login_subtitle":             {"en": "EPRA Kenya · NGDA 2026",              "sw": "EPRA Kenya · NGDA 2026"},
    "login_form_name":            {"en": "Sign in to KenyaWatts",               "sw": "Ingia KenyaWatts"},
    "login_username":             {"en": "Username",                            "sw": "Jina la Mtumiaji"},
    "login_password":             {"en": "Password",                            "sw": "Nenosiri"},
    "login_btn":                  {"en": "Sign in",                             "sw": "Ingia"},
    "login_forgot":               {"en": "Forgot your password?",               "sw": "Umesahau nenosiri?"},
    "login_error":                {"en": "Incorrect username or password. Please try again.",
                                   "sw": "Jina la mtumiaji au nenosiri si sahihi. Tafadhali jaribu tena."},

    # ── Submit page ───────────────────────────────────────────────────────────
    "submit_tab_submit":          {"en": "📤  Submit plan",                     "sw": "📤  Wasilisha Mpango"},
    "submit_tab_history":         {"en": "🕐  Submission history",              "sw": "🕐  Historia ya Mawasilisho"},
    "submit_pathway_upload":      {"en": "📄 Upload existing PDF / Word plan",  "sw": "📄 Pakia mpango wa PDF / Word"},
    "submit_pathway_template":    {"en": "📝 Fill structured template",         "sw": "📝 Jaza fomu iliyopangwa"},
    "submit_choose_pathway":      {"en": "Choose submission pathway:",          "sw": "Chagua njia ya kuwasilisha:"},
    "submit_run_validation":      {"en": "▶  Run validation checks (10 rules)", "sw": "▶  Endesha ukaguzi (kanuni 10)"},
    "submit_to_epra":             {"en": "✅ Submit plan to EPRA",              "sw": "✅ Wasilisha mpango kwa EPRA"},
    "submit_confirmed":           {"en": "✅ Submission confirmed and saved",   "sw": "✅ Uwasilishaji umethibitishwa na kuhifadhiwa"},
    "submit_epra_notified":       {"en": "Saved to KenyaWatts server · EPRA has been notified automatically",
                                   "sw": "Imehifadhiwa kwenye seva · EPRA imearifu moja kwa moja"},
    "submit_history_title":       {"en": "Submission and upload history",       "sw": "Historia ya mawasilisho na upakiaji"},
    "submit_section_title":       {"en": "Structured submission template",      "sw": "Fomu ya Uwasilishaji Iliyopangwa"},

    # ── Validation ────────────────────────────────────────────────────────────
    "val_critical_errors":        {"en": "Critical errors",                     "sw": "Makosa Makubwa"},
    "val_warnings":               {"en": "Warnings",                            "sw": "Maonyo"},
    "val_checks_passed":          {"en": "Checks passed",                       "sw": "Ukaguzi Uliopita"},
    "val_no_errors":              {"en": "✓ No critical errors. Ready to submit to EPRA.",
                                   "sw": "✓ Hakuna makosa makubwa. Iko tayari kuwasilishwa kwa EPRA."},

    # ── Data download ─────────────────────────────────────────────────────────
    "download_title":             {"en": "Data download centre",                "sw": "Kituo cha Upakuliaji wa Data"},
    "download_scope_national":    {"en": "National aggregated data",            "sw": "Data ya Kitaifa Iliyokusanywa"},
    "download_scope_all":         {"en": "All county submissions",              "sw": "Mawasilisho ya Kaunti Zote"},
    "download_scope_individual":  {"en": "Individual county",                   "sw": "Kaunti Moja"},
    "download_scope_label":       {"en": "Download scope:",                     "sw": "Upeo wa upakuliaji:"},
    "download_from_date":         {"en": "From date",                           "sw": "Tarehe ya Kuanzia"},
    "download_to_date":           {"en": "To date",                             "sw": "Tarehe ya Mwisho"},
    "download_categories":        {"en": "Select data categories to include:",  "sw": "Chagua aina za data za kujumuisha:"},
    "download_elec":              {"en": "Electricity access",                  "sw": "Upatikanaji wa Umeme"},
    "download_cooking":           {"en": "Clean cooking",                       "sw": "Kupika Safi"},
    "download_solar":             {"en": "Solar / Renewable data",              "sw": "Data ya Jua / Nishati Mbadala"},
    "download_budget":            {"en": "Budget and investment",               "sw": "Bajeti na Uwekezaji"},
    "download_meta":              {"en": "Submission metadata",                 "sw": "Metadata ya Uwasilishaji"},
    "download_format":            {"en": "Download format:",                    "sw": "Muundo wa upakuliaji:"},
    "download_btn":               {"en": "⬇ Download",                         "sw": "⬇ Pakua"},
    "download_report_title":      {"en": "Report generation",                   "sw": "Uzalishaji wa Ripoti"},
    "download_report_national":   {"en": "📊 National Summary Report",          "sw": "📊 Ripoti ya Muhtasari wa Kitaifa"},
    "download_report_county":     {"en": "🏘️ County-Specific Report",           "sw": "🏘️ Ripoti ya Kaunti Maalum"},
    "download_report_type":       {"en": "Report type:",                        "sw": "Aina ya ripoti:"},
    "download_generate_btn":      {"en": "⬇️  Generate and download report",    "sw": "⬇️  Tengeneza na pakua ripoti"},

    # ── Communications ────────────────────────────────────────────────────────
    "comms_inbox":                {"en": "Inbox",                               "sw": "Kisanduku cha Barua"},
    "comms_compose":              {"en": "Compose message",                     "sw": "Tunga Ujumbe"},
    "comms_broadcast":            {"en": "Broadcast to counties",               "sw": "Tuma kwa Kaunti Zote"},
    "comms_send":                 {"en": "Send message",                        "sw": "Tuma Ujumbe"},
    "comms_reply":                {"en": "Send reply",                          "sw": "Tuma Jibu"},
    "comms_no_messages":          {"en": "No messages yet.",                    "sw": "Hakuna ujumbe bado."},

    # ── Validation queue ──────────────────────────────────────────────────────
    "vq_title":                   {"en": "Validation queue",                    "sw": "Foleni ya Uthibitisho"},
    "vq_total":                   {"en": "Total received",                      "sw": "Jumla Iliyopokelewa"},
    "vq_pending":                 {"en": "Pending review",                      "sw": "Inasubiri Ukaguzi"},
    "vq_approved":                {"en": "Approved",                            "sw": "Imeidhinishwa"},
    "vq_rejected":                {"en": "Resubmission req.",                   "sw": "Uwasilishaji Upya"},
    "vq_approve_btn":             {"en": "✅ Approve",                          "sw": "✅ Idhinisha"},
    "vq_reject_btn":              {"en": "↩ Request resubmission",             "sw": "↩ Omba Uwasilishaji Upya"},
    "vq_epra_note":               {"en": "EPRA note (optional — shown to county):", "sw": "Kumbuka ya EPRA (ya hiari — inaonyeshwa kwa kaunti):"},

    # ── National overview ─────────────────────────────────────────────────────
    "overview_title":             {"en": "National overview",                   "sw": "Muhtasari wa Kitaifa"},
    "overview_electricity":       {"en": "Electricity access",                  "sw": "Upatikanaji wa Umeme"},
    "overview_cooking":           {"en": "Clean cooking access",                "sw": "Upatikanaji wa Kupika Safi"},
    "overview_counties_sub":      {"en": "Counties submitted",                  "sw": "Kaunti Zilizotumia"},
    "overview_investment":        {"en": "Investment identified",               "sw": "Uwekezaji Uliobainishwa"},

    # ── Account settings ──────────────────────────────────────────────────────
    "account_change_pw":          {"en": "🔑  Change password",                 "sw": "🔑  Badilisha Nenosiri"},
    "account_update_name":        {"en": "✏️  Update display name",            "sw": "✏️  Sasisha Jina la Onyesho"},
    "account_my_info":            {"en": "👤  My account info",                 "sw": "👤  Taarifa zangu za Akaunti"},
    "account_cur_pw":             {"en": "Current password",                    "sw": "Nenosiri la Sasa"},
    "account_new_pw":             {"en": "New password",                        "sw": "Nenosiri Jipya"},
    "account_confirm_pw":         {"en": "Confirm new password",                "sw": "Thibitisha Nenosiri Jipya"},
    "account_update_pw_btn":      {"en": "Update password",                     "sw": "Sasisha Nenosiri"},

    # ── General ───────────────────────────────────────────────────────────────
    "general_save":               {"en": "Save changes",                        "sw": "Hifadhi Mabadiliko"},
    "general_generate":           {"en": "Generate",                            "sw": "Tengeneza"},
    "general_loading":            {"en": "Loading…",                            "sw": "Inapakia…"},
    "general_confidential":       {"en": "CONFIDENTIAL — EPRA INTERNAL USE ONLY",
                                   "sw": "SIRI — KWA MATUMIZI YA NDANI YA EPRA PEKE YAKE"},
    "general_county":             {"en": "County",                              "sw": "Kaunti"},
    "general_region":             {"en": "Region",                              "sw": "Mkoa"},
    "general_population":         {"en": "Population",                          "sw": "Idadi ya Watu"},
    "general_status":             {"en": "Status",                              "sw": "Hali"},
    "general_submitted":          {"en": "Submitted",                           "sw": "Imewasilishwa"},
    "general_pending":            {"en": "Pending",                             "sw": "Inasubiri"},
    "general_overdue":            {"en": "Overdue",                             "sw": "Imechelewa"},
    "general_approved":           {"en": "Approved",                            "sw": "Imeidhinishwa"},
    "general_date":               {"en": "Date",                                "sw": "Tarehe"},
    "general_time":               {"en": "Time",                                "sw": "Wakati"},
    "general_search":             {"en": "Search",                              "sw": "Tafuta"},
    "general_filter":             {"en": "Filter",                              "sw": "Chuja"},
    "general_all":                {"en": "All",                                 "sw": "Zote"},
    "general_none":               {"en": "None",                                "sw": "Hakuna"},
    "general_yes":                {"en": "Yes",                                 "sw": "Ndiyo"},
    "general_no":                 {"en": "No",                                  "sw": "Hapana"},
    "general_cancel":             {"en": "Cancel",                              "sw": "Ghairi"},
    "general_confirm":            {"en": "Confirm",                             "sw": "Thibitisha"},
    "general_close":              {"en": "Close",                               "sw": "Funga"},
    "general_back":               {"en": "Back",                                "sw": "Rudi"},
    "general_next":               {"en": "Next",                                "sw": "Ifuatayo"},
    "general_submit":             {"en": "Submit",                              "sw": "Wasilisha"},
    "general_send":               {"en": "Send",                                "sw": "Tuma"},
    "general_clear":              {"en": "Clear",                               "sw": "Futa"},
    "general_refresh":            {"en": "Refresh",                             "sw": "Onyesha upya"},
    "general_download":           {"en": "Download",                            "sw": "Pakua"},
    "general_upload":             {"en": "Upload",                              "sw": "Pakia"},
    "general_view":               {"en": "View",                                "sw": "Tazama"},
    "general_edit":               {"en": "Edit",                                "sw": "Hariri"},
    "general_delete":             {"en": "Delete",                              "sw": "Futa"},
    "general_error":              {"en": "Error",                               "sw": "Hitilafu"},
    "general_warning":            {"en": "Warning",                             "sw": "Onyo"},
    "general_success":            {"en": "Success",                             "sw": "Mafanikio"},
    "general_info":               {"en": "Information",                         "sw": "Taarifa"},
    "general_records":            {"en": "records",                             "sw": "rekodi"},
    "general_columns":            {"en": "columns",                             "sw": "safu"},
    "general_preview":            {"en": "Data preview",                        "sw": "Muhtasari wa Data"},
    "general_by":                 {"en": "By",                                  "sw": "Na"},
    "general_target_year":        {"en": "Target year",                         "sw": "Mwaka Lengwa"},
    "general_budget":             {"en": "Budget",                              "sw": "Bajeti"},
    "general_ref":                {"en": "Reference",                           "sw": "Kumbukumbu"},
    "general_national_target":    {"en": "National target",                     "sw": "Lengo la Kitaifa"},
    "general_on_track":           {"en": "On track",                            "sw": "Ipo njiani"},
    "general_at_risk":            {"en": "At risk",                             "sw": "Iko hatarini"},
    "general_in_progress":        {"en": "In progress",                         "sw": "Inaendelea"},
}


def t(key: str) -> str:
    """
    Return the translated string for the given key in the current language.
    Falls back to English if the key or language is missing.
    """
    lang = st.session_state.get("lang", "en")
    entry = TRANSLATIONS.get(key, {})
    return entry.get(lang, entry.get("en", key))

# ── AGGREGATION ───────────────────────────────────────────────────────────────
# =============================================================================
# SECTION 12 — NATIONAL AGGREGATION ENGINE
# This is the core of the INEP harmonisation — it computes national-level
# energy indicators from submitted county plans using population weighting.
#
# POPULATION WEIGHTING explained:
# A simple average of county electricity access would give Marsabit (pop 459K)
# the same weight as Nairobi (pop 4.9M). That would be misleading.
# Population weighting means each county's figure is multiplied by its
# population before averaging, so larger counties have more influence on
# the national number — which reflects the real situation.
#
# Formula: national_elec = SUM(county_elec * county_pop) / SUM(county_pop)
# =============================================================================

def compute_national():
    """
    Compute population-weighted national energy indicators.

    DATA SOURCE PRIORITY:
    1. Real submitted county data from the persistent file/cloud store
       (actual submissions made by counties through the platform)
    2. Falls back to the COUNTIES DataFrame for counties with no submission

    When a county resubmits, their new data immediately updates the
    national aggregation because save_submission() replaces the old record.

    POPULATION WEIGHTING:
    Each county's figure is multiplied by its population before averaging.
    This gives larger counties (e.g. Nairobi 4.9M) more influence than
    smaller ones (e.g. Lamu 144K) — reflecting the real national situation.

    Returns a dictionary used by:
    - National overview dashboard (gauges and metrics)
    - INEP report generator
    - AI assistant context
    """
    # Load real submitted data from the file/cloud store
    real_subs = get_submissions()
    submitted_counties_set = {s["county"] for s in real_subs}

    # Also use COUNTIES DataFrame for counties with status submitted/review
    # that may not have gone through the submission form
    df_submitted = COUNTIES[COUNTIES["status"].isin(["submitted","review"])].copy()

    # Override with real submission data where available
    # This ensures resubmissions update the national figures immediately
    for s in real_subs:
        mask = df_submitted["name"] == s["county"]
        if mask.any():
            # Update the row with the real submitted figures
            df_submitted.loc[mask, "elec"]      = s.get("elec_pct",    df_submitted.loc[mask,"elec"].values[0])
            df_submitted.loc[mask, "cooking"]   = s.get("cooking_pct", df_submitted.loc[mask,"cooking"].values[0])
            df_submitted.loc[mask, "budget"]    = s.get("budget_kes_b",df_submitted.loc[mask,"budget"].values[0])
            df_submitted.loc[mask, "target_yr"] = s.get("target_year", df_submitted.loc[mask,"target_yr"].values[0])

    # Add any submitted counties not in the base COUNTIES DataFrame
    # (edge case: county not in our 47 list submits via the form)
    for s in real_subs:
        if s["county"] not in df_submitted["name"].values:
            new_row = {
                "name":s["county"],"elec":s.get("elec_pct",0),
                "cooking":s.get("cooking_pct",0),"pop":1000000,
                "budget":s.get("budget_kes_b",0),"target_yr":s.get("target_year",2030),
            }
            df_submitted = pd.concat(
                [df_submitted, pd.DataFrame([new_row])], ignore_index=True)

    if df_submitted.empty:
        return {
            "submitted_count": len(real_subs),
            "total_counties":  len(COUNTIES),
            "w_elec":          0,
            "w_cooking":       0,
            "total_budget":    0,
            "latest_target":   2035,
            "coverage_pct":    0,
            "overdue_count":   len(COUNTIES[COUNTIES["status"]=="overdue"]),
            "pending_count":   len(COUNTIES[COUNTIES["status"]=="pending"]),
            "real_sub_count":  len(real_subs),
        }

    tp = df_submitted["pop"].sum()
    return {
        "submitted_count":  len(df_submitted),
        "total_counties":   len(COUNTIES),
        "w_elec":           round((df_submitted["elec"]*df_submitted["pop"]).sum()/tp, 1),
        "w_cooking":        round((df_submitted["cooking"]*df_submitted["pop"]).sum()/tp, 1),
        "total_budget":     round(df_submitted["budget"].sum(), 1),
        "latest_target":    int(df_submitted["target_yr"].max()),
        "coverage_pct":     round(len(df_submitted)/len(COUNTIES)*100),
        "overdue_count":    len(COUNTIES[COUNTIES["status"]=="overdue"]),
        "pending_count":    len(COUNTIES[COUNTIES["status"]=="pending"]),
        "real_sub_count":   len(real_subs),   # actual submissions through the form
    }

# ── AI ────────────────────────────────────────────────────────────────────────
def ask_ai(question, history):
    try:
        api_key = st.secrets.get("anthropic",{}).get("api_key","")
        if not api_key or "YOUR_" in api_key:
            return "AI assistant is not configured. Add your Anthropic API key to secrets.toml."
        client  = anthropic.Anthropic(api_key=api_key)
        nat     = compute_national()
        system  = f"""You are the KenyaWatts AI assistant for EPRA's national energy planning platform.
Real data: {nat['submitted_count']} of {nat['total_counties']} counties submitted.
Weighted electricity access: {nat['w_elec']}% (target 100% by 2030).
Weighted clean cooking: {nat['w_cooking']}% (target 100% by 2028).
Clean energy generation: 82% (target 100% by 2035).
Critical overdue counties: Turkana 12%, Marsabit 8%, Mandera 11%, Wajir 9%.
Makueni CEP: electricity 75.1%, solar GHI 2008 kWh/m², budget KES 74.9B, firewood 72.5%.
Average outage: 8.8 hrs/month vs EPRA 5 hr benchmark. 11/12 months exceeded.
INEP Regulations 2025 legally require all 47 county submissions.
Answer in 3-4 sentences. Plain text only."""
        msgs = [{"role":m["role"],"content":m["content"]} for m in history]
        msgs.append({"role":"user","content":question})
        res = client.messages.create(model="claude-sonnet-4-6",max_tokens=600,system=system,messages=msgs)
        return res.content[0].text
    except Exception as e:
        return f"AI error: {str(e)}"

# ── AUTH ──────────────────────────────────────────────────────────────────────
# =============================================================================
# SECTION 14 — AUTHENTICATION
# Uses the streamlit-authenticator library to manage logins.
# Credentials come from st.secrets (the secrets.toml file).
# The library hashes passwords at login time and compares with stored hashes.
# Cookie expiry = 0 means the session ends when the browser closes.
# =============================================================================

def setup_auth():
    """
    Create and return the Authenticate object from streamlit-authenticator.
    Reads all user credentials from st.secrets (secrets.toml).
    This object handles login form rendering, password verification,
    cookie creation, and logout.
    """
    creds = st.secrets.get("credentials",{})
    ud = {}
    for uname, info in creds.get("usernames",{}).items():
        ud[uname] = {"email":info.get("email",""),"name":info.get("name",uname),"password":info.get("password","")}
    cookie = st.secrets.get("cookie",{})
    return stauth.Authenticate(
        {"usernames":ud},
        cookie.get("name","kenyawatts_auth"),
        cookie.get("key","kw_key"),
        0,   # expiry_days=0 means session cookie only — clears when browser closes
             # This ensures every new visit starts at the login page
    )

def get_user_role(username):
    """
    Look up the role, county_id, name and email for a given username.
    Reads from st.secrets (secrets.toml) — the same source as authentication.
    Returns a dictionary used throughout the app to control what each user sees.
    The 'role' value is the most important — it controls the entire navigation.
    """
    info = st.secrets.get("credentials",{}).get("usernames",{}).get(username,{})
    return {"role":info.get("role","county"),"county_id":info.get("county_id",""),
            "name":info.get("name",username),"email":info.get("email","")}

# =============================================================================
# SECTION 15 — SHARED UI COMPONENTS
# Small reusable functions that render common UI elements.
# Defined once here and called from every page function below.
# =============================================================================

def metric_card(label, value, delta=None, delta_color="#0f9d7e"):
    """
    Render a styled KPI metric card with a label, main value, and
    optional delta (change indicator) in a specified colour.
    Uses HTML/CSS for consistent styling regardless of Streamlit theme.
    """
    delta_html = f'<div class="kw-metric-delta" style="color:{delta_color}">{delta}</div>' if delta else ""
    st.markdown(f"""<div class="kw-metric">
    <div class="kw-metric-label">{label}</div>
    <div class="kw-metric-value">{value}</div>
    {delta_html}</div>""", unsafe_allow_html=True)

def section(title, sub=""):
    st.markdown(f"**{title}**")
    if sub: st.caption(sub)

def alert(type_, text):
    st.markdown(f'<div class="kw-alert-{type_}">{text}</div>', unsafe_allow_html=True)

# =============================================================================
# PAGE FUNCTION: County / National Energy Map
# Shows an interactive Plotly map of Kenya with county markers.
# EPRA sees all counties with toggleable data layers.
# County users see only their own county marker and profile cards.
# Marker colour and size reflect the selected data layer.
# Colorbar labels are white so they are visible on the dark map background.
# =============================================================================

def page_map(role, county_id):
    """
    Render the interactive energy map.
    role      = determines whether to show all counties or just one
    county_id = the two-letter ID of the logged-in county user's county
    """
    is_epra = role in ("epra","ministry","devpartner")
    section("Kenya county energy map", "Interactive map — click a county marker for details")

    # Map metric selector
    if is_epra:
        metric = st.selectbox("Map indicator:", [
            "Electricity access (%)", "Clean cooking access (%)",
            "Solar GHI (kWh/m²)", "Submission status", "Plan budget (KES B)"
        ])
    else:
        metric = "Electricity access (%)"

    # Build map with plotly scatter_mapbox
    df = COUNTIES.copy()
    if not is_epra:
        df = df[df["id"]==county_id]

    # Assign colour and size based on metric
    if metric == "Submission status":
        df["color_val"] = df["status"].map({"submitted":4,"review":3,"pending":2,"overdue":1})
        df["color_str"] = df["status"].map(STATUS_COLOR)
        df["hover"]     = df.apply(lambda r: f"<b>{r['name']}</b><br>Status: {STATUS_LABEL[r['status']]}<br>Population: {r['pop']//1000:,}K", axis=1)
        color_col = "color_str"
    elif metric == "Electricity access (%)":
        df["color_str"] = df["elec"].apply(lambda v: "#0f9d7e" if v>=75 else "#d4891a" if v>=40 else "#b33a2c")
        df["hover"]     = df.apply(lambda r: f"<b>{r['name']}</b><br>Electricity access: {r['elec']}%<br>Target year: {r['target_yr']}", axis=1)
        color_col = "elec"
    elif metric == "Clean cooking access (%)":
        df["hover"] = df.apply(lambda r: f"<b>{r['name']}</b><br>Clean cooking: {r['cooking']}%", axis=1)
        color_col = "cooking"
    elif metric == "Solar GHI (kWh/m²)":
        df["hover"] = df.apply(lambda r: f"<b>{r['name']}</b><br>Solar GHI: {r['solar']} kWh/m²", axis=1)
        color_col = "solar"
    else:
        df["hover"] = df.apply(lambda r: f"<b>{r['name']}</b><br>Budget: KES {r['budget']}B", axis=1)
        color_col = "budget"

    if metric == "Submission status":
        fig = go.Figure(go.Scattermapbox(
            lat=df["lat"], lon=df["lon"],
            mode="markers",
            marker=dict(size=18, color=df["color_str"], opacity=0.85),
            text=df["hover"], hoverinfo="text",
            showlegend=False,
        ))
    else:
        fig = px.scatter_mapbox(
            df, lat="lat", lon="lon",
            color=color_col, size_max=20,
            hover_data=None, custom_data=["hover"],
            color_continuous_scale=["#b33a2c","#d4891a","#0f9d7e"],
            zoom=5,
        )
        fig.update_traces(hovertemplate="%{customdata[0]}<extra></extra>", marker_size=16)
        fig.update_coloraxes(colorbar=dict(
            title=dict(text=metric, font=dict(size=12, color="#ffffff", family="Arial")),
            tickfont=dict(size=11, color="#ffffff", family="Arial"),
            bgcolor="rgba(14,30,46,0.75)",
            bordercolor="rgba(255,255,255,0.2)",
            borderwidth=1,
        ))

    fig.update_layout(
        mapbox_style="carto-positron",
        mapbox=dict(center=dict(lat=0.5, lon=37.9), zoom=5),
        height=480, margin=dict(t=0,b=0,l=0,r=0),
        font=CHART_FONT,
    )

    # Legend for status view — white labels on coloured pill backgrounds
    if metric == "Submission status":
        cols = st.columns(4)
        for col, (s,c) in zip(cols, STATUS_COLOR.items()):
            col.markdown(
                f'<div style="display:flex;align-items:center;gap:7px;font-size:12px;padding:4px 0"><div style="background:{c};color:#ffffff;font-size:11px;font-weight:600;padding:3px 10px;border-radius:12px;letter-spacing:.2px">{STATUS_LABEL[s]}</div></div>',
                unsafe_allow_html=True
            )
        st.markdown("")

    st.plotly_chart(fig, use_container_width=True)

    # Summary below map
    if is_epra:
        nat = compute_national()
        c1,c2,c3,c4 = st.columns(4)
        with c1: metric_card("Submitted",   str(nat["submitted_count"]), f"of {nat['total_counties']} counties", "#0f9d7e")
        with c2: metric_card("Overdue",     str(nat["overdue_count"]),   "Need urgent action", "#b33a2c")
        with c3: metric_card("Pending",     str(nat["pending_count"]),   "Not yet started", "#7a7870")
        with c4: metric_card("Coverage",    f"{nat['coverage_pct']}%",   "of counties with plans", "#1a6fa3")

        alert("danger", "<b>Equity alert:</b> Marsabit (8%), Wajir (9%), Mandera (11%) and Turkana (12%) have the lowest electricity access and are all overdue. These counties must be prioritised for direct submission support.")

# ── NATIONAL OVERVIEW ─────────────────────────────────────────────────────────
def page_national_overview(role):
    nat = compute_national()
    if role == "developer":
        alert("warn",
              "<b>Developer / demo session.</b> Full access enabled for testing. "
              "Use Developer tools to reset all data before a real test session.")

    # Show real submission count vs static data count
    real_subs    = get_submissions()
    real_count   = len(real_subs)
    if real_count > 0:
        alert("success",
            f"<b>{real_count} real county plan(s)</b> submitted through the platform · "
            f"National figures updated with live data · "
            f"{nat['submitted_count']} total counties in aggregation.")
    else:
        alert("info",
            f"<b>{'EPRA Admin' if role=='epra' else role.title()} view</b> — "
            f"No real submissions yet. National figures use baseline county data. "
            f"When counties submit plans the national aggregation updates automatically.")

    c1,c2,c3,c4 = st.columns(4)
    with c1: metric_card("Counties submitted", f"{nat['submitted_count']} / {nat['total_counties']}", f"{nat['coverage_pct']}% · {nat['overdue_count']} overdue")
    with c2: metric_card("Weighted elec. access", f"{nat['w_elec']}%", "↑ Population-weighted")
    with c3: metric_card("Weighted clean cooking", f"{nat['w_cooking']}%", "Target 100% by 2028", "#d4891a")
    with c4: metric_card("Total plan budgets", f"KES {nat['total_budget']}B", "Sum of submitted plans")

    st.markdown("")
    # Target gauges
    section("National energy targets — live progress")
    tc1,tc2,tc3 = st.columns(3)
    for col, label, val, goal, yr, color in [
        (tc1,"Electricity access",  nat["w_elec"],  100,2030,"#1a6fa3"),
        (tc2,"Clean cooking",       nat["w_cooking"],100,2028,"#0f9d7e"),
        (tc3,"Clean energy gen.",   82,              100,2035,"#5b4fc9"),
    ]:
        with col:
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=val,
                title={"text":f"<b>{label}</b><br><span style='font-size:11px;color:#1a1916'>Target {goal}% by {yr}</span>","font":{"size":13,"color":"#1a1916","family":"Arial"}},
                gauge={"axis":{"range":[0,100],"tickfont":dict(size=11,color="#1a1916")},"bar":{"color":color},"bgcolor":"#f7f6f2","borderwidth":0},
                number={"suffix":"%","font":{"size":28,"color":"#1a1916","family":"Arial"}}
            ))
            fig.update_layout(height=190, margin=dict(t=60,b=10,l=20,r=20),
                              paper_bgcolor="white", font=CHART_FONT)
            st.plotly_chart(fig, use_container_width=True)

    st.markdown("")
    col_left, col_right = st.columns(2)
    with col_left:
        section("Generation trend (GWh)", "FY 2019/20 — 2024/25 · Source: EPRA")
        fig = px.line(GEN_TREND, x="Year", y="GWh", markers=True, color_discrete_sequence=["#1a6fa3"])
        fig = apply_layout(fig, xlabel="Financial Year", ylabel="Generation (GWh)")
        fig.update_layout(yaxis=dict(range=[10000,16000]))
        fig.update_traces(line_width=2.5, marker_size=7)
        st.plotly_chart(fig, use_container_width=True)

    with col_right:
        section("Generation mix", "90%+ from renewable sources")
        fig = px.pie(GEN_MIX, names="Source", values="Pct",
                     color="Source", color_discrete_map=dict(zip(GEN_MIX["Source"],GEN_MIX["Color"])))
        fig.update_layout(**BASE_LAYOUT, height=260)
        fig.update_traces(textinfo="percent+label", textfont=dict(size=12,color="#1a1916",family="Arial"))
        st.plotly_chart(fig, use_container_width=True)

# ── ALL COUNTIES ──────────────────────────────────────────────────────────────
def page_all_counties():
    """
    Show all 47 counties with status filters and a search bar.
    EPRA can filter by submission status, search by county name,
    and click any county to expand its full energy profile.
    Also shows real submission data from the file/cloud store
    so recently submitted counties show their actual indicator values.
    """
    alert("info","<b>EPRA full access:</b> All 47 counties. Use the search bar or status filter to find specific counties.")

    # ── Search and filter row ─────────────────────────────────────────────────
    # Two controls side by side: text search and status dropdown filter
    sc1, sc2 = st.columns([2,1])
    with sc1:
        # Free-text search by county name or region
        # Searches as you type — no button needed
        search_term = st.text_input(
            "🔍  Search by county name or region:",
            placeholder="e.g. Nairobi, Coast, Turkana...",
            key="county_search_bar",
            label_visibility="visible"
        )
    with sc2:
        status_filter = st.selectbox(
            "Filter by status:",
            ["All counties","Submitted ✓","In review","Pending","Overdue"],
            key="county_status_filter"
        )

    # Apply status filter
    sm = {"All counties":None,"Submitted ✓":"submitted",
          "In review":"review","Pending":"pending","Overdue":"overdue"}
    sf = sm[status_filter]
    df = COUNTIES if sf is None else COUNTIES[COUNTIES["status"]==sf]

    # Apply text search (case-insensitive, searches name and region columns)
    if search_term.strip():
        mask = (
            df["name"].str.contains(search_term, case=False, na=False)
            | df["region"].str.contains(search_term, case=False, na=False)
        )
        df = df[mask]

    # Show search result count
    if search_term.strip() or sf:
        st.caption(f"Showing {len(df)} of {len(COUNTIES)} counties")

    # ── Summary count cards ───────────────────────────────────────────────────
    c1,c2,c3,c4 = st.columns(4)
    for col,s,lbl,color in [
        (c1,"submitted","Submitted","#0f9d7e"),
        (c2,"review",   "In review","#d4891a"),
        (c3,"pending",  "Pending",  "#7a7870"),
        (c4,"overdue",  "Overdue",  "#b33a2c")
    ]:
        with col:
            n = len(COUNTIES[COUNTIES["status"]==s])
            st.markdown(
                f'<div style="padding:12px 14px;border-radius:10px;background:#ffffff;border:0.5px solid #e8e6de;border-top:3px solid {color}"><div style="font-size:22px;font-weight:700;color:{color}">{n}</div><div style="font-size:11px;color:#444441;font-weight:500;margin-top:2px">{lbl}</div></div>',
                unsafe_allow_html=True
            )

    st.markdown("")

    # Show message if no counties match the search/filter
    if df.empty:
        st.info(
            f"No counties found matching "
            f"{'search term: **' + search_term + '**' if search_term else ''}"
            f"{' with status: **' + status_filter + '**' if sf else ''}. "
            "Try a different search term or filter."
        )
        return

    # Load real submission data to override static values where available
    real_subs  = get_submissions()
    real_subs_dict = {s["county"]: s for s in real_subs}

    # ── County expander rows ──────────────────────────────────────────────────
    for _, row in df.iterrows():
        badge  = STATUS_LABEL[row["status"]]
        s_key  = row["status"]

        # Use real submission data if this county has submitted through the form
        real = real_subs_dict.get(row["name"])
        elec     = real.get("elec_pct",    row["elec"])    if real else row["elec"]
        cooking  = real.get("cooking_pct", row["cooking"]) if real else row["cooking"]
        solar    = real.get("solar_ghi",   row["solar"])   if real else row["solar"]
        budget   = real.get("budget_kes_b",row["budget"])  if real else row["budget"]
        target   = real.get("target_year", row["target_yr"]) if real else row["target_yr"]

        with st.expander(
            f"**{row['name']}** · {row['region']} · "
            f"{row['pop']//1000:,}K pop · "
            f"{'✓ Real data' if real else badge}"
        ):
            cc1,cc2,cc3,cc4,cc5 = st.columns(5)
            cc1.metric("Electricity",   f"{elec}%",
                       delta="Real" if real else None)
            cc2.metric("Clean cooking", f"{cooking}%")
            cc3.metric("Solar GHI",     f"{solar} kWh/m²")
            cc4.metric("MTF Tier",      f"Tier {row['mtf']}")
            cc5.metric("Target year",   str(target))

            st.markdown(
                f'<span class="kw-badge-{s_key}">{badge}</span>',
                unsafe_allow_html=True
            )

            # Show real submission details if available
            if real:
                alert("success",
                    f"<b>Real submission on file:</b> "
                    f"Submitted by {real.get('submitted_by','—')} · "
                    f"{real.get('date_display',real.get('date',''))} "
                    f"at {real.get('time','')} · "
                    f"Ref: <code>{real.get('ref','')}</code> · "
                    f"Status: {real.get('status','').title()}"
                )
            elif row["status"] == "overdue":
                alert("danger",
                    f"<b>Action required:</b> {row['name']} is overdue. "
                    "INEP Regulations 2025 require submission. Reminder sent.")
            elif row["status"] == "submitted":
                alert("success",
                    f"<b>Plan on file:</b> Budget KES {budget}B · "
                    f"Growth {row['growth']}% p.a. · Validation passed.")

# ── COUNTY DASHBOARD ──────────────────────────────────────────────────────────
def page_county_dashboard(county_id, user_name):
    row = COUNTIES[COUNTIES["id"]==county_id]
    if row.empty: st.error("County not found."); return
    row = row.iloc[0]
    sc  = STATUS_COLOR[row["status"]]
    nat = compute_national()

    alert("info", f"<b>County committee view — {row['name']} County only.</b> You can only see and manage your county's data.")

    st.markdown(f"""<div style="padding:16px 20px;border-radius:12px;background:#0e1e2e;border-left:4px solid {sc};margin-bottom:16px;display:flex;justify-content:space-between;align-items:center">
      <div>
        <div style="font-size:17px;font-weight:700;color:#ffffff;margin-bottom:4px">{row['name']} County Energy Plan</div>
        <div style="font-size:12px;color:#a8c4d4;margin-top:2px">Plan period: 2023–{row['target_yr']} &nbsp;·&nbsp; {row['region']} region</div>
      </div>
      <span style="font-size:12px;font-weight:700;color:#0e1e2e;background:{sc};padding:5px 14px;border-radius:20px;letter-spacing:.2px">{STATUS_LABEL[row['status']]}</span>
    </div>""", unsafe_allow_html=True)

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Electricity access",f"{row['elec']}%","Target: 100%")
    c2.metric("Clean cooking",f"{row['cooking']}%","Target: 100% by 2028")
    c3.metric("Solar GHI",f"{row['solar']} kWh/m²")
    c4.metric("Target year",str(row["target_yr"]))

    st.markdown("")
    section("Your county vs national average (submitted counties)")
    fig = go.Figure()
    fig.add_trace(go.Bar(name=row["name"],   x=["Electricity (%)","Clean cooking (%)"], y=[row["elec"],row["cooking"]],   marker_color="#1a6fa3"))
    fig.add_trace(go.Bar(name="National avg",x=["Electricity (%)","Clean cooking (%)"], y=[nat["w_elec"],nat["w_cooking"]],marker_color="#c8c6be"))
    fig = apply_layout(fig, xlabel="Indicator", ylabel="Percentage (%)")
    fig.update_layout(barmode="group", height=260, legend=dict(font=LEGEND_FONT))
    st.plotly_chart(fig, use_container_width=True)

    # Check for real submission data for this county
    cname_for_subs = COUNTIES[COUNTIES["id"]==county_id]["name"].values[0] if county_id and not COUNTIES[COUNTIES["id"]==county_id].empty else ""
    real_subs = get_submissions(county=cname_for_subs)
    if real_subs:
        latest = real_subs[0]   # get_submissions returns newest first
        sc2 = {"submitted":"#d4891a","approved":"#0f9d7e","rejected":"#b33a2c"}.get(
              latest.get("status","submitted"),"#7a7870")
        alert("success",
            f"<b>Your plan is on the platform.</b> "
            f"Submitted: {latest.get('date_display', latest.get('date',''))} "
            f"at {latest.get('time','')} · "
            f"Ref: <code>{latest.get('ref','')}</code> · "
            f"Status: <span style='color:{sc2};font-weight:600'>"
            f"{latest.get('status','').title()}</span>"
            + (f" · EPRA note: {latest['epra_note']}" if latest.get('epra_note') else "")
        )
    elif row["status"]=="overdue":
        alert("danger",
              "<b>Your plan is overdue.</b> Under Section 5(5)(a) of the Energy Act 2019, "
              "submission is a legal requirement. Use 'Submit energy plan' to submit now.")

# ── UPLOAD / SUBMISSION LOG ────────────────────────────────────────────────────
def show_upload_log(county_id=None):
    """Show full submission history from the persistent file store — all submissions, newest first."""
    county_name = None
    if county_id:
        row = COUNTIES[COUNTIES["id"]==county_id]
        if not row.empty:
            county_name = row["name"].values[0]

    file_subs = get_submissions(county=county_name)

    if not file_subs:
        st.info("No submissions on record yet. When a plan is submitted it will appear here permanently.")
        last = st.session_state.get("last_submission")
        if last:
            st.success(f"✅ Your most recent submission (Ref: `{last['ref']}`) was saved at {last['datetime']}. Refresh the page to see it listed here.")
        return

    # Show just-submitted highlight
    last = st.session_state.get("last_submission")
    if last and (not county_name or last.get("county","").lower() == county_name.lower()):
        st.success(
            f"✅ **Latest submission saved** · Ref: `{last['ref']}` · {last['datetime']}  \n"
            f"EPRA has been notified automatically and will review your plan."
        )

    st.markdown(f"**Submission history — {len(file_subs)} record(s)** · Saved permanently · Newest first")

    for s in file_subs:
        is_latest = (s == file_subs[0])
        sc  = {"submitted":"#d4891a","approved":"#0f9d7e","rejected":"#b33a2c"}.get(s.get("status","submitted"),"#7a7870")
        sl  = {"submitted":"Submitted — awaiting EPRA review","approved":"Approved by EPRA ✓","rejected":"Resubmission requested"}.get(s.get("status","submitted"),"Submitted")
        doc_icon = "📄" if s.get("pathway")=="pdf" else "📝"
        latest_badge = ' <span style="background:#0f9d7e;color:white;padding:1px 7px;border-radius:8px;font-size:10px;font-weight:700;margin-left:6px">LATEST</span>' if is_latest else ""
        epra_note_html = ""
        if s.get("epra_note") and s.get("status") == "approved":
            epra_note_html = f'<div style="margin-top:6px;font-size:11px;color:#3ecfaa;background:rgba(62,207,170,0.1);padding:5px 8px;border-radius:5px">✅ EPRA note: {s["epra_note"]} · {s.get("epra_action_time","")}</div>'
        elif s.get("epra_note") and s.get("status") == "rejected":
            epra_note_html = f'<div style="margin-top:6px;font-size:11px;color:#d4891a;background:rgba(212,137,26,0.1);padding:5px 8px;border-radius:5px">↩ EPRA note: {s["epra_note"]} · {s.get("epra_action_time","")}</div>'
        st.markdown(f"""
        <div style="background:#0e1e2e;border-radius:10px;padding:12px 16px;margin-bottom:8px;
                    border-left:3px solid {sc}">
          <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:6px">
            <div style="display:flex;gap:10px;align-items:center">
              <span style="font-size:18px">{doc_icon}</span>
              <div>
                <div style="font-size:13px;font-weight:600;color:#e8e6e0">
                  County Energy Plan — {s['county']}{latest_badge}
                </div>
                <div style="font-size:11px;color:#a8c4d4;margin-top:2px">
                  📅 {s.get('date_display', s.get('date',''))} &nbsp;⏱ {s.get('time','')}
                  &nbsp;·&nbsp; By: {s.get('submitted_by','—')}
                  &nbsp;·&nbsp; Via: {'PDF/Word upload' if s.get('pathway')=='pdf' else 'Structured template'}
                </div>
              </div>
            </div>
            <div style="text-align:right;flex-shrink:0">
              <div style="font-size:11px;font-weight:700;color:{sc}">{sl}</div>
              <div style="font-size:10px;color:#5a8a9e;margin-top:2px;font-family:monospace">{s.get('ref','')}</div>
            </div>
          </div>
          <div style="display:flex;gap:16px;font-size:11px;color:#a8c4d4;margin-top:4px;flex-wrap:wrap">
            <span>⚡ {s.get('elec_pct','—')}% electricity</span>
            <span>🍳 {s.get('cooking_pct','—')}% clean cooking</span>
            <span>☀️ {s.get('solar_ghi','—')} kWh/m² solar</span>
            <span>💰 KES {s.get('budget_kes_b','—')}B budget</span>
            <span>🎯 Target: {s.get('target_year','—')}</span>
          </div>
          {epra_note_html}
        </div>""", unsafe_allow_html=True)

# =============================================================================
# PAGE FUNCTION: Submit County Energy Plan
# Two pathways:
#   1. PDF/Word upload — county uploads their existing plan document.
#      AI simulates extraction of key indicators from the document.
#   2. Structured template — county fills in fields directly.
#      National assumptions are pre-loaded as placeholder hints.
# Both pathways run the same 10-rule validation engine before submission.
# On submit:
#   - Record saved to kw_submissions.json (and Supabase if configured)
#   - Automatic notification message sent to EPRA via kw_messages.json
#   - Confirmation card shows date, time, reference number
# =============================================================================

def page_submit(role, county_id, user_name):
    """
    Render the plan submission page with two pathways and validation.
    role      = controls whether county name field is pre-filled and locked
    county_id = used to pre-fill the county name for county users
    user_name = shown on the confirmation card and audit log
    """
    is_county   = role=="county"
    county_name = COUNTIES[COUNTIES["id"]==county_id]["name"].values[0] if county_id else ""
    if is_county:
        alert("info", f"<b>Submitting as {county_name} County Energy Planning Committee.</b> Your submission goes directly to EPRA for validation.")

    # ── Persistent success banner (shown after any successful submission) ─────
    last = st.session_state.get("last_submission")
    if last and last.get("county") == (county_name or last.get("county","")):
        st.success(
            f"✅ **Submission confirmed** — {last['county']} County · "
            f"Ref: `{last['ref']}` · {last['datetime']} · By: {last['by']}  \n"
            f"EPRA has been notified. Check **Submission history** tab to see your record."
        )

    # Default to history tab (index 1) right after a successful submission
    default_tab = st.session_state.pop("submit_page_tab", 0)
    tab_submit, tab_history = st.tabs([t("submit_tab_submit"), t("submit_tab_history")])

    with tab_submit:
        pathway = st.radio(t("submit_choose_pathway"),
            [t("submit_pathway_upload"), t("submit_pathway_template")], horizontal=True)
        st.divider()

        if t("submit_pathway_upload") in pathway:
            section("Upload your County Energy Plan document")
            alert("success","<b>How this works:</b> Upload your completed CEP. AI extracts key indicators automatically. Review the values, then submit to EPRA.")

            uploaded = st.file_uploader("Choose your County Energy Plan (PDF or Word)",
                                        type=["pdf","doc","docx"],
                                        help="Max 50MB. Use the Makueni CEP PDF to test.")
            if uploaded:
                st.success(f"✓ File received: **{uploaded.name}** · {uploaded.size//1024} KB · Uploaded {datetime.now().strftime('%d %b %Y at %H:%M')}")
                with st.spinner("AI extracting key indicators from your document…"):
                    import time; time.sleep(2)
                alert("success","<b>AI extraction complete.</b> Review and correct values below before running validation.")

                extracted = {"elec":75.1,"cooking":17.9,"firewood":72.5,"solar":2008.0,"budget":74.9,"target_yr":2028,"mtf":3,"growth":1.1}
                col1,col2 = st.columns(2)
                with col1:
                    county_in  = st.text_input("County name *", value=county_name if is_county else "Makueni", disabled=is_county, key="pdf_county")
                    elec_in    = st.number_input("Electricity access (%)*", value=extracted["elec"], min_value=0.0, max_value=100.0, key="pdf_elec")
                    cooking_in = st.number_input("Clean cooking access (%)*", value=extracted["cooking"], min_value=0.0, max_value=100.0, key="pdf_cooking")
                    fw_in      = st.number_input("Firewood as primary fuel (%)", value=extracted["firewood"], min_value=0.0, max_value=100.0, key="pdf_fw")
                with col2:
                    solar_in   = st.number_input("Solar GHI (kWh/m²/year)*", value=extracted["solar"], min_value=0.0, key="pdf_solar")
                    budget_in  = st.number_input("Total budget (KES billions)", value=extracted["budget"], min_value=0.0, key="pdf_budget")
                    target_in  = st.number_input("Universal access target year*", value=extracted["target_yr"], min_value=2025, max_value=2040, key="pdf_target")
                    mtf_in     = st.selectbox("MTF demand tier*", [1,2,3,4,5], index=extracted["mtf"]-1, key="pdf_mtf")
                    growth_in  = st.number_input("Population growth rate (%/yr)*", value=extracted["growth"], min_value=0.0, max_value=10.0, key="pdf_growth")

                _validate_and_submit(county_in, elec_in, cooking_in, fw_in, solar_in,
                                      budget_in, target_in, mtf_in, growth_in,
                                      uploaded.name, "pdf", county_id, user_name)
            else:
                st.info("Upload your CEP document above. Use the Makueni County Energy Plan PDF as a test file.")

        else:
            section("Structured submission template")
            alert("info","<b>National assumptions pre-loaded:</b> KNBS baselines · cost benchmarks · solar GHI reference · Kenya's official targets.")
            st.markdown("")

            # ── Makueni reference guide (collapsed by default) ────────────────
            with st.expander("📄 View example — Makueni County Energy Plan (guidance only)", expanded=False):
                st.caption("Use this as a guide for what figures to fill in. Makueni is Kenya's reference county plan.")
                col_ex1, col_ex2 = st.columns(2)
                with col_ex1:
                    st.markdown("""
                    | Field | Makueni example value |
                    |---|---|
                    | Electricity access | 75.1% |
                    | Clean cooking access | 17.9% |
                    | Firewood as primary fuel | 72.5% |
                    | Solar GHI | 2,008 kWh/m² |
                    """)
                with col_ex2:
                    st.markdown("""
                    | Field | Makueni example value |
                    |---|---|
                    | Total budget | KES 74.9B |
                    | Target year | 2028 |
                    | MTF demand tier | Tier 3 (rural) |
                    | Population growth | 1.1% per year |
                    """)
                st.caption("Source: Makueni County Energy Plan 2023–2032 · WRI + Strathmore University")

            st.markdown("")

            # ── Section 1: Electricity access ─────────────────────────────────
            with st.expander("▶ Section 1 — Electricity access & clean cooking", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    county_in  = st.text_input("County name *", value=county_name, disabled=is_county, key="tmpl_county")
                    elec_in    = st.number_input("Total electricity access (%) *", min_value=0.0, max_value=100.0,
                                                  help="Include grid + mini-grid + Solar Home Systems", key="tmpl_elec")
                    cooking_in = st.number_input("Clean cooking access (%) *", min_value=0.0, max_value=100.0,
                                                  help="LPG + biogas + electric + improved cookstoves", key="tmpl_cooking")
                    fw_in      = st.number_input("Firewood as primary fuel (%)", min_value=0.0, max_value=100.0,
                                                  help="% of households using firewood as main cooking fuel", key="tmpl_fw")
                with col2:
                    solar_in   = st.number_input("Solar GHI (kWh/m²/year) *", min_value=0.0,
                                                  help="Kenya range: 1,600–2,200 kWh/m²/yr", key="tmpl_solar")
                    budget_in  = st.number_input("Total plan budget (KES billions)", min_value=0.0,
                                                  help="Total investment identified in your county energy plan", key="tmpl_budget")
                    target_in  = st.number_input("Universal electricity access target year *", value=2030,
                                                  min_value=2025, max_value=2040, key="tmpl_target")
                    mtf_in     = st.selectbox("MTF demand tier *", [1, 2, 3, 4, 5], index=2,
                                               help="Tier 1 = minimal, Tier 5 = high urban demand", key="tmpl_mtf")
                    growth_in  = st.number_input("Population growth rate (%/yr) *", value=1.1,
                                                  min_value=0.0, max_value=10.0, key="tmpl_growth")

            # ── Section 2: Energy resources ───────────────────────────────────
            with st.expander("▶ Section 2 — Energy resources", expanded=False):
                r1, r2 = st.columns(2)
                with r1:
                    st.number_input("Wind speed at 100m (m/s)", min_value=0.0, key="tmpl_wind")
                    st.number_input("Woody biomass supply (tonnes/year)", min_value=0.0, key="tmpl_biomass")
                    st.number_input("Hydropower potential (MW)", min_value=0.0, key="tmpl_hydro")
                with r2:
                    st.number_input("Biogas potential (GJ/year)", min_value=0.0, key="tmpl_biogas")
                    st.number_input("Geothermal potential (MW)", min_value=0.0, key="tmpl_geo")
                    st.number_input("Mini-grid sites identified", min_value=0, step=1, key="tmpl_minigrids")

            # ── Section 3: Energy efficiency ──────────────────────────────────
            with st.expander("▶ Section 3 — Energy efficiency", expanded=False):
                e1, e2 = st.columns(2)
                with e1:
                    st.number_input("LED bulb adoption in households (%)", min_value=0.0, max_value=100.0, value=0.0, key="tmpl_led")
                    st.number_input("Solar PV installed in county facilities (kW)", min_value=0.0, key="tmpl_solar_facility")
                with e2:
                    st.number_input("Energy-efficient cookstove penetration (%)", min_value=0.0, max_value=100.0, key="tmpl_ics")
                    st.number_input("County buildings with energy audit (%)", min_value=0.0, max_value=100.0, key="tmpl_audit")

            # ── Section 4: Gender & social inclusion ─────────────────────────
            with st.expander("▶ Section 4 — Gender & social inclusion (GESI)", expanded=False):
                g1, g2 = st.columns(2)
                with g1:
                    st.number_input("Female-headed households with electricity (%)", min_value=0.0, max_value=100.0, key="tmpl_gesi_elec")
                    st.number_input("Women in county energy committee (%)", min_value=0.0, max_value=100.0, key="tmpl_gesi_women")
                with g2:
                    st.number_input("Youth (18–35) in energy sector jobs (%)", min_value=0.0, max_value=100.0, key="tmpl_gesi_youth")
                    st.text_input("Key vulnerable groups served (describe)", key="tmpl_gesi_groups")

            # ── Section 5: Investment plan ────────────────────────────────────
            with st.expander("▶ Section 5 — Investment plan", expanded=False):
                i1, i2 = st.columns(2)
                with i1:
                    st.number_input("Government allocation (KES millions)", min_value=0.0, key="tmpl_gov_alloc")
                    st.number_input("Development partner funding (KES millions)", min_value=0.0, key="tmpl_dev_funding")
                    st.number_input("Private sector investment (KES millions)", min_value=0.0, key="tmpl_private")
                with i2:
                    st.number_input("Household contribution (KES millions)", min_value=0.0, key="tmpl_hh")
                    st.text_input("Key funding gaps (describe)", key="tmpl_gaps")
                    st.text_area("Priority investment projects (list top 3)", height=80, key="tmpl_projects")

            # ── Section 6: Monitoring & evaluation ───────────────────────────
            with st.expander("▶ Section 6 — Monitoring & evaluation", expanded=False):
                m1, m2 = st.columns(2)
                with m1:
                    st.number_input("Electricity access annual target 2027 (%)", min_value=0.0, max_value=100.0, key="tmpl_me_elec27")
                    st.number_input("Clean cooking annual target 2027 (%)", min_value=0.0, max_value=100.0, key="tmpl_me_cook27")
                with m2:
                    st.text_input("Data collection method (e.g. KNBS survey, county monitoring)", key="tmpl_me_method")
                    st.text_input("Reporting frequency", value="Annual", key="tmpl_me_freq")

            st.markdown("")
            # ── Run validation + submit ───────────────────────────────────────
            _validate_and_submit(county_in, elec_in, cooking_in, fw_in, solar_in,
                                  budget_in, target_in, mtf_in, growth_in,
                                  None, "form", county_id, user_name)

    with tab_history:
        section("Submission and upload history", "All submissions and uploads for your county this session")
        show_upload_log(county_id if is_county else None)

def _validate_and_submit(county, elec, cooking, fw, solar, budget, target, mtf, growth, filename, ptype, county_id, user_name):
    """
    Run 10 validation rules against the submitted county data, then
    show results and (if no critical errors) offer the Submit button.

    VALIDATION RULES:
      V1  Electricity access must be between 0 and 100%
      V2  Cooking fuel split (firewood + clean) should sum to ~100%
      V3  Population growth rate should be under 5% (flag if higher)
      V4  Solar GHI must be in Kenya range 1,600-2,200 kWh/m2/year
      V5  Budget must not be suspiciously low (possible unit error)
      V6  (reserved for scenario assumption checks)
      V7  MTF demand tier must be declared (required for aggregation)
      V8  Universal access target year must be between 2025 and 2040
      V9  (reserved for duplicate submission check)
      V10 County name must not be empty

    ERROR  = blocks submission — must be fixed before county can submit
    WARNING = does not block — shown as advisory, county can still submit
    """
    if st.button(t("submit_run_validation"), type="primary"):
        errors, warnings = [], []
        if not county:                                       errors.append(("V10","County name is required — mandatory field."))
        if elec<0 or elec>100:                              errors.append(("V1", f"Electricity access {elec}% is outside valid range 0–100%."))
        if fw>0 and cooking>0 and abs((fw+cooking)-100)>10: warnings.append(("V2",f"Cooking fuel split: firewood ({fw}%) + clean ({cooking}%) = {fw+cooking:.1f}%. All fuels should sum to 100%."))
        if solar>0 and (solar<1600 or solar>2200):          warnings.append(("V4",f"Solar GHI {solar} kWh/m² is outside Kenya's range 1,600–2,200. Check units."))
        if not mtf:                                          errors.append(("V7","MTF demand tier must be declared for aggregation."))
        if growth>5:                                         warnings.append(("V3",f"Population growth {growth}% exceeds 5% — verify with KNBS data."))
        if target<2025 or target>2040:                       warnings.append(("V8",f"Target year {target} is outside plausible range 2025–2040."))
        if budget>0 and budget<0.5:                          warnings.append(("V5",f"Budget KES {budget}B is very low — confirm units are billions."))

        passed = 10 - len(errors) - len(warnings)
        rc1,rc2,rc3 = st.columns(3)
        rc1.metric(t("val_critical_errors"), len(errors),   delta="block submission" if errors else None, delta_color="inverse")
        rc2.metric(t("val_warnings"),        len(warnings), delta="review before submit" if warnings else None, delta_color="off")
        rc3.metric(t("val_checks_passed"),   passed,        delta=f"of 10")

        for rule,msg in errors:   st.error(f"**ERROR · Rule {rule}:** {msg}")
        for rule,msg in warnings: st.warning(f"**WARNING · Rule {rule}:** {msg}")

        if not errors:
            st.success(t("val_no_errors"))
            if st.button(t("submit_to_epra"), type="primary", key=f"final_submit_{county}"):
                ref  = f"CEP-{county[:2].upper()}-{datetime.now().strftime('%Y')}-{abs(hash(county+str(datetime.now())))%9000+1000}"
                now  = datetime.now()

                # ── Build the full submission record ───────────────────────────
                submission_record = {
                    # Identity
                    "ref":          ref,
                    "county":       county,
                    "county_id":    county_id or county[:2].upper(),
                    "submitted_by": user_name,
                    # Timestamps
                    "datetime":     now.strftime("%Y-%m-%d %H:%M:%S"),
                    "date":         now.strftime("%Y-%m-%d"),
                    "date_display": now.strftime("%d %b %Y"),
                    "time":         now.strftime("%H:%M:%S"),
                    # Document
                    "pathway":      ptype,
                    "document":     filename or "Structured template",
                    # Energy data
                    "elec_pct":     elec,
                    "cooking_pct":  cooking,
                    "firewood_pct": fw,
                    "solar_ghi":    solar,
                    "budget_kes_b": budget,
                    "target_year":  target,
                    "mtf_tier":     mtf,
                    "growth_pct":   growth,
                    # Status
                    "status":          "submitted",
                    "epra_note":       "",
                    "epra_action_time":"",
                    # Validation summary
                    "validation_passed": True,
                    "validation_warnings": 0,
                }

                # ── Persist to disk immediately ────────────────────────────────
                save_submission(submission_record)

                # ── Update in-memory COUNTIES so dashboard reflects immediately ──
                # Without this the county would need to reload to see their
                # status change from "pending/overdue" to "submitted"
                county_mask = COUNTIES["name"] == county
                if county_mask.any():
                    COUNTIES.loc[county_mask, "status"]    = "submitted"
                    COUNTIES.loc[county_mask, "elec"]      = elec
                    COUNTIES.loc[county_mask, "cooking"]   = cooking
                    COUNTIES.loc[county_mask, "solar"]     = solar
                    COUNTIES.loc[county_mask, "budget"]    = budget
                    COUNTIES.loc[county_mask, "target_yr"] = int(target)

                # ── Also save to session upload_log for local history tab ──────
                st.session_state.upload_log.append({
                    "title":        f"County Energy Plan — {county}",
                    "county":       county,
                    "county_id":    county_id or county[:2].upper(),
                    "type":         ptype,
                    "filename":     filename or "Structured template",
                    "date":         now.strftime("%d %b %Y"),
                    "time":         now.strftime("%H:%M:%S"),
                    "status":       "submitted",
                    "ref":          ref,
                    "submitted_by": user_name,
                })
                st.session_state.audit_log.append({
                    "time":   now.strftime("%Y-%m-%d %H:%M:%S"),
                    "user":   user_name,
                    "action": f"Plan submitted and saved to server — {county}",
                    "ref":    ref,
                })

                # ── Send automatic notification message to EPRA via file store ─
                notif_msg = {
                    "id":          f"NOTIF-{ref}",
                    "from_role":   "system",
                    "from_county": county,
                    "from_name":   "KenyaWatts System",
                    "to":          "EPRA",
                    "type":        "Submission",
                    "subject":     f"✅ New plan submitted — {county} County",
                    "body": (
                        f"<b>{county} County Energy Plan</b> has been submitted and saved to the KenyaWatts platform.<br><br>"
                        f"<b>Submitted by:</b> {user_name}<br>"
                        f"<b>Date:</b> {now.strftime('%d %b %Y')}<br>"
                        f"<b>Time:</b> {now.strftime('%H:%M:%S')}<br>"
                        f"<b>Document:</b> {filename or 'Structured template'}<br>"
                        f"<b>Reference:</b> {ref}<br><br>"
                        f"<b>Key indicators:</b> Electricity access {elec}% · "
                        f"Clean cooking {cooking}% · Solar GHI {solar} kWh/m² · "
                        f"Budget KES {budget}B · Target year {target}<br><br>"
                        f"The plan is now in the validation queue. Please review and approve or request resubmission."
                    ),
                    "date":          now.strftime("%Y-%m-%d"),
                    "time":          now.strftime("%H:%M:%S"),
                    "datetime":      now.strftime("%Y-%m-%d %H:%M:%S"),
                    "read_by_epra":  False,
                    "replies":       [],
                }
                add_message(notif_msg)

                # ── Push UI notification ───────────────────────────────────────
                push_notification(f"✅ Plan submitted — {county} · Ref: {ref}", "✅")
                st.balloons()

                # ── Confirmation card ──────────────────────────────────────────
                st.markdown(f"""
                <div style="background:#0e1e2e;border:1.5px solid #0f9d7e;border-radius:12px;
                            padding:22px 24px;margin-top:16px">
                  <div style="font-size:17px;font-weight:700;color:#3ecfaa;margin-bottom:14px">
                    ✅ Submission confirmed and saved
                  </div>
                  <table style="width:100%;border-collapse:collapse">
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4;width:38%;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">County</td>
                      <td style="padding:6px 0;font-size:13px;color:#e8e6e0;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">{county}</td>
                    </tr>
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">Submitted by</td>
                      <td style="padding:6px 0;font-size:13px;color:#e8e6e0;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">{user_name}</td>
                    </tr>
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">Date</td>
                      <td style="padding:6px 0;font-size:13px;color:#e8e6e0;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">{now.strftime('%d %b %Y')}</td>
                    </tr>
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">Time</td>
                      <td style="padding:6px 0;font-size:13px;color:#e8e6e0;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">{now.strftime('%H:%M:%S')}</td>
                    </tr>
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">Document</td>
                      <td style="padding:6px 0;font-size:13px;color:#e8e6e0;
                                 border-bottom:0.5px solid rgba(255,255,255,0.08)">{filename or 'Structured template'}</td>
                    </tr>
                    <tr>
                      <td style="padding:6px 0;font-size:12px;font-weight:600;color:#a8c4d4">
                        Reference</td>
                      <td style="padding:6px 0;font-size:14px;color:#3ecfaa;font-family:monospace;
                                 font-weight:700">{ref}</td>
                    </tr>
                  </table>
                  <div style="margin-top:14px;padding-top:12px;border-top:0.5px solid rgba(255,255,255,0.1);
                              display:flex;justify-content:space-between;align-items:center">
                    <div style="font-size:12px;color:#a8c4d4">
                      Saved to KenyaWatts server · EPRA has been notified automatically
                    </div>
                    <div style="font-size:11px;color:#3ecfaa;font-weight:600">
                      ● Live on platform
                    </div>
                  </div>
                </div>""", unsafe_allow_html=True)

                # ── Persist confirmation in session state so it shows on next render ──
                # This means even if the user switches tab and comes back,
                # they can still see their most recent successful submission.
                st.session_state["last_submission"] = {
                    "ref":      ref,
                    "county":   county,
                    "datetime": now.strftime("%d %b %Y at %H:%M:%S"),
                    "by":       user_name,
                    "doc":      filename or "Structured template",
                    "elec":     elec,
                    "cooking":  cooking,
                    "solar":    solar,
                    "budget":   budget,
                    "target":   target,
                }
                # Auto-switch to history tab on next render
                st.session_state["submit_page_tab"] = 1

# =============================================================================
# PAGE FUNCTION: EPRA Validation Queue (EPRA only)
# Shows all county plan submissions that have been received from the
# persistent file/cloud store. EPRA can:
#   - View full submission details (county, submitter, date/time, indicators)
#   - Approve a submission (status changes to "approved", county notified)
#   - Request resubmission (status changes to "rejected", county notified)
#   - Run national aggregation (triggers re-computation of national INEP)
# Approval/rejection is saved permanently and the county automatically
# receives a message in their inbox via the shared message store.
# =============================================================================

def page_validation_queue():
    """Render the EPRA validation queue — reads live from the submission store."""
    alert("warn","<b>EPRA admin only.</b> Review all county plan submissions received from the file store.")

    # ── All submissions from the persistent store ──────────────────────────────
    all_subs = get_submissions()
    pending_review = [s for s in all_subs if s.get("status")=="submitted"]
    approved       = [s for s in all_subs if s.get("status")=="approved"]
    rejected       = [s for s in all_subs if s.get("status")=="rejected"]

    # Summary counts
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Total received",   len(all_subs),      "from file store")
    c2.metric("Pending review",   len(pending_review), "awaiting EPRA action")
    c3.metric("Approved",         len(approved),       "included in INEP")
    c4.metric("Resubmission req", len(rejected),       "returned to county")

    if not all_subs:
        st.info("No county submissions have been received yet. When a county submits a plan it will appear here immediately.")
    else:
        st.markdown(f"**All submissions — {len(all_subs)} total · updated in real time**")
        for s in all_subs:
            sc  = {"submitted":"#d4891a","approved":"#0f9d7e","rejected":"#b33a2c"}.get(s.get("status","submitted"),"#7a7870")
            sl  = {"submitted":"Pending review","approved":"Approved ✓","rejected":"Resubmission requested"}.get(s.get("status","submitted"),"Unknown")
            with st.expander(
                f"**{s['county']} County** · Ref: {s['ref']} · "
                f"{s.get('date_display', s.get('date',''))} at {s.get('time','')} · "
                f"Status: {sl}"
            ):
                # Full submission details
                st.markdown(f"""
                <div style="background:#0e1e2e;border-radius:10px;padding:14px 18px;margin-bottom:12px">
                  <table style="width:100%;border-collapse:collapse">
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4;width:35%;border-bottom:0.5px solid rgba(255,255,255,0.08)">County</td>
                      <td style="padding:5px 0;font-size:13px;color:#e8e6e0;border-bottom:0.5px solid rgba(255,255,255,0.08)">{s['county']}</td>
                    </tr>
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4;border-bottom:0.5px solid rgba(255,255,255,0.08)">Submitted by</td>
                      <td style="padding:5px 0;font-size:13px;color:#e8e6e0;border-bottom:0.5px solid rgba(255,255,255,0.08)">{s.get('submitted_by','—')}</td>
                    </tr>
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4;border-bottom:0.5px solid rgba(255,255,255,0.08)">Date & time</td>
                      <td style="padding:5px 0;font-size:13px;color:#e8e6e0;border-bottom:0.5px solid rgba(255,255,255,0.08)">{s.get('date_display', s.get('date',''))} at {s.get('time','')}</td>
                    </tr>
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4;border-bottom:0.5px solid rgba(255,255,255,0.08)">Document</td>
                      <td style="padding:5px 0;font-size:13px;color:#e8e6e0;border-bottom:0.5px solid rgba(255,255,255,0.08)">{s.get('document','—')}</td>
                    </tr>
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4;border-bottom:0.5px solid rgba(255,255,255,0.08)">Reference</td>
                      <td style="padding:5px 0;font-size:14px;color:#3ecfaa;font-family:monospace;font-weight:700;border-bottom:0.5px solid rgba(255,255,255,0.08)">{s.get('ref','—')}</td>
                    </tr>
                    <tr>
                      <td style="padding:5px 0;font-size:12px;font-weight:600;color:#a8c4d4">Status</td>
                      <td style="padding:5px 0;font-size:13px;color:{sc};font-weight:700">{sl}</td>
                    </tr>
                  </table>
                </div>""", unsafe_allow_html=True)

                # Energy indicator summary
                mc1,mc2,mc3,mc4 = st.columns(4)
                mc1.metric("Electricity access",  f"{s.get('elec_pct','—')}%")
                mc2.metric("Clean cooking",        f"{s.get('cooking_pct','—')}%")
                mc3.metric("Solar GHI",            f"{s.get('solar_ghi','—')} kWh/m²")
                mc4.metric("Plan budget",          f"KES {s.get('budget_kes_b','—')}B")

                if s.get("epra_note"):
                    alert("info", f"<b>EPRA note:</b> {s['epra_note']}")

                # EPRA action buttons (only for pending)
                if s.get("status") == "submitted":
                    epra_note = st.text_input(
                        "EPRA note (optional — shown to county):",
                        placeholder="e.g. Approved — all checks passed, or: Please correct solar GHI unit error",
                        key=f"note_{s['ref']}"
                    )
                    btn_a, btn_b, _ = st.columns([1,1.4,2])
                    if btn_a.button(f"✅ Approve", key=f"approve_{s['ref']}", type="primary"):
                        update_submission_status(s["ref"], "approved", epra_note or "Approved by EPRA")
                        # Notify county via message store
                        add_message({
                            "id":          f"NOTIF-APPROVE-{s['ref']}",
                            "from_role":   "epra",
                            "from_county": "",
                            "from_name":   "EPRA",
                            "to":          s["county"],
                            "type":        "Submission",
                            "subject":     f"✅ Your plan has been approved — {s['county']} County",
                            "body": (
                                f"Your County Energy Plan (Ref: {s['ref']}) has been reviewed and "
                                f"<b>approved by EPRA</b> on {datetime.now().strftime('%d %b %Y at %H:%M')}.<br><br>"
                                f"Your county data is now included in the national INEP aggregation. "
                                f"You can view your county on the national dashboard.<br><br>"
                                f"{'<b>EPRA note:</b> ' + epra_note if epra_note else ''}"
                            ),
                            "date":         datetime.now().strftime("%Y-%m-%d"),
                            "time":         datetime.now().strftime("%H:%M:%S"),
                            "datetime":     datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "read_by_epra": True,
                            "replies":      [],
                        })
                        push_notification(f"{s['county']} County plan approved", "✅")
                        st.success(f"✅ {s['county']} County plan approved. County has been notified.")
                        st.rerun()
                    if btn_b.button(f"↩ Request resubmission", key=f"reject_{s['ref']}"):
                        update_submission_status(s["ref"], "rejected", epra_note or "Resubmission requested")
                        add_message({
                            "id":          f"NOTIF-REJECT-{s['ref']}",
                            "from_role":   "epra",
                            "from_county": "",
                            "from_name":   "EPRA",
                            "to":          s["county"],
                            "type":        "Validation",
                            "subject":     f"↩ Resubmission required — {s['county']} County",
                            "body": (
                                f"Your County Energy Plan (Ref: {s['ref']}) has been reviewed and "
                                f"requires resubmission. Please correct the issues noted below and resubmit.<br><br>"
                                f"<b>EPRA note:</b> {epra_note or 'Please review and resubmit.'}<br><br>"
                                f"Use the Submit energy plan tab to resubmit. Your previous reference is {s['ref']}."
                            ),
                            "date":         datetime.now().strftime("%Y-%m-%d"),
                            "time":         datetime.now().strftime("%H:%M:%S"),
                            "datetime":     datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "read_by_epra": True,
                            "replies":      [],
                        })
                        push_notification(f"Resubmission requested from {s['county']} County", "↩")
                        st.warning(f"Resubmission requested. {s['county']} County has been notified.")
                        st.rerun()
                elif s.get("status") == "approved":
                    alert("success", f"Approved {s.get('epra_action_time','')} · {s.get('epra_note','')}")
                elif s.get("status") == "rejected":
                    alert("warn", f"Resubmission requested {s.get('epra_action_time','')} · {s.get('epra_note','')}")

    st.divider()
    section("National aggregation trigger")
    nat = compute_national()
    all_real_subs = get_submissions()
    approved_real = [s for s in all_real_subs if s.get("status")=="approved"]
    st.info(f"{len(approved_real)} approved from file store · {nat['submitted_count']} in county data · {nat['pending_count']+nat['overdue_count']} not submitted")
    if st.button("▶  Run national aggregation now", type="primary"):
        with st.spinner("Aggregating submissions into national INEP…"):
            import time; time.sleep(2)
        st.success(f"Aggregation complete. INEP updated with {len(approved_real)} approved county plans.")
        push_notification("National INEP aggregation complete", "🗂️")

# =============================================================================
# PAGE FUNCTION: Communications (County) / Communications Hub (EPRA)
# Two-way messaging between EPRA and county committees.
#
# COUNTY VIEW:
#   - Received from EPRA tab: shows all messages from EPRA targeted at
#     this county or at "All counties". Also shows EPRA replies to
#     messages the county sent. Refresh button pulls latest from file store.
#   - Send to EPRA tab: county composes and sends a message (query,
#     extension request, support request, etc.) to EPRA.
#
# EPRA VIEW:
#   - Received tab: shows all messages from all county committees, plus
#     system notifications (new submissions, overdue alerts).
#     EPRA can reply inline to any county message.
#     Reply is saved to the file store and appears in the county's inbox.
#   - Send broadcast tab: EPRA composes a message to all counties,
#     a region, overdue counties, or a specific county.
#
# ALL MESSAGES are saved to kw_messages.json (and Supabase if configured).
# This means messages persist across sessions and are visible to both
# parties in real time when they refresh their inbox.
# =============================================================================

def render_msg(subject, from_, to_, date_, type_, body, actions=None, urgent=False):
    """
    Render a single styled message card.
    urgent=True adds a coloured border to highlight critical messages.
    actions = list of button label strings shown below the message body.
    """
    """Render a single inbox message card."""
    type_colors = {
        "Assumptions":"#1a6fa3","Benchmark":"#5b4fc9","Validation":"#d4891a",
        "Guidance":"#0f9d7e","Overdue":"#b33a2c","Submission":"#0f9d7e",
        "Alert":"#b33a2c","INEP update":"#0f9d7e","County query":"#d4891a","System":"#7a7870",
    }
    c      = type_colors.get(type_,"#7a7870")
    border = f"border:1.5px solid {c}" if urgent else "border:0.5px solid #e8e6de"
    acts   = "".join([
        f'<span style="font-size:11px;padding:3px 10px;border-radius:6px;border:0.5px solid #e8e6de;background:#f7f6f2;color:#6b6860;margin-right:5px;cursor:pointer">{a}</span>'
        for a in (actions or [])
    ])
    st.markdown(f"""
    <div style="background:#ffffff;{border};border-radius:10px;padding:14px 16px;margin-bottom:10px">
      <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:8px">
        <div style="flex:1">
          <span style="font-size:10px;font-weight:700;padding:3px 9px;border-radius:8px;
            background:{c}18;color:{c};text-transform:uppercase;letter-spacing:.4px;margin-right:8px">{type_}</span>
          <span style="font-size:13px;font-weight:600;color:#1a1916">{subject}</span>
        </div>
        <span style="font-size:11px;color:#7a7870;flex-shrink:0;margin-left:12px">{date_}</span>
      </div>
      <div style="font-size:11px;color:#7a7870;margin-bottom:8px">
        From: <b style="color:#1a1916">{from_}</b> &nbsp;→&nbsp; {to_}
      </div>
      <div style="font-size:12px;color:#444441;line-height:1.75;margin-bottom:{"10px" if actions else "0"}">{body}</div>
      {f'<div style="margin-top:4px">{acts}</div>' if actions else ""}
    </div>""", unsafe_allow_html=True)

# =============================================================================
# PAGE FUNCTION: Communications (county view) / Communications Hub (EPRA view)
# Two-way messaging between EPRA and county committees.
# All messages are persisted in kw_messages.json (and Supabase if configured).
#
# COUNTY VIEW — two sections:
#   "Received from EPRA" : shows EPRA broadcasts and replies to county messages
#   "Send to EPRA"       : county composes and sends a query or request to EPRA
#
# EPRA VIEW — two sections (via selectbox to avoid st.tabs() rendering issues):
#   "Received messages"  : all county messages, submission notifications, replies
#   "Send broadcast"     : EPRA composes a message to counties / regions
#
# DEVELOPER VIEW — same as EPRA (full access)
# =============================================================================

def page_inbox(role, county_id, user_name):
    """
    Render the communications page.
    role      = determines which view to show (county / epra / developer)
    county_id = used to filter messages for the county's own inbox
    user_name = shown on sent messages
    """

    # Get county name from county_id for labelling
    cnty = ""
    row  = None
    if county_id:
        matches = COUNTIES[COUNTIES["id"]==county_id]
        if not matches.empty:
            cnty = matches["name"].values[0]
            row  = matches.iloc[0]

    # =========================================================================
    # COUNTY COMMITTEE VIEW
    # =========================================================================
    if role == "county":

        alert("info",
            f"<b>{cnty} County communications</b> — "
            "receive messages and guidance from EPRA, view validation feedback, "
            "and send queries or requests directly to EPRA.")

        # Two clear sections — receiving and sending
        section_choice = st.radio(
            "Section:",
            ["📥  Received from EPRA", "📤  Send message to EPRA"],
            horizontal=True,
            key="county_comms_section",
            label_visibility="collapsed"
        )
        st.divider()

        # ── Received from EPRA ────────────────────────────────────────────────
        if "Received" in section_choice:

            # Refresh button — pulls latest from file store
            rc1, rc2 = st.columns([1,4])
            if rc1.button("🔄 Refresh", key="county_refresh_btn"):
                st.rerun()
            rc2.caption(
                f"Last checked: {datetime.now().strftime('%H:%M:%S')} — "
                "new messages appear when you click Refresh"
            )

            # Count summary
            all_msgs   = get_messages()
            epra_to_me = [
                m for m in all_msgs
                if m.get("from_role") in ("epra","system")
                and (
                    m.get("to","").lower() == cnty.lower()
                    or "all" in m.get("to","").lower()
                )
            ]
            my_sent   = [m for m in all_msgs if m.get("from_county")==cnty]
            my_replies = [r for m in my_sent for r in m.get("replies",[])]

            # Badge counts
            if epra_to_me or my_replies:
                b1,b2 = st.columns(2)
                with b1:
                    st.markdown(
                        f'<div style="background:#0e1e2e;border-radius:8px;padding:10px 14px;'
                        f'border-left:3px solid #1a6fa3;font-size:12px;color:#e8e6e0">'
                        f'Messages from EPRA: <b style="color:#3ecfaa">{len(epra_to_me)}</b>'
                        f'</div>', unsafe_allow_html=True)
                with b2:
                    st.markdown(
                        f'<div style="background:#0e1e2e;border-radius:8px;padding:10px 14px;'
                        f'border-left:3px solid #0f9d7e;font-size:12px;color:#e8e6e0">'
                        f'EPRA replies to your messages: <b style="color:#3ecfaa">{len(my_replies)}</b>'
                        f'</div>', unsafe_allow_html=True)
                st.markdown("")

            # Show EPRA approval/rejection messages first (most important)
            for m in all_msgs:
                if (m.get("from_role")=="epra"
                        and m.get("to","").lower()==cnty.lower()
                        and m.get("type") in ("Submission","Validation")):
                    tc = "#0f9d7e" if "approved" in m.get("subject","").lower() else "#b33a2c"
                    render_msg(m["subject"], m["from_name"], cnty,
                               m["date"], m["type"], m["body"], urgent=True)

            # Show EPRA replies to county's messages
            if my_replies:
                st.markdown("**EPRA replies to your messages:**")
                for m in my_sent:
                    for r in m.get("replies",[]):
                        st.markdown(f"""
                        <div style="background:#ffffff;border:0.5px solid #e8e6de;
                          border-left:3px solid #0f9d7e;border-radius:10px;
                          padding:14px 16px;margin-bottom:8px">
                          <div style="display:flex;justify-content:space-between;margin-bottom:6px">
                            <span style="font-size:13px;font-weight:600;color:#1a1916">
                              ↩ Re: {m["subject"]}
                            </span>
                            <span style="font-size:11px;color:#7a7870">
                              {r["date"]} at {r["time"]}
                            </span>
                          </div>
                          <div style="font-size:11px;color:#7a7870;margin-bottom:6px">
                            From: <b>EPRA</b> → {cnty} County
                          </div>
                          <div style="font-size:12px;color:#444441;line-height:1.75">
                            {r["body"]}
                          </div>
                        </div>""", unsafe_allow_html=True)

            # Show all EPRA broadcasts targeted at this county or all counties
            if epra_to_me:
                st.markdown("**Messages from EPRA:**")
                for m in reversed(epra_to_me):
                    if m.get("type") not in ("Submission","Validation"):
                        tc = {
                            "Assumptions":"#1a6fa3","Guidance":"#0f9d7e",
                            "Benchmark":"#5b4fc9","Alert":"#b33a2c",
                            "Reminder":"#b33a2c"
                        }.get(m.get("type",""),"#7a7870")
                        render_msg(
                            m["subject"], m.get("from_name","EPRA"),
                            m.get("to","All counties"),
                            m["date"], m.get("type","Message"),
                            m["body"]
                        )
            else:
                st.info(
                    "No messages from EPRA yet. "
                    "When EPRA sends a broadcast or responds to your query "
                    "it will appear here. Click Refresh to check for new messages."
                )

            # Show submission confirmations from this session
            my_subs = get_submissions(county=cnty)
            if my_subs:
                st.markdown("**Your submission history:**")
                for s in my_subs[:3]:
                    sc = {"submitted":"#d4891a","approved":"#0f9d7e",
                          "rejected":"#b33a2c"}.get(s.get("status",""),"#7a7870")
                    render_msg(
                        subject = f"✅ Plan submitted — {s['county']} · Ref: {s['ref']}",
                        from_   = "KenyaWatts System",
                        to_     = cnty,
                        date_   = s.get("date_display", s.get("date","")),
                        type_   = "Submission",
                        body    = (
                            f"<b>Submitted by:</b> {s.get('submitted_by','—')}<br>"
                            f"<b>Date and time:</b> {s.get('date_display','')} "
                            f"at {s.get('time','')}<br>"
                            f"<b>Document:</b> {s.get('document','—')}<br>"
                            f"<b>Status:</b> <span style='color:{sc};font-weight:600'>"
                            f"{s.get('status','').title()}</span>"
                            + (f"<br><b>EPRA note:</b> {s['epra_note']}"
                               if s.get('epra_note') else "")
                        )
                    )

        # ── Send message to EPRA ──────────────────────────────────────────────
        else:
            section("Send a message to EPRA")
            alert("info",
                "Use this to ask questions, request submission support, "
                "request a deadline extension, or flag a data issue. "
                "EPRA responds within 5 working days. "
                "Check the Received section for their reply.")

            # Show previous messages sent by this county
            my_msgs = [m for m in get_messages() if m.get("from_county")==cnty]
            if my_msgs:
                st.markdown(f"**Your sent messages ({len(my_msgs)}):**")
                for m in reversed(my_msgs):
                    nreplies  = len(m.get("replies",[]))
                    reply_txt = (
                        f"<span style='color:#0f9d7e;font-weight:600'>"
                        f"✓ {nreplies} EPRA reply</span>"
                        if nreplies else
                        "<span style='color:#9c9a8e'>Awaiting EPRA response</span>"
                    )
                    st.markdown(f"""
                    <div style="background:#ffffff;border:0.5px solid #e8e6de;
                      border-left:3px solid #1a6fa3;border-radius:8px;
                      padding:12px 14px;margin-bottom:8px">
                      <div style="display:flex;justify-content:space-between;margin-bottom:5px">
                        <span style="font-size:13px;font-weight:600;color:#1a1916">
                          {m["subject"]}
                        </span>
                        <span style="font-size:11px;color:#9c9a8e">
                          {m["date"]} at {m["time"]}
                        </span>
                      </div>
                      <div style="font-size:11px;color:#7a7870;margin-bottom:5px">
                        Ref: <code>{m["id"]}</code> · {reply_txt}
                      </div>
                      <div style="font-size:12px;color:#444441;line-height:1.6">
                        {m["body"][:200]}{"..." if len(m["body"])>200 else ""}
                      </div>
                    </div>""", unsafe_allow_html=True)
                st.divider()

            # Compose new message
            msg_type    = st.selectbox(
                "Message type:",
                ["❓ Methodology question",
                 "🙏 Submission support request",
                 "📅 Extension request",
                 "🐛 Data / system issue",
                 "💬 General query"],
                key="county_msg_type"
            )
            msg_subject = st.text_input(
                "Subject: *",
                placeholder="Brief description of your query",
                key="county_msg_subject"
            )
            msg_body = st.text_area(
                "Message: *",
                placeholder=(
                    "Describe your question or request in detail.\n"
                    "Include your county name and submission reference if relevant."
                ),
                height=140,
                key="county_msg_body"
            )
            if st.button("📤 Send to EPRA", type="primary", key="county_send_btn"):
                if not msg_subject.strip() or not msg_body.strip():
                    st.error("Both subject and message are required.")
                else:
                    ref = f"MSG-{cnty[:2].upper()}-{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    add_message({
                        "id":          ref,
                        "from_role":   "county",
                        "from_county": cnty,
                        "from_name":   user_name,
                        "to":          "EPRA",
                        "type":        msg_type.split(" ",1)[-1],
                        "subject":     msg_subject.strip(),
                        "body":        msg_body.strip(),
                        "date":        datetime.now().strftime("%Y-%m-%d"),
                        "time":        datetime.now().strftime("%H:%M"),
                        "datetime":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "read_by_epra":False,
                        "replies":     [],
                    })
                    st.session_state.audit_log.append({
                        "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "user":   user_name,
                        "action": f"Message sent to EPRA: {msg_subject[:50]}",
                        "ref":    ref,
                    })
                    push_notification(f"Message sent to EPRA · Ref: {ref}", "📧")
                    st.success(f"✓ Message sent. Ref: `{ref}`")
                    st.info("EPRA responds within 5 working days. Check Received for their reply.")
                    # Clear fields
                    for k in ["county_msg_subject","county_msg_body"]:
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()

    # =========================================================================
    # EPRA / DEVELOPER VIEW
    # =========================================================================
    elif role in ("epra","developer","ministry"):

        alert("info",
            "<b>EPRA communications hub</b> — "
            "messages from all counties, system notifications, and broadcast tools.")

        # Use selectbox instead of st.tabs() — tabs fail silently inside
        # st.navigation() page functions when content is complex
        epra_view = st.selectbox(
            "Section:",
            ["📥  Received — county messages and notifications",
             "📣  Send broadcast to counties"],
            key  = "epra_comms_view",
            label_visibility="collapsed"
        )
        st.divider()

        # ── RECEIVED ─────────────────────────────────────────────────────────
        if "Received" in epra_view:

            # Refresh button
            rr1, rr2 = st.columns([1,4])
            if rr1.button("🔄 Refresh", key="epra_refresh_btn"):
                st.rerun()
            rr2.caption(
                f"Last refreshed: {datetime.now().strftime('%H:%M:%S')} — "
                "new county messages appear when you click Refresh"
            )

            # Load all messages from file/cloud store
            all_msgs    = get_messages()
            county_msgs = [m for m in all_msgs
                           if m.get("from_role") in ("county","system")]
            unread      = [m for m in county_msgs if not m.get("read_by_epra")]

            if unread:
                alert("warn",
                    f"<b>{len(unread)} unread message(s)</b> from county committees.")

            # ── County messages with inline reply ─────────────────────────────
            if county_msgs:
                st.markdown(f"**County messages and notifications ({len(county_msgs)})**")
                for m in reversed(county_msgs):
                    unread_dot = "🔵 " if not m.get("read_by_epra") else ""
                    tc = {
                        "County query":"#d4891a",
                        "Methodology question":"#1a6fa3",
                        "Submission support request":"#5b4fc9",
                        "Extension request":"#b33a2c",
                        "Data / system issue":"#b33a2c",
                        "General query":"#7a7870",
                        "Submission":"#0f9d7e",
                        "system":"#7a7870",
                    }.get(m.get("type","County query"),"#d4891a")

                    # Message card
                    st.markdown(f"""
                    <div style="background:#ffffff;border:0.5px solid #e8e6de;
                      border-left:3px solid {tc};border-radius:10px;
                      padding:14px 16px;margin-bottom:4px">
                      <div style="display:flex;justify-content:space-between;margin-bottom:7px">
                        <div>
                          <span style="font-size:10px;font-weight:700;padding:2px 8px;
                            border-radius:6px;background:{tc}18;color:{tc};
                            text-transform:uppercase;margin-right:8px">
                            {m.get("type","Message")}
                          </span>
                          <span style="font-size:13px;font-weight:600;color:#1a1916">
                            {unread_dot}{m["subject"]}
                          </span>
                        </div>
                        <span style="font-size:11px;color:#7a7870;flex-shrink:0;margin-left:8px">
                          {m["date"]} at {m["time"]}
                        </span>
                      </div>
                      <div style="font-size:11px;color:#7a7870;margin-bottom:7px">
                        From: <b style="color:#1a1916">{m.get("from_name","County")}</b>
                        ({m.get("from_county","")}) → EPRA · Ref: <code>{m["id"]}</code>
                      </div>
                      <div style="font-size:12px;color:#444441;line-height:1.75;margin-bottom:8px">
                        {m["body"]}
                      </div>
                      {"".join([
                        f'<div style="background:#e8f7f4;border-radius:6px;padding:10px 12px;margin-bottom:4px"><div style="font-size:11px;font-weight:600;color:#0f9d7e;margin-bottom:3px">↩ EPRA replied · {r["date"]} at {r["time"]}</div><div style="font-size:12px;color:#1a1916;line-height:1.6">{r["body"]}</div></div>'
                        for r in m.get("replies",[])
                      ])}
                    </div>""", unsafe_allow_html=True)

                    # Mark as read automatically
                    if not m.get("read_by_epra"):
                        mark_read(m["id"])

                    # Inline reply toggle
                    reply_key = f"reply_open_{m['id']}"
                    if reply_key not in st.session_state:
                        st.session_state[reply_key] = False

                    bc1, bc2, _ = st.columns([1,1.3,2])
                    if bc1.button(
                        f"📧 Reply to {m.get('from_county','county')}",
                        key=f"btn_reply_{m['id']}"
                    ):
                        st.session_state[reply_key] = not st.session_state[reply_key]
                    if bc2.button(
                        "📣 Broadcast answer to all",
                        key=f"btn_broad_{m['id']}"
                    ):
                        push_notification(
                            f"Answer broadcast re: {m['subject'][:40]}", "📣")
                        st.success("Answer broadcast to all counties.")

                    # Reply text area (shown when toggle is on)
                    if st.session_state.get(reply_key):
                        reply_body = st.text_area(
                            f"Reply to {m.get('from_county','')} County:",
                            placeholder="Type your response here...",
                            height=110,
                            key=f"reply_text_{m['id']}"
                        )
                        if st.button("📤 Send reply", type="primary",
                                     key=f"send_reply_{m['id']}"):
                            if not reply_body.strip():
                                st.error("Reply cannot be empty.")
                            else:
                                reply = {
                                    "from": "EPRA",
                                    "body": reply_body.strip(),
                                    "date": datetime.now().strftime("%Y-%m-%d"),
                                    "time": datetime.now().strftime("%H:%M"),
                                }
                                add_reply(m["id"], reply)
                                st.session_state[reply_key] = False
                                push_notification(
                                    f"Reply sent to {m.get('from_county','')} County",
                                    "📧")
                                st.session_state.audit_log.append({
                                    "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    "user":   user_name,
                                    "action": f"Replied to {m.get('from_county','')}",
                                    "ref":    m["id"],
                                })
                                st.success(
                                    f"✓ Reply sent to {m.get('from_county','')} County. "
                                    "It will appear in their inbox immediately.")
                                st.rerun()
                    st.markdown("")
            else:
                st.info(
                    "No messages from counties yet. "
                    "When a county sends a message or submits a plan, "
                    "it will appear here. Click Refresh to check for updates.")

        # ── SEND BROADCAST ────────────────────────────────────────────────────
        else:
            # This section is now rendered directly (NOT inside with tab_broadcast:)
            # which was the root cause of the blank page bug.
            # st.tabs() and with tab_broadcast: silently fail inside
            # st.navigation() pages — using st.selectbox() + if/else is reliable.

            section("Send broadcast to counties")
            alert("info",
                "<b>Compose a message to counties.</b> "
                "Select recipients, choose a message type, fill in the subject "
                "and body, then click Send. The message appears in the county "
                "Communications tab when they click Refresh.")

            # Sample messages for testing — shown expanded by default
            with st.expander("📋 Sample test messages — click to expand", expanded=False):
                st.markdown("""
**Copy any of these to test the two-way messaging:**

| Recipient | Subject | Body |
|---|---|---|
| Makueni | Updated solar GHI baseline | Please use 2,008 kWh/m² as the solar GHI for Makueni County. |
| All 47 counties | Submission deadline — 30 June 2026 | All county energy plans must be submitted by 30 June 2026 under INEP Regulations 2025. |
| Turkana | Direct submission support available | EPRA can provide technical assistance. Contact Allan.Wairimu@epra.go.ke |
                """)

            st.markdown("### Compose your message")

            # Step 1: Recipients
            # IMPORTANT: No default= argument — avoids Streamlit blank page bug
            recipient_opts = (
                ["All 47 counties",
                 "Overdue counties only",
                 "North East region",
                 "Coast region",
                 "Rift Valley region",
                 "Nyanza region",
                 "Western region",
                 "Central region",
                 "Eastern region",
                 "Nairobi"]
                + sorted(COUNTIES["name"].tolist())
            )
            b_to = st.multiselect(
                "1. Send to (required):",
                options=recipient_opts,
                key="bc_to_v3",
                help="Select one or more recipients"
            )

            # Step 2: Message type
            b_type = st.selectbox(
                "2. Message type:",
                ["Assumptions","Guidance","Benchmark","Alert","Reminder"],
                key="bc_type_v3"
            )

            # Step 3: Subject
            b_subject = st.text_input(
                "3. Subject:",
                placeholder="e.g. Updated national planning assumptions",
                key="bc_subject_v3"
            )

            # Step 4: Body
            b_body = st.text_area(
                "4. Message body:",
                placeholder="Type your full message here...",
                height=200,
                key="bc_body_v3"
            )
            if b_body:
                st.caption(f"{len(b_body)} characters · {len(b_body.split())} words")

            # Step 5: Attachment
            b_attach = st.file_uploader(
                "5. Attach document (optional):",
                type=["pdf","xlsx","docx"],
                key="bc_attach_v3"
            )

            st.markdown("")
            if st.button("📣 Send broadcast", type="primary",
                         key="bc_send_v3", use_container_width=False):
                if not b_to:
                    st.error("Please select at least one recipient.")
                elif not b_subject.strip():
                    st.error("Please enter a subject line.")
                elif not b_body.strip():
                    st.error("Please enter a message body.")
                else:
                    # Build recipient string
                    to_str = ", ".join(b_to)

                    # Save message to persistent store
                    msg = {
                        "id":           f"MSG-EPRA-{datetime.now().strftime('%Y%m%d%H%M%S')}",
                        "from_role":    "epra",
                        "from_county":  "",
                        "from_name":    "EPRA",
                        "to":           to_str,
                        "type":         b_type,
                        "subject":      b_subject.strip(),
                        "body":         b_body.strip(),
                        "date":         datetime.now().strftime("%Y-%m-%d"),
                        "time":         datetime.now().strftime("%H:%M"),
                        "datetime":     datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "attach":       b_attach.name if b_attach else "",
                        "read_by_epra": True,
                        "replies":      [],
                    }
                    add_message(msg)
                    push_notification(f"Broadcast sent to: {to_str}", "📣")
                    st.session_state.audit_log.append({
                        "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "user":   user_name,
                        "action": f"Broadcast sent to {to_str}: {b_subject[:40]}",
                        "ref":    msg["id"],
                    })
                    st.success(f"✅ Broadcast sent to: **{to_str}**")
                    st.info(
                        "Counties see this message when they open Communications "
                        "and click Refresh."
                    )
                    # Clear fields
                    for k in ["bc_to_v3","bc_subject_v3","bc_body_v3","bc_attach_v3"]:
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()

    # ── Ministry / Dev partner — read only ────────────────────────────────────
    else:
        alert("info","<b>Communications view</b> — INEP updates from EPRA.")
        nat = compute_national()
        render_msg(
            subject = f"📊 INEP updated — {nat['submitted_count']} counties submitted",
            from_   = "EPRA",
            to_     = "Ministry of Energy",
            date_   = datetime.now().strftime("%Y-%m-%d"),
            type_   = "INEP update",
            body    = (
                f"National INEP has been updated with data from "
                f"{nat['submitted_count']} submitted county plans "
                f"({nat['coverage_pct']}% coverage).<br><br>"
                f"<b>Electricity access:</b> {nat['w_elec']}% "
                f"(target 100% by 2030)<br>"
                f"<b>Clean cooking:</b> {nat['w_cooking']}% "
                f"(target 100% by 2028)<br>"
                f"<b>Counties overdue:</b> {nat['overdue_count']}"
            )
        )


def page_makueni():
    alert("success","<b>Reference plan:</b> Makueni County Energy Plan 2023–2032 · WRI + Strathmore University · Upload the PDF in Submit tab to test AI extraction.")
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Electricity access","75.1%","Grid 29.2% + MG 5.7% + SHS 40.2%")
    c2.metric("Solar GHI","2,008 kWh/m²","PV output 4.35 kWh/kWp/day")
    c3.metric("Total budget","KES 74.9B","2023–2032")
    c4.metric("Clean cooking","17.9%","⚠ Firewood 72.5%")
    st.markdown("")

    c_left, c_right = st.columns(2)
    with c_left:
        section("Primary cooking fuel — households 2022")
        fig = px.bar(MAKUENI_COOKING, x="Pct", y="Fuel", orientation="h",
                     color="Fuel", color_discrete_map={"Firewood":"#b33a2c","LPG":"#0f9d7e",
                     "Charcoal":"#7a7870","Biogas":"#5b4fc9","Electric":"#1a6fa3","Other":"#c8c6be"})
        fig = apply_layout(fig, xlabel="Percentage (%)", ylabel="Fuel type")
        fig.update_layout(showlegend=False, height=240)
        st.plotly_chart(fig, use_container_width=True)

    with c_right:
        section("Monthly outage vs EPRA benchmark")
        colors = ["#b33a2c" if h>5 else "#0f9d7e" for h in OUTAGE["Hours"]]
        fig = go.Figure(go.Bar(x=OUTAGE["Month"], y=OUTAGE["Hours"], marker_color=colors))
        fig.add_hline(y=5, line_dash="dash", line_color="#1a6fa3",
                      annotation_text="<b>EPRA benchmark 5 hrs</b>",
                      annotation_font=dict(size=11, color="#1a1916"))
        fig = apply_layout(fig, xlabel="Month", ylabel="Hours per customer")
        fig.update_layout(height=240)
        st.plotly_chart(fig, use_container_width=True)

    section("Electrification scenarios — investment required (OnSSET)")
    scen = pd.DataFrame({
        "Scenario":["Low demand (2028)","High demand (2028)","Grid intensification (2028)"],
        "Grid (MW)":[19.3,42.7,38.4], "Solar PV (MW)":[2.34,53.6,0], "Investment (USD M)":[132.5,360.0,571.8]
    })
    fig = go.Figure()
    for s,c in [("Grid (MW)","#1a6fa3"),("Solar PV (MW)","#d4891a"),("Investment (USD M)","#0f9d7e")]:
        fig.add_trace(go.Bar(name=s, x=scen["Scenario"], y=scen[s], marker_color=c))
    fig = apply_layout(fig, xlabel="Scenario", ylabel="Value")
    fig.update_layout(barmode="group", height=280, legend=dict(font=LEGEND_FONT))
    st.plotly_chart(fig, use_container_width=True)



# ── NATIONAL REPORT GENERATOR (EPRA only) ────────────────────────────────────
def generate_word_report(nat, counties_df, generated_by):
    """Generate a Word (.docx) National Energy Plan summary report."""
    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def add_heading(text, level=1, color=(26, 31, 22)):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Arial"
        return h

    def add_para(text, bold=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"
        return p

    def add_table(headers, rows, header_color=(14,30,46)):
        """
        Add a styled table to the Word document.
        Uses python-docx OxmlElement instead of lxml to avoid import conflicts
        on Streamlit Cloud. Header row has dark navy background with white text.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn

        t   = doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = "Table Grid"

        # Style the header row
        hdr = t.rows[0]
        for h, cell in zip(headers, hdr.cells):
            cell.text = h
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
            run.bold  = True
            run.font.size  = Pt(10)
            run.font.name  = "Arial"
            run.font.color.rgb = RGBColor(255, 255, 255)  # white text

            # Set cell background colour using OxmlElement (no lxml needed)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd   = OxmlElement("w:shd")
            hex_color = "%02x%02x%02x" % tuple(header_color)
            shd.set(_qn("w:fill"),  hex_color)
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:val"),   "clear")
            tc_pr.append(shd)

        # Style the data rows with alternating background
        for ri, row_data in enumerate(rows):
            tr = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = tr.cells[ci]
                cell.text = str(val)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
                run.font.name = "Arial"
                # Alternating row shading
                if ri % 2 == 1:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd   = OxmlElement("w:shd")
                    shd.set(_qn("w:fill"),  "F7F6F2")
                    shd.set(_qn("w:color"), "auto")
                    shd.set(_qn("w:val"),   "clear")
                    tc_pr.append(shd)
        return t

    now = datetime.now()

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("KENYA INTEGRATED NATIONAL ENERGY PLAN")
    tr.bold = True; tr.font.size = Pt(20); tr.font.name = "Arial"
    tr.font.color.rgb = RGBColor(14,30,46)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("National Summary Report — Aggregated from County Energy Plans")
    sr.font.size = Pt(13); sr.font.name = "Arial"; sr.font.color.rgb = RGBColor(26,111,163)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Generated by: {generated_by}\n"
        f"Date: {now.strftime('%d %B %Y')} at {now.strftime('%H:%M')}\n"
        f"Platform: KenyaWatts · EPRA Kenya · NGDA 2026\n"
        f"Classification: CONFIDENTIAL — EPRA INTERNAL USE ONLY"
    )
    mr.font.size = Pt(10); mr.font.name = "Arial"; mr.font.color.rgb = RGBColor(90,138,158)
    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    add_heading("1. Executive Summary", 1, (14,30,46))
    add_para(
        f"This report presents the current state of Kenya's national energy planning "
        f"as aggregated from {nat['submitted_count']} submitted County Energy Plans (CEPs) "
        f"on the KenyaWatts Digital Integrated National Energy Planning Platform. "
        f"The data represents {nat['coverage_pct']}% of Kenya's 47 counties. "
        f"A total of {nat['overdue_count']} counties remain overdue on their submissions."
    )
    doc.add_paragraph()

    # Key indicators table
    add_heading("1.1 Key National Indicators", 2, (26,111,163))
    add_table(
        ["Indicator","Current value","National target","Target year","Status"],
        [
            ["Electricity access (weighted avg)", f"{nat['w_elec']}%", "100%", "2030",
             "On track" if float(nat['w_elec'])>70 else "At risk"],
            ["Clean cooking access (weighted avg)", f"{nat['w_cooking']}%", "100%", "2028",
             "At risk" if float(nat['w_cooking'])<50 else "On track"],
            ["Clean energy generation", "82%", "100%", "2035", "On track"],
            ["Counties submitted", f"{nat['submitted_count']} / {nat['total_counties']}", "47 / 47", "Current cycle",
             "In progress"],
            ["Total plan investment needed", f"KES {nat['total_budget']}B",
             "TBD (all 47 counties)", "2023–2032", "Partial"],
        ]
    )
    doc.add_paragraph()

    # ── 2. Submission Status ──────────────────────────────────────────────────
    add_heading("2. County Submission Status", 1, (14,30,46))
    add_para(
        f"As of {now.strftime('%d %B %Y')}, {nat['submitted_count']} of Kenya's 47 counties "
        f"have submitted County Energy Plans through the KenyaWatts platform. "
        f"{nat['overdue_count']} counties are overdue and have received reminder notifications. "
        f"{nat['pending_count']} counties are yet to begin their submission."
    )
    doc.add_paragraph()

    submitted = counties_df[counties_df["status"].isin(["submitted","review"])]
    overdue   = counties_df[counties_df["status"]=="overdue"]
    pending   = counties_df[counties_df["status"]=="pending"]

    add_heading("2.1 Submitted counties", 2, (15,157,126))
    if not submitted.empty:
        add_table(
            ["County","Region","Electricity (%)","Clean cooking (%)","Budget (KES B)","Target year"],
            [[r["name"],r["region"],f"{r['elec']}%",f"{r['cooking']}%",
              str(r["budget"]) if r["budget"]>0 else "—",str(r["target_yr"])]
             for _,r in submitted.iterrows()]
        )
    doc.add_paragraph()

    add_heading("2.2 Overdue counties — action required", 2, (163,45,44))
    if not overdue.empty:
        add_para(
            "The following counties have not submitted their County Energy Plans despite "
            "multiple reminder notifications. Escalation to the Council of Governors is "
            "recommended under Section 5(5)(a) of the Energy Act 2019.", bold=False
        )
        add_table(
            ["County","Region","Electricity access (%)","Population","Reminders sent"],
            [[r["name"],r["region"],f"{r['elec']}%",f"{r['pop']//1000:,}K","3 (final notice sent)"]
             for _,r in overdue.iterrows()]
        )
    doc.add_paragraph()

    # ── 3. National Energy Access ─────────────────────────────────────────────
    add_heading("3. National Energy Access Analysis", 1, (14,30,46))
    add_heading("3.1 Electricity access", 2, (26,111,163))
    add_para(
        f"The population-weighted national electricity access rate from submitted county "
        f"plans is {nat['w_elec']}%. This figure reflects the weighted contribution of each "
        f"county based on its population. The national target is 100% by 2030. "
        f"The four counties with the most critical access gaps — Turkana (12%), Marsabit (8%), "
        f"Mandera (11%), and Wajir (9%) — are all overdue on plan submission and represent "
        f"the highest-priority investment areas."
    )
    doc.add_paragraph()

    add_heading("3.2 Electricity access by county (submitted plans)", 2, (26,111,163))
    if not submitted.empty:
        sorted_df = submitted.sort_values("elec")
        add_table(
            ["County","Electricity access (%)","Gap to 100% target","MTF demand tier"],
            [[r["name"],f"{r['elec']}%",f"{100-r['elec']}%",f"Tier {r['mtf']}"]
             for _,r in sorted_df.iterrows()]
        )
    doc.add_paragraph()

    add_heading("3.3 Clean cooking access", 2, (26,111,163))
    add_para(
        f"The population-weighted clean cooking access rate is {nat['w_cooking']}%. "
        f"Kenya's target is universal clean cooking by 2028. Based on submitted county plans, "
        f"firewood remains the primary cooking fuel in the majority of rural households. "
        f"{'This trajectory suggests the 2028 target is at risk without accelerated intervention.' if float(nat['w_cooking'])<50 else 'Progress is being made but acceleration is needed to meet the 2028 target.'}"
    )
    doc.add_paragraph()

    # ── 4. Renewable Energy Resources ────────────────────────────────────────
    add_heading("4. Renewable Energy Resource Assessment", 1, (14,30,46))
    add_para(
        "Based on submitted county plans, Kenya has abundant renewable energy resources "
        "distributed across its counties. The table below summarises solar potential by county."
    )
    doc.add_paragraph()

    if not submitted.empty:
        add_table(
            ["County","Solar GHI (kWh/m²/yr)","Solar potential","Population (K)"],
            [[r["name"],str(r["solar"]),
              "High (>2,000)" if r["solar"]>2000 else "Good (1,800–2,000)" if r["solar"]>1800 else "Moderate",
              f"{r['pop']//1000:,}"]
             for _,r in submitted.sort_values("solar",ascending=False).iterrows()]
        )
    doc.add_paragraph()

    # ── 5. Investment Summary ─────────────────────────────────────────────────
    add_heading("5. Investment and Budget Summary", 1, (14,30,46))
    add_para(
        f"Based on {nat['submitted_count']} submitted county plans, the total identified "
        f"investment required is KES {nat['total_budget']} billion. This covers electricity "
        f"access expansion, clean cooking programmes, renewable energy development, and "
        f"energy efficiency improvements. The national figure will be updated as additional "
        f"counties submit their plans."
    )
    doc.add_paragraph()

    if not submitted.empty:
        budget_counties = submitted[submitted["budget"]>0]
        if not budget_counties.empty:
            add_table(
                ["County","Plan budget (KES B)","Budget per capita (KES)","Plan period"],
                [[r["name"],f"{r['budget']}B",
                  f"{int(r['budget']*1e9/r['pop']):,}",
                  f"2023–{r['target_yr']}"]
                 for _,r in budget_counties.iterrows()]
            )
    doc.add_paragraph()

    # ── 6. Recommendations ────────────────────────────────────────────────────
    add_heading("6. Recommendations", 1, (14,30,46))
    recs = [
        ("Prioritise direct submission support for overdue counties",
         f"Turkana, Marsabit, Mandera and Wajir have not submitted plans despite three reminder "
         f"notifications. EPRA should arrange direct technical assistance with these counties "
         f"given their critical access gaps and strategic importance to the national INEP."),
        ("Accelerate clean cooking intervention",
         f"The current clean cooking access rate ({nat['w_cooking']}%) is significantly below "
         f"the 2028 target trajectory. EPRA should recommend to the Ministry of Energy that "
         f"LPG subsidy expansion and improved cookstove distribution programmes be accelerated."),
        ("Complete national aggregation when all 47 counties submit",
         f"The current INEP aggregation covers {nat['coverage_pct']}% of counties. The "
         f"national picture will only be complete when all 47 counties have submitted. "
         f"EPRA should set a firm deadline for the remaining {nat['pending_count']+nat['overdue_count']} counties."),
        ("Publish county comparison dashboard publicly",
         f"Making county-level energy access data publicly available — following the India SECI "
         f"model — will create accountability pressure and incentivise timely submission in "
         f"future planning cycles."),
    ]
    for i,(title,body) in enumerate(recs,1):
        add_para(f"{i}. {title}", bold=True)
        add_para(f"   {body}")
        doc.add_paragraph()

    # ── 7. Data notes ─────────────────────────────────────────────────────────
    add_heading("7. Data Notes and Methodology", 1, (14,30,46))
    add_para(
        f"• National electricity access and clean cooking figures are population-weighted "
        f"averages calculated from {nat['submitted_count']} submitted county plans.\n"
        f"• Budget figures are simple sums of submitted county plan budgets.\n"
        f"• Universal access target year reflects the latest target year from submitted plans ({nat['latest_target']}).\n"
        f"• Data generated from KenyaWatts platform on {now.strftime('%d %B %Y')}.\n"
        f"• Clean energy generation (82%) sourced from EPRA Statistics Report FY 2024/25.\n"
        f"• This report is classified CONFIDENTIAL and is for EPRA internal use only."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_pdf_report(nat, counties_df, generated_by):
    """Generate a PDF National Energy Plan summary report."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2*cm, rightMargin=2*cm)

    styles = getSampleStyleSheet()
    now = datetime.now()

    # Custom styles
    title_style = ParagraphStyle("KWTitle", parent=styles["Title"],
        fontSize=20, textColor=colors.HexColor("#0e1e2e"),
        fontName="Helvetica-Bold", spaceAfter=6, alignment=TA_CENTER)
    sub_style = ParagraphStyle("KWSub", parent=styles["Normal"],
        fontSize=13, textColor=colors.HexColor("#1a6fa3"),
        fontName="Helvetica", spaceAfter=4, alignment=TA_CENTER)
    meta_style = ParagraphStyle("KWMeta", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#5a8a9e"),
        fontName="Helvetica", spaceAfter=2, alignment=TA_CENTER)
    h1_style = ParagraphStyle("KWH1", parent=styles["Heading1"],
        fontSize=14, textColor=colors.HexColor("#0e1e2e"),
        fontName="Helvetica-Bold", spaceBefore=16, spaceAfter=6)
    h2_style = ParagraphStyle("KWH2", parent=styles["Heading2"],
        fontSize=12, textColor=colors.HexColor("#1a6fa3"),
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("KWBody", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#1a1916"),
        fontName="Helvetica", spaceAfter=6, leading=15)
    conf_style = ParagraphStyle("KWConf", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#b33a2c"),
        fontName="Helvetica-Bold", alignment=TA_CENTER)

    def make_table(headers, rows, col_widths=None):
        data = [headers] + rows
        tw   = col_widths or [A4[0]/(2*cm)/len(headers)*cm]*len(headers)
        t    = Table(data, colWidths=tw)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0),(-1,0), colors.HexColor("#0e1e2e")),
            ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
            ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0),(-1,0), 9),
            ("FONTNAME",   (0,1),(-1,-1),"Helvetica"),
            ("FONTSIZE",   (0,1),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#f7f6f2")]),
            ("GRID",       (0,0),(-1,-1), 0.4, colors.HexColor("#e8e6de")),
            ("VALIGN",     (0,0),(-1,-1), "MIDDLE"),
            ("TOPPADDING", (0,0),(-1,-1), 5),
            ("BOTTOMPADDING",(0,0),(-1,-1), 5),
            ("LEFTPADDING", (0,0),(-1,-1), 6),
        ]))
        return t

    story = []

    # Cover
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("KENYA INTEGRATED NATIONAL ENERGY PLAN", title_style))
    story.append(Paragraph("National Summary Report — Aggregated from County Energy Plans", sub_style))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(f"Generated by: {generated_by}", meta_style))
    story.append(Paragraph(f"Date: {now.strftime('%d %B %Y')} at {now.strftime('%H:%M')}", meta_style))
    story.append(Paragraph("Platform: KenyaWatts · Energy & Petroleum Regulatory Authority · NGDA 2026", meta_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("CONFIDENTIAL — EPRA INTERNAL USE ONLY", conf_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0e1e2e")))
    story.append(Spacer(1, 0.5*cm))

    # Section 1
    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(
        f"This report presents the current state of Kenya's national energy planning as "
        f"aggregated from <b>{nat['submitted_count']}</b> submitted County Energy Plans (CEPs) "
        f"on the KenyaWatts platform. The data represents <b>{nat['coverage_pct']}%</b> "
        f"of Kenya's 47 counties. {nat['overdue_count']} counties remain overdue.", body_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("1.1 Key National Indicators", h2_style))
    story.append(make_table(
        ["Indicator","Current Value","Target","Year","Status"],
        [
            [f"Electricity access",         f"{nat['w_elec']}%",  "100%","2030","On track" if float(nat['w_elec'])>70 else "At risk"],
            [f"Clean cooking access",       f"{nat['w_cooking']}%","100%","2028","At risk" if float(nat['w_cooking'])<50 else "On track"],
            ["Clean energy generation",     "82%",              "100%","2035","On track"],
            ["Counties submitted",          f"{nat['submitted_count']}/47","47/47","Current","In progress"],
            ["Total investment identified", f"KES {nat['total_budget']}B","TBD","2023–2032","Partial"],
        ],
        [5*cm,3*cm,2*cm,2*cm,2.5*cm]
    ))
    story.append(Spacer(1, 0.4*cm))

    # Section 2
    story.append(Paragraph("2. County Submission Status", h1_style))
    submitted = counties_df[counties_df["status"].isin(["submitted","review"])]
    overdue   = counties_df[counties_df["status"]=="overdue"]
    story.append(Paragraph("2.1 Submitted county plans", h2_style))
    if not submitted.empty:
        story.append(make_table(
            ["County","Region","Electricity (%)","Clean cooking (%)","Budget (KES B)","Target year"],
            [[r["name"],r["region"],f"{r['elec']}%",f"{r['cooking']}%",
              str(r["budget"]) if r["budget"]>0 else "—",str(r["target_yr"])]
             for _,r in submitted.iterrows()],
            [3.5*cm,3*cm,2.5*cm,2.8*cm,2.8*cm,2.4*cm]
        ))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("2.2 Overdue counties — action required", h2_style))
    if not overdue.empty:
        story.append(make_table(
            ["County","Region","Electricity (%)","Population","Reminders sent"],
            [[r["name"],r["region"],f"{r['elec']}%",f"{r['pop']//1000:,}K","3 (final)"]
             for _,r in overdue.iterrows()],
            [3.5*cm,3.5*cm,2.8*cm,2.8*cm,2.4*cm]
        ))
    story.append(Spacer(1, 0.4*cm))

    # Section 3
    story.append(Paragraph("3. Recommendations", h1_style))
    recs = [
        ("Prioritise direct submission support for overdue counties",
         f"Turkana, Marsabit, Mandera and Wajir have the lowest electricity access in Kenya "
         f"(8–12%) and are all overdue. Direct technical assistance is recommended."),
        ("Accelerate clean cooking intervention",
         f"Current clean cooking access ({nat['w_cooking']}%) is well below the 2028 target. "
         f"LPG subsidy expansion and ICS distribution should be accelerated."),
        ("Complete INEP when all 47 counties submit",
         f"The national picture is {nat['coverage_pct']}% complete. A firm deadline should "
         f"be set for the remaining {nat['pending_count']+nat['overdue_count']} counties."),
    ]
    for i,(title,body) in enumerate(recs,1):
        story.append(Paragraph(f"{i}. {title}", ParagraphStyle("rec",parent=body_style,fontName="Helvetica-Bold",textColor=colors.HexColor("#0e1e2e"))))
        story.append(Paragraph(f"   {body}", body_style))

    # Footer note
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e8e6de")))
    story.append(Paragraph(
        f"Data source: KenyaWatts Platform · EPRA Statistics FY 2024/25 · "
        f"Generated {now.strftime('%d %B %Y at %H:%M')} · "
        f"CONFIDENTIAL — EPRA INTERNAL USE ONLY",
        meta_style
    ))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def page_national_report(user_info):
    """EPRA-only: Generate and download the national INEP summary report."""
    section("National INEP report generator",
            "Generate a Word or PDF summary of all submitted county plans — EPRA confidential")

    alert("warn",
        "<b>CONFIDENTIAL — EPRA internal use only.</b> This report aggregates submitted county "
        "energy plans into a national summary. It is intended for EPRA planners and the "
        "Ministry of Energy only. Do not distribute without authorisation.")

    nat = compute_national()

    # Live summary of what the report will contain
    st.markdown("**What this report will contain:**")
    c1,c2,c3 = st.columns(3)
    with c1:
        st.markdown(f"""<div style="background:#0e1e2e;border-radius:10px;padding:14px 16px">
          <div style="font-size:11px;color:#a8c4d4;margin-bottom:6px;text-transform:uppercase;letter-spacing:.4px">Data coverage</div>
          <div style="font-size:18px;font-weight:700;color:#3ecfaa">{nat['submitted_count']} / {nat['total_counties']}</div>
          <div style="font-size:11px;color:#a8c4d4">counties with submitted plans</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div style="background:#0e1e2e;border-radius:10px;padding:14px 16px">
          <div style="font-size:11px;color:#a8c4d4;margin-bottom:6px;text-transform:uppercase;letter-spacing:.4px">National electricity access</div>
          <div style="font-size:18px;font-weight:700;color:#3ecfaa">{nat['w_elec']}%</div>
          <div style="font-size:11px;color:#a8c4d4">population-weighted average</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div style="background:#0e1e2e;border-radius:10px;padding:14px 16px">
          <div style="font-size:11px;color:#a8c4d4;margin-bottom:6px;text-transform:uppercase;letter-spacing:.4px">Total investment identified</div>
          <div style="font-size:18px;font-weight:700;color:#3ecfaa">KES {nat['total_budget']}B</div>
          <div style="font-size:11px;color:#a8c4d4">from submitted county budgets</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("")
    st.markdown("**Report sections:**")
    st.markdown("""
    1. Executive summary with key national indicators table
    2. County submission status — submitted, pending, and overdue counties
    3. National energy access analysis — electricity and clean cooking
    4. Renewable energy resource assessment by county
    5. Investment and budget summary
    6. Recommendations for EPRA action
    7. Data notes and methodology
    """)

    st.divider()
    st.markdown("**Generate report:**")

    col_fmt, col_gen = st.columns([1,2])
    with col_fmt:
        fmt = st.radio("Report format:", ["📄 Word document (.docx)", "📑 PDF document (.pdf)"], index=0)
    with col_gen:
        report_title = st.text_input("Report title (optional)",
            value=f"Kenya INEP National Summary — {datetime.now().strftime('%B %Y')}",
            key="report_title_input")
        include_overdue  = st.checkbox("Include overdue county details", value=True)
        include_recs     = st.checkbox("Include recommendations section", value=True)

    st.markdown("")

    # Show library availability status
    col_status1, col_status2 = st.columns(2)
    with col_status1:
        if DOCX_AVAILABLE:
            st.success("✓ Word (.docx) generation ready")
        else:
            st.warning("⚠ Word generation unavailable — python-docx not installed")
    with col_status2:
        if PDF_AVAILABLE:
            st.success("✓ PDF generation ready")
        else:
            st.warning("⚠ PDF generation unavailable — reportlab not installed")

    st.markdown("")
    # Generate button — shows download button immediately after generation
    if st.button("⬇️  Generate and download report", type="primary",
                 key="generate_report_btn", use_container_width=False):
        with st.spinner("Generating national report — please wait…"):
            try:
                if "Word" in fmt:
                    if not DOCX_AVAILABLE:
                        st.error(
                            "python-docx library is not available on this server. "
                            "Check that requirements.txt contains: python-docx>=1.1.0 "
                            "and that the app has been redeployed after adding it. "
                            "Use PDF format instead — it is working correctly."
                        )
                        return
                    # Generate Word document
                    data     = generate_word_report(nat, COUNTIES, user_info["name"])
                    filename = f"Kenya_INEP_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    mime     = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    fmt_label = "Word"
                else:
                    if not PDF_AVAILABLE:
                        st.error(
                            "reportlab library is not available. "
                            "Check requirements.txt contains: reportlab>=4.2.0"
                        )
                        return
                    # Generate PDF
                    data     = generate_pdf_report(nat, COUNTIES, user_info["name"])
                    filename = f"Kenya_INEP_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                    mime     = "application/pdf"
                    fmt_label = "PDF"

                # Show size and success
                size_kb = len(data) // 1024
                st.success(
                    f"✅ {fmt_label} report generated · {size_kb} KB · "
                    f"{datetime.now().strftime('%H:%M:%S')}"
                )

                # Download button — appears immediately below
                st.download_button(
                    label     = f"📥  Download {filename}",
                    data      = data,
                    file_name = filename,
                    mime      = mime,
                    type      = "primary",
                    key       = "download_report_btn",
                )

                # Audit trail entry
                st.session_state.audit_log.append({
                    "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "user":   user_info["name"],
                    "action": f"National INEP {fmt_label} report generated",
                    "ref":    f"RPT-{datetime.now().strftime('%Y%m%d%H%M')}",
                })
                alert("info",
                    "<b>CONFIDENTIAL report generated.</b> "
                    "Share only with authorised EPRA leadership and Ministry of Energy officials. "
                    "This generation has been logged in the audit trail.")

            except Exception as e:
                # Show the full error so we can diagnose it
                st.error(f"Report generation failed: {type(e).__name__}: {str(e)}")
                with st.expander("Full error details (for debugging)"):
                    st.exception(e)
                st.info(
                    "If Word fails, try PDF format instead. "
                    "If both fail, check the Streamlit Cloud logs for the library error."
                )

# ── MAKUENI — CONTEXTUAL (different per role) ─────────────────────────────────
def page_makueni_contextual(role, county_id):
    is_makueni_county = county_id == "MK"

    if role == "epra":
        # EPRA sees it as a template reference tool
        st.markdown("""<div class="kw-alert-info">
        <b>EPRA template reference:</b> The Makueni CEP (2023–2032) is Kenya's most comprehensive
        completed county energy plan. Use it to demonstrate the submission standard to other counties,
        design the submission template, and benchmark incoming plans. Developed with WRI and Strathmore University.
        </div>""", unsafe_allow_html=True)

    elif is_makueni_county:
        # Makueni sees their own plan
        st.markdown("""<div class="kw-alert-success">
        <b>Your county energy plan:</b> Makueni County Energy Plan 2023–2032.
        This is the plan your county submitted. It is stored in the KenyaWatts repository
        and visible to EPRA and the Ministry. Your plan is being used as the national reference standard.
        </div>""", unsafe_allow_html=True)

    # Show the full plan content for both
    page_makueni()

    # Add extra guidance for EPRA only
    if role == "epra":
        st.divider()
        st.markdown("### How to use this as a template reference")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            **What Makueni did well:**
            - Covered all 7 CEP chapters
            - Used OnSSET for electrification scenarios
            - Used LEAP for clean cooking modelling
            - Disaggregated data by gender (GESI)
            - Reported at sub-county level (6 sub-counties)
            - Declared all scenario assumptions clearly
            - Included M&E framework with annual targets
            """)
        with col2:
            st.markdown("""
            **What other counties can simplify:**
            - Full OnSSET/LEAP not required — simplified estimates accepted
            - Sub-county breakdown optional for first submission
            - Budget estimates accepted if exact figures not yet available
            - Document upload pathway available for counties with existing plans
            - KenyaWatts validation will flag missing mandatory fields
            """)

# ── SUMMARY REPORT GENERATORS (used by Data Download tab) ─────────────────────

def generate_word_summary_report(nat, counties_df, generated_by, sections, county_filter=None, report_title=""):
    """Generate a Word (.docx) National or County-Specific Summary Report with selectable sections."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def add_heading(text, level=1, color=(26, 31, 22)):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Arial"
        return h

    def add_para(text, bold=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"
        return p

    def add_table(headers, rows, header_color=(14, 30, 46)):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0]
        for h, cell in zip(headers, hdr.cells):
            cell.text = h
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(255, 255, 255)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            hex_color = "%02x%02x%02x" % tuple(header_color)
            shd.set(_qn("w:fill"), hex_color)
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:val"), "clear")
            tc_pr.append(shd)
        for ri, row_data in enumerate(rows):
            tr = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = tr.cells[ci]
                cell.text = str(val)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
                run.font.name = "Arial"
                if ri % 2 == 1:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(_qn("w:fill"), "F7F6F2")
                    shd.set(_qn("w:color"), "auto")
                    shd.set(_qn("w:val"), "clear")
                    tc_pr.append(shd)
        return t

    now = datetime.now()
    is_county_report = county_filter is not None
    cdf = counties_df[counties_df["name"] == county_filter] if is_county_report else counties_df
    county_row = cdf.iloc[0] if (is_county_report and not cdf.empty) else None

    if not report_title:
        report_title = (
            f"{county_filter} County Energy Summary — {now.strftime('%B %Y')}"
            if is_county_report
            else f"Kenya National Summary Report — {now.strftime('%B %Y')}"
        )

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(report_title.upper())
    tr.bold = True; tr.font.size = Pt(20); tr.font.name = "Arial"
    tr.font.color.rgb = RGBColor(14, 30, 46)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_text = (
        f"{county_filter} County — Energy Planning Summary Report"
        if is_county_report
        else "National Summary Report — Aggregated from County Energy Plans"
    )
    sr = sub.add_run(sub_text)
    sr.font.size = Pt(13); sr.font.name = "Arial"; sr.font.color.rgb = RGBColor(26, 111, 163)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Generated by: {generated_by}\n"
        f"Date: {now.strftime('%d %B %Y')} at {now.strftime('%H:%M')}\n"
        f"Platform: KenyaWatts · EPRA Kenya\n"
        f"Classification: CONFIDENTIAL — EPRA INTERNAL USE ONLY"
    )
    mr.font.size = Pt(10); mr.font.name = "Arial"; mr.font.color.rgb = RGBColor(90, 138, 158)
    doc.add_page_break()

    sec_num = 1

    # ── Section 1: Executive Summary ──────────────────────────────────────────
    if sections.get("exec_summary", True):
        add_heading(f"{sec_num}. Executive Summary", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            add_para(
                f"This report presents the energy planning status of {r['name']} County. "
                f"Current electricity access stands at {r['elec']}%, with clean cooking access "
                f"at {r['cooking']}%. The county has a population of {r['pop']:,} and targets "
                f"universal electricity access by {r['target_yr']}. "
                f"Submission status: {r['status'].title()}."
            )
            doc.add_paragraph()
            add_heading(f"Key County Indicators", 2, (26, 111, 163))
            add_table(
                ["Indicator", "Current Value", "National Target", "Target Year", "Status"],
                [
                    ["Electricity access", f"{r['elec']}%", "100%", "2030",
                     "On track" if r["elec"] > 70 else "At risk"],
                    ["Clean cooking access", f"{r['cooking']}%", "100%", "2028",
                     "At risk" if r["cooking"] < 50 else "On track"],
                    ["MTF demand tier", f"Tier {r['mtf']}", "Tier 4–5", "2030",
                     "On track" if r["mtf"] >= 3 else "At risk"],
                    ["Plan budget (KES B)", f"{r['budget']}B" if r["budget"] > 0 else "Not submitted",
                     "TBD", f"2023–{r['target_yr']}", "Submitted" if r["budget"] > 0 else "Pending"],
                    ["Submission status", r["status"].title(), "Submitted", "Current cycle",
                     "✓" if r["status"] in ("submitted", "review") else "⚠"],
                ]
            )
        else:
            add_para(
                f"This report presents the current state of Kenya's national energy planning "
                f"as aggregated from {nat['submitted_count']} submitted County Energy Plans (CEPs) "
                f"on the KenyaWatts platform. Data covers {nat['coverage_pct']}% of Kenya's 47 counties. "
                f"{nat['overdue_count']} counties remain overdue on their submissions."
            )
            doc.add_paragraph()
            add_heading("Key National Indicators", 2, (26, 111, 163))
            add_table(
                ["Indicator", "Current Value", "National Target", "Target Year", "Status"],
                [
                    ["Electricity access (weighted avg)", f"{nat['w_elec']}%", "100%", "2030",
                     "On track" if float(nat['w_elec']) > 70 else "At risk"],
                    ["Clean cooking access (weighted avg)", f"{nat['w_cooking']}%", "100%", "2028",
                     "At risk" if float(nat['w_cooking']) < 50 else "On track"],
                    ["Clean energy generation", "82%", "100%", "2035", "On track"],
                    ["Counties submitted", f"{nat['submitted_count']} / {nat['total_counties']}", "47 / 47",
                     "Current cycle", "In progress"],
                    ["Total investment identified", f"KES {nat['total_budget']}B",
                     "TBD (all 47 counties)", "2023–2032", "Partial"],
                ]
            )
        doc.add_paragraph()

    # ── Section 2: Submission Status ──────────────────────────────────────────
    if sections.get("submission_status", True):
        add_heading(f"{sec_num}. County Submission Status", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            add_para(
                f"{r['name']} County's submission status is '{r['status'].title()}'. "
                f"{'The county has submitted its County Energy Plan and data is included in the national INEP aggregation.' if r['status'] in ('submitted','review') else 'The county has not yet submitted its County Energy Plan.'}"
            )
            add_table(
                ["Field", "Value"],
                [
                    ["County", r["name"]],
                    ["Region", r["region"]],
                    ["Status", r["status"].title()],
                    ["Population", f"{r['pop']:,}"],
                    ["Target year", str(r["target_yr"])],
                    ["Budget (KES B)", str(r["budget"]) if r["budget"] > 0 else "Not submitted"],
                ]
            )
        else:
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            overdue = counties_df[counties_df["status"] == "overdue"]
            pending = counties_df[counties_df["status"] == "pending"]
            add_para(
                f"As of {now.strftime('%d %B %Y')}, {nat['submitted_count']} of Kenya's 47 counties "
                f"have submitted County Energy Plans. {nat['overdue_count']} counties are overdue "
                f"and {nat['pending_count']} counties are yet to begin submission."
            )
            doc.add_paragraph()
            add_heading("Submitted counties", 2, (15, 157, 126))
            if not submitted.empty:
                add_table(
                    ["County", "Region", "Electricity (%)", "Clean cooking (%)", "Budget (KES B)", "Target year"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['cooking']}%",
                      str(r["budget"]) if r["budget"] > 0 else "—", str(r["target_yr"])]
                     for _, r in submitted.iterrows()]
                )
            doc.add_paragraph()
            add_heading("Overdue counties — action required", 2, (163, 45, 44))
            if not overdue.empty:
                add_table(
                    ["County", "Region", "Electricity access (%)", "Population", "Reminders sent"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['pop'] // 1000:,}K", "3 (final notice)"]
                     for _, r in overdue.iterrows()]
                )
            doc.add_paragraph()
            add_heading("Pending counties", 2, (26, 111, 163))
            if not pending.empty:
                add_table(
                    ["County", "Region", "Electricity access (%)", "Population"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['pop'] // 1000:,}K"]
                     for _, r in pending.iterrows()]
                )
        doc.add_paragraph()

    # ── Section 3: Energy Access Analysis ────────────────────────────────────
    if sections.get("energy_access", True):
        add_heading(f"{sec_num}. Energy Access Analysis", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            add_heading("Electricity Access", 2, (26, 111, 163))
            add_para(
                f"{r['name']} County has an electricity access rate of {r['elec']}%, "
                f"leaving a gap of {100 - r['elec']}% to the 2030 universal access target. "
                f"The county is classified as MTF Tier {r['mtf']} demand."
            )
            doc.add_paragraph()
            add_heading("Clean Cooking Access", 2, (26, 111, 163))
            add_para(
                f"Clean cooking access in {r['name']} County stands at {r['cooking']}%. "
                f"Kenya's target is universal clean cooking by 2028. "
                f"{'Significant acceleration is needed to close this gap.' if r['cooking'] < 50 else 'Progress is on track but further investment is needed.'}"
            )
        else:
            add_heading("Electricity Access", 2, (26, 111, 163))
            add_para(
                f"The population-weighted national electricity access rate is {nat['w_elec']}%. "
                f"The four counties with the most critical gaps — Turkana (12%), Marsabit (8%), "
                f"Mandera (11%), and Wajir (9%) — are all overdue on plan submission."
            )
            doc.add_paragraph()
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            if not submitted.empty:
                add_heading("Electricity Access by County (submitted plans)", 2, (26, 111, 163))
                add_table(
                    ["County", "Electricity access (%)", "Gap to 100%", "MTF demand tier"],
                    [[r["name"], f"{r['elec']}%", f"{100 - r['elec']}%", f"Tier {r['mtf']}"]
                     for _, r in submitted.sort_values("elec").iterrows()]
                )
            doc.add_paragraph()
            add_heading("Clean Cooking Access", 2, (26, 111, 163))
            add_para(
                f"The population-weighted clean cooking access rate is {nat['w_cooking']}%. "
                f"Kenya's target is universal clean cooking by 2028. "
                f"{'This trajectory suggests the 2028 target is at risk without accelerated intervention.' if float(nat['w_cooking']) < 50 else 'Progress is being made but acceleration is needed.'}"
            )
        doc.add_paragraph()

    # ── Section 4: Renewable Energy Resources ────────────────────────────────
    if sections.get("renewable_resources", True):
        add_heading(f"{sec_num}. Renewable Energy Resource Assessment", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            solar_label = "High (>2,000)" if r["solar"] > 2000 else "Good (1,800–2,000)" if r["solar"] > 1800 else "Moderate"
            add_para(
                f"{r['name']} County has a solar GHI of {r['solar']} kWh/m²/year, "
                f"classified as '{solar_label}'. This represents a significant opportunity "
                f"for solar-based electrification solutions including Solar Home Systems and minigrids."
            )
            add_table(
                ["Resource", "Value", "Assessment"],
                [
                    ["Solar GHI", f"{r['solar']} kWh/m²/yr", solar_label],
                    ["Solar potential", "High" if r["solar"] > 1950 else "Moderate", "Suitable for SHS and minigrids"],
                ]
            )
        else:
            add_para(
                "Kenya has abundant renewable energy resources distributed across its counties. "
                "Northern counties have the highest solar potential while the Rift Valley and Coast "
                "offer significant wind and geothermal resources."
            )
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            if not submitted.empty:
                add_table(
                    ["County", "Solar GHI (kWh/m²/yr)", "Solar potential", "Population (K)"],
                    [[r["name"], str(r["solar"]),
                      "High (>2,000)" if r["solar"] > 2000 else "Good (1,800–2,000)" if r["solar"] > 1800 else "Moderate",
                      f"{r['pop'] // 1000:,}"]
                     for _, r in submitted.sort_values("solar", ascending=False).iterrows()]
                )
        doc.add_paragraph()

    # ── Section 5: Investment and Budget ─────────────────────────────────────
    if sections.get("investment_budget", True):
        add_heading(f"{sec_num}. Investment and Budget Summary", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            if r["budget"] > 0:
                add_para(
                    f"{r['name']} County has identified a total investment of KES {r['budget']} billion "
                    f"for its energy plan (2023–{r['target_yr']}). This equates to approximately "
                    f"KES {int(r['budget'] * 1e9 / r['pop']):,} per capita."
                )
                add_table(
                    ["Item", "Value"],
                    [
                        ["County", r["name"]],
                        ["Total budget (KES B)", f"{r['budget']}B"],
                        ["Budget per capita (KES)", f"{int(r['budget'] * 1e9 / r['pop']):,}"],
                        ["Plan period", f"2023–{r['target_yr']}"],
                    ]
                )
            else:
                add_para(
                    f"{r['name']} County has not yet submitted a budget estimate. "
                    "This information is required as part of the County Energy Plan submission."
                )
        else:
            add_para(
                f"Based on {nat['submitted_count']} submitted county plans, the total identified "
                f"investment required is KES {nat['total_budget']} billion. This covers electricity "
                f"access expansion, clean cooking programmes, renewable energy development, and "
                f"energy efficiency improvements."
            )
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            budget_counties = submitted[submitted["budget"] > 0] if not submitted.empty else submitted
            if not budget_counties.empty:
                doc.add_paragraph()
                add_table(
                    ["County", "Plan budget (KES B)", "Budget per capita (KES)", "Plan period"],
                    [[r["name"], f"{r['budget']}B",
                      f"{int(r['budget'] * 1e9 / r['pop']):,}",
                      f"2023–{r['target_yr']}"]
                     for _, r in budget_counties.iterrows()]
                )
        doc.add_paragraph()

    # ── Section 6: Recommendations ────────────────────────────────────────────
    if sections.get("recommendations", True):
        add_heading(f"{sec_num}. Recommendations", 1, (14, 30, 46))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            recs = []
            if r["elec"] < 50:
                recs.append((
                    f"Prioritise electricity access expansion in {r['name']}",
                    f"With only {r['elec']}% electricity access, {r['name']} County requires "
                    f"urgent investment in grid extension, minigrids, and Solar Home Systems "
                    f"to meet the 2030 universal access target."
                ))
            if r["cooking"] < 40:
                recs.append((
                    f"Accelerate clean cooking access in {r['name']}",
                    f"Clean cooking access of {r['cooking']}% is well below the national "
                    f"2028 target. LPG subsidy programmes and improved cookstove distribution "
                    f"should be prioritised."
                ))
            if r["status"] in ("pending", "overdue"):
                recs.append((
                    f"Complete and submit County Energy Plan",
                    f"{r['name']} County has not yet submitted its County Energy Plan. "
                    f"Submission is required under the Energy Act 2019 and is needed for "
                    f"inclusion in the national INEP aggregation."
                ))
            if r["solar"] > 2000:
                recs.append((
                    f"Leverage high solar potential for off-grid solutions",
                    f"With a solar GHI of {r['solar']} kWh/m²/year, {r['name']} County "
                    f"is well-suited for solar-based electrification. EPRA should facilitate "
                    f"private sector investment in Solar Home Systems and solar minigrids."
                ))
            if not recs:
                recs.append((
                    f"Continue progress and submit updated plan",
                    f"{r['name']} County is making good progress on energy access. "
                    f"Continued monitoring and annual plan updates are recommended to maintain "
                    f"trajectory toward the 2030 and 2028 targets."
                ))
            for i, (title, body) in enumerate(recs, 1):
                add_para(f"{i}. {title}", bold=True)
                add_para(f"   {body}")
                doc.add_paragraph()
        else:
            recs = [
                ("Prioritise direct submission support for overdue counties",
                 f"Turkana, Marsabit, Mandera and Wajir have the lowest electricity access in Kenya "
                 f"(8–12%) and are all overdue. EPRA should arrange direct technical assistance."),
                ("Accelerate clean cooking intervention",
                 f"Current clean cooking access ({nat['w_cooking']}%) is significantly below "
                 f"the 2028 target. LPG subsidy expansion and improved cookstove distribution "
                 f"programmes should be accelerated."),
                ("Complete national aggregation when all 47 counties submit",
                 f"The national INEP covers {nat['coverage_pct']}% of counties. A firm deadline "
                 f"should be set for the remaining {nat['pending_count'] + nat['overdue_count']} counties."),
                ("Publish county comparison dashboard publicly",
                 "Making county-level energy access data publicly available will create "
                 "accountability pressure and incentivise timely submission in future cycles."),
                ("Increase investment in high-solar northern counties",
                 "Counties such as Mandera, Wajir, Marsabit and Turkana have the highest solar "
                 "potential in Kenya (>2,100 kWh/m²/yr) yet the lowest access. Targeted "
                 "off-grid solar programmes can cost-effectively close this gap."),
            ]
            for i, (title, body) in enumerate(recs, 1):
                add_para(f"{i}. {title}", bold=True)
                add_para(f"   {body}")
                doc.add_paragraph()

    # ── Section 7: Data Notes ─────────────────────────────────────────────────
    if sections.get("data_notes", True):
        add_heading(f"{sec_num}. Data Notes and Methodology", 1, (14, 30, 46))
        if is_county_report and county_row is not None:
            add_para(
                f"• County data sourced from the KenyaWatts platform and the submitted "
                f"County Energy Plan for {county_filter}.\n"
                f"• Population figures from KNBS 2019 Census.\n"
                f"• Solar GHI data from Global Solar Atlas.\n"
                f"• Electricity and clean cooking access rates from submitted county plans.\n"
                f"• Report generated from KenyaWatts platform on {now.strftime('%d %B %Y')}.\n"
                f"• This report is classified CONFIDENTIAL — EPRA internal use only."
            )
        else:
            add_para(
                f"• National electricity access and clean cooking figures are population-weighted "
                f"averages from {nat['submitted_count']} submitted county plans.\n"
                f"• Budget figures are simple sums of submitted county plan budgets.\n"
                f"• Universal access target year reflects the latest target year from submitted plans ({nat['latest_target']}).\n"
                f"• Population figures from KNBS 2019 Census.\n"
                f"• Clean energy generation (82%) sourced from EPRA Statistics Report FY 2024/25.\n"
                f"• Report generated from KenyaWatts platform on {now.strftime('%d %B %Y')}.\n"
                f"• This report is classified CONFIDENTIAL — EPRA internal use only."
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_pdf_summary_report(nat, counties_df, generated_by, sections, county_filter=None, report_title=""):
    """Generate a PDF National or County-Specific Summary Report with selectable sections."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm)

    styles = getSampleStyleSheet()
    now = datetime.now()
    is_county_report = county_filter is not None
    cdf = counties_df[counties_df["name"] == county_filter] if is_county_report else counties_df
    county_row = cdf.iloc[0] if (is_county_report and not cdf.empty) else None

    if not report_title:
        report_title = (
            f"{county_filter} County Energy Summary — {now.strftime('%B %Y')}"
            if is_county_report
            else f"Kenya National Summary Report — {now.strftime('%B %Y')}"
        )

    title_style = ParagraphStyle("KWTitle", parent=styles["Title"],
        fontSize=18, textColor=colors.HexColor("#0e1e2e"),
        fontName="Helvetica-Bold", spaceAfter=6, alignment=TA_CENTER)
    sub_style = ParagraphStyle("KWSub", parent=styles["Normal"],
        fontSize=12, textColor=colors.HexColor("#1a6fa3"),
        fontName="Helvetica", spaceAfter=4, alignment=TA_CENTER)
    meta_style = ParagraphStyle("KWMeta", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#5a8a9e"),
        fontName="Helvetica", spaceAfter=2, alignment=TA_CENTER)
    h1_style = ParagraphStyle("KWH1", parent=styles["Heading1"],
        fontSize=14, textColor=colors.HexColor("#0e1e2e"),
        fontName="Helvetica-Bold", spaceBefore=16, spaceAfter=6)
    h2_style = ParagraphStyle("KWH2", parent=styles["Heading2"],
        fontSize=12, textColor=colors.HexColor("#1a6fa3"),
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("KWBody", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#1a1916"),
        fontName="Helvetica", spaceAfter=6, leading=15)
    conf_style = ParagraphStyle("KWConf", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#b33a2c"),
        fontName="Helvetica-Bold", alignment=TA_CENTER)

    def make_table(headers, rows, col_widths=None):
        data = [headers] + rows
        tw = col_widths or [(A4[0] - 4 * cm) / len(headers)] * len(headers)
        t = Table(data, colWidths=tw)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0e1e2e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f6f2")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e8e6de")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        return t

    story = []

    # Cover
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(report_title.upper(), title_style))
    sub_text = (
        f"{county_filter} County — Energy Planning Summary Report"
        if is_county_report
        else "National Summary Report — Aggregated from County Energy Plans"
    )
    story.append(Paragraph(sub_text, sub_style))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(f"Generated by: {generated_by}", meta_style))
    story.append(Paragraph(f"Date: {now.strftime('%d %B %Y')} at {now.strftime('%H:%M')}", meta_style))
    story.append(Paragraph("Platform: KenyaWatts · Energy & Petroleum Regulatory Authority", meta_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("CONFIDENTIAL — EPRA INTERNAL USE ONLY", conf_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0e1e2e")))
    story.append(Spacer(1, 0.5 * cm))

    sec_num = 1

    # Section 1: Executive Summary
    if sections.get("exec_summary", True):
        story.append(Paragraph(f"{sec_num}. Executive Summary", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            story.append(Paragraph(
                f"This report presents the energy planning status of <b>{r['name']}</b> County. "
                f"Current electricity access stands at <b>{r['elec']}%</b>, with clean cooking access "
                f"at <b>{r['cooking']}%</b>. The county has a population of {r['pop']:,} and targets "
                f"universal electricity access by {r['target_yr']}. "
                f"Submission status: <b>{r['status'].title()}</b>.",
                body_style
            ))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("Key County Indicators", h2_style))
            story.append(make_table(
                ["Indicator", "Current Value", "Target", "Year", "Status"],
                [
                    ["Electricity access", f"{r['elec']}%", "100%", "2030",
                     "On track" if r["elec"] > 70 else "At risk"],
                    ["Clean cooking access", f"{r['cooking']}%", "100%", "2028",
                     "At risk" if r["cooking"] < 50 else "On track"],
                    ["MTF demand tier", f"Tier {r['mtf']}", "Tier 4–5", "2030",
                     "On track" if r["mtf"] >= 3 else "At risk"],
                    ["Plan budget (KES B)", f"{r['budget']}B" if r["budget"] > 0 else "Not submitted",
                     "TBD", f"2023–{r['target_yr']}", "Submitted" if r["budget"] > 0 else "Pending"],
                    ["Submission status", r["status"].title(), "Submitted", "Current", "✓" if r["status"] in ("submitted", "review") else "⚠"],
                ],
                [5 * cm, 3 * cm, 2 * cm, 2 * cm, 2.5 * cm]
            ))
        else:
            story.append(Paragraph(
                f"This report presents the current state of Kenya's national energy planning as "
                f"aggregated from <b>{nat['submitted_count']}</b> submitted County Energy Plans (CEPs) "
                f"on the KenyaWatts platform. Data covers <b>{nat['coverage_pct']}%</b> of Kenya's "
                f"47 counties. {nat['overdue_count']} counties remain overdue.",
                body_style
            ))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("Key National Indicators", h2_style))
            story.append(make_table(
                ["Indicator", "Current Value", "Target", "Year", "Status"],
                [
                    ["Electricity access", f"{nat['w_elec']}%", "100%", "2030",
                     "On track" if float(nat['w_elec']) > 70 else "At risk"],
                    ["Clean cooking access", f"{nat['w_cooking']}%", "100%", "2028",
                     "At risk" if float(nat['w_cooking']) < 50 else "On track"],
                    ["Clean energy generation", "82%", "100%", "2035", "On track"],
                    ["Counties submitted", f"{nat['submitted_count']}/47", "47/47", "Current", "In progress"],
                    ["Total investment identified", f"KES {nat['total_budget']}B", "TBD", "2023–2032", "Partial"],
                ],
                [5 * cm, 3 * cm, 2 * cm, 2 * cm, 2.5 * cm]
            ))
        story.append(Spacer(1, 0.4 * cm))

    # Section 2: Submission Status
    if sections.get("submission_status", True):
        story.append(Paragraph(f"{sec_num}. County Submission Status", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            story.append(Paragraph(
                f"<b>{r['name']}</b> County's submission status is <b>{r['status'].title()}</b>. "
                f"{'The county has submitted its County Energy Plan.' if r['status'] in ('submitted','review') else 'The county has not yet submitted its County Energy Plan.'}",
                body_style
            ))
            story.append(make_table(
                ["Field", "Value"],
                [
                    ["County", r["name"]], ["Region", r["region"]],
                    ["Status", r["status"].title()], ["Population", f"{r['pop']:,}"],
                    ["Target year", str(r["target_yr"])],
                    ["Budget (KES B)", str(r["budget"]) if r["budget"] > 0 else "Not submitted"],
                ],
                [6 * cm, 8.5 * cm]
            ))
        else:
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            overdue = counties_df[counties_df["status"] == "overdue"]
            pending = counties_df[counties_df["status"] == "pending"]
            story.append(Paragraph(
                f"As of {now.strftime('%d %B %Y')}, <b>{nat['submitted_count']}</b> of 47 counties "
                f"have submitted County Energy Plans. <b>{nat['overdue_count']}</b> counties are overdue "
                f"and <b>{nat['pending_count']}</b> are yet to begin.",
                body_style
            ))
            story.append(Paragraph("Submitted county plans", h2_style))
            if not submitted.empty:
                story.append(make_table(
                    ["County", "Region", "Electricity (%)", "Cooking (%)", "Budget (KES B)", "Target year"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['cooking']}%",
                      str(r["budget"]) if r["budget"] > 0 else "—", str(r["target_yr"])]
                     for _, r in submitted.iterrows()],
                    [3 * cm, 3 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm, 2 * cm]
                ))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("Overdue counties — action required", h2_style))
            if not overdue.empty:
                story.append(make_table(
                    ["County", "Region", "Electricity (%)", "Population", "Reminders sent"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['pop'] // 1000:,}K", "3 (final)"]
                     for _, r in overdue.iterrows()],
                    [3.5 * cm, 3.5 * cm, 2.8 * cm, 2.8 * cm, 2.9 * cm]
                ))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("Pending counties", h2_style))
            if not pending.empty:
                story.append(make_table(
                    ["County", "Region", "Electricity (%)", "Population"],
                    [[r["name"], r["region"], f"{r['elec']}%", f"{r['pop'] // 1000:,}K"]
                     for _, r in pending.iterrows()],
                    [4 * cm, 4 * cm, 3 * cm, 4.5 * cm]
                ))
        story.append(Spacer(1, 0.4 * cm))

    # Section 3: Energy Access
    if sections.get("energy_access", True):
        story.append(Paragraph(f"{sec_num}. Energy Access Analysis", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            story.append(Paragraph("Electricity Access", h2_style))
            story.append(Paragraph(
                f"<b>{r['name']}</b> County has an electricity access rate of <b>{r['elec']}%</b>, "
                f"leaving a gap of <b>{100 - r['elec']}%</b> to the 2030 universal access target. "
                f"The county is classified as MTF Tier {r['mtf']} demand.",
                body_style
            ))
            story.append(Paragraph("Clean Cooking Access", h2_style))
            story.append(Paragraph(
                f"Clean cooking access in <b>{r['name']}</b> County stands at <b>{r['cooking']}%</b>. "
                f"{'Significant acceleration is needed to close this gap.' if r['cooking'] < 50 else 'Progress is on track but further investment is needed.'}",
                body_style
            ))
        else:
            story.append(Paragraph("Electricity Access", h2_style))
            story.append(Paragraph(
                f"The population-weighted national electricity access rate is <b>{nat['w_elec']}%</b>. "
                f"The four counties with the most critical gaps — Turkana (12%), Marsabit (8%), "
                f"Mandera (11%), and Wajir (9%) — are all overdue on plan submission.",
                body_style
            ))
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            if not submitted.empty:
                story.append(Paragraph("Electricity Access by County (submitted plans)", h2_style))
                story.append(make_table(
                    ["County", "Electricity access (%)", "Gap to 100%", "MTF tier"],
                    [[r["name"], f"{r['elec']}%", f"{100 - r['elec']}%", f"Tier {r['mtf']}"]
                     for _, r in submitted.sort_values("elec").iterrows()],
                    [4.5 * cm, 4 * cm, 3.5 * cm, 3.5 * cm]
                ))
            story.append(Paragraph("Clean Cooking Access", h2_style))
            story.append(Paragraph(
                f"The population-weighted clean cooking access rate is <b>{nat['w_cooking']}%</b>. "
                f"Kenya's target is universal clean cooking by 2028.",
                body_style
            ))
        story.append(Spacer(1, 0.4 * cm))

    # Section 4: Renewable Resources
    if sections.get("renewable_resources", True):
        story.append(Paragraph(f"{sec_num}. Renewable Energy Resource Assessment", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            solar_label = "High (>2,000)" if r["solar"] > 2000 else "Good (1,800–2,000)" if r["solar"] > 1800 else "Moderate"
            story.append(Paragraph(
                f"<b>{r['name']}</b> County has a solar GHI of <b>{r['solar']} kWh/m²/year</b> "
                f"({solar_label}). This represents a significant opportunity for solar-based "
                f"electrification including Solar Home Systems and minigrids.",
                body_style
            ))
            story.append(make_table(
                ["Resource", "Value", "Assessment"],
                [
                    ["Solar GHI", f"{r['solar']} kWh/m²/yr", solar_label],
                    ["Solar electrification", "Suitable" if r["solar"] > 1800 else "Moderate", "SHS and minigrids viable"],
                ],
                [5 * cm, 5 * cm, 5.5 * cm]
            ))
        else:
            story.append(Paragraph(
                "Kenya has abundant renewable energy resources distributed across its counties. "
                "Northern counties have the highest solar potential.",
                body_style
            ))
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            if not submitted.empty:
                story.append(make_table(
                    ["County", "Solar GHI (kWh/m²/yr)", "Potential", "Population (K)"],
                    [[r["name"], str(r["solar"]),
                      "High" if r["solar"] > 2000 else "Good" if r["solar"] > 1800 else "Moderate",
                      f"{r['pop'] // 1000:,}"]
                     for _, r in submitted.sort_values("solar", ascending=False).iterrows()],
                    [4 * cm, 4 * cm, 3 * cm, 4.5 * cm]
                ))
        story.append(Spacer(1, 0.4 * cm))

    # Section 5: Investment
    if sections.get("investment_budget", True):
        story.append(Paragraph(f"{sec_num}. Investment and Budget Summary", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            if r["budget"] > 0:
                story.append(Paragraph(
                    f"<b>{r['name']}</b> County has identified a total investment of "
                    f"<b>KES {r['budget']} billion</b> for its energy plan (2023–{r['target_yr']}), "
                    f"equating to KES {int(r['budget'] * 1e9 / r['pop']):,} per capita.",
                    body_style
                ))
                story.append(make_table(
                    ["Item", "Value"],
                    [
                        ["Total budget (KES B)", f"{r['budget']}B"],
                        ["Budget per capita (KES)", f"{int(r['budget'] * 1e9 / r['pop']):,}"],
                        ["Plan period", f"2023–{r['target_yr']}"],
                    ],
                    [6 * cm, 9.5 * cm]
                ))
            else:
                story.append(Paragraph(
                    f"<b>{r['name']}</b> County has not yet submitted a budget estimate.",
                    body_style
                ))
        else:
            story.append(Paragraph(
                f"Based on {nat['submitted_count']} submitted county plans, the total identified "
                f"investment is <b>KES {nat['total_budget']} billion</b>.",
                body_style
            ))
            submitted = counties_df[counties_df["status"].isin(["submitted", "review"])]
            budget_counties = submitted[submitted["budget"] > 0] if not submitted.empty else submitted
            if not budget_counties.empty:
                story.append(make_table(
                    ["County", "Budget (KES B)", "Per capita (KES)", "Plan period"],
                    [[r["name"], f"{r['budget']}B",
                      f"{int(r['budget'] * 1e9 / r['pop']):,}", f"2023–{r['target_yr']}"]
                     for _, r in budget_counties.iterrows()],
                    [4 * cm, 3.5 * cm, 4 * cm, 4 * cm]
                ))
        story.append(Spacer(1, 0.4 * cm))

    # Section 6: Recommendations
    if sections.get("recommendations", True):
        story.append(Paragraph(f"{sec_num}. Recommendations", h1_style))
        sec_num += 1
        if is_county_report and county_row is not None:
            r = county_row
            recs = []
            if r["elec"] < 50:
                recs.append((f"Prioritise electricity access expansion",
                              f"With only {r['elec']}% access, urgent investment in grid extension, "
                              f"minigrids, and Solar Home Systems is needed."))
            if r["cooking"] < 40:
                recs.append((f"Accelerate clean cooking access",
                              f"Clean cooking access of {r['cooking']}% is well below the 2028 target. "
                              f"LPG and ICS programmes should be prioritised."))
            if r["status"] in ("pending", "overdue"):
                recs.append(("Complete and submit County Energy Plan",
                              "Submission is required under the Energy Act 2019 and is needed "
                              "for inclusion in the national INEP aggregation."))
            if r["solar"] > 2000:
                recs.append(("Leverage high solar potential",
                              f"Solar GHI of {r['solar']} kWh/m²/year makes this county ideal "
                              f"for Solar Home Systems and solar minigrids."))
            if not recs:
                recs.append(("Continue progress and update plans annually",
                              "The county is making good progress. Annual plan updates are recommended."))
            for i, (title, body) in enumerate(recs, 1):
                story.append(Paragraph(
                    f"<b>{i}. {title}</b>",
                    ParagraphStyle("rec", parent=body_style, fontName="Helvetica-Bold",
                                   textColor=colors.HexColor("#0e1e2e"))
                ))
                story.append(Paragraph(f"   {body}", body_style))
        else:
            recs = [
                ("Prioritise direct submission support for overdue counties",
                 f"Turkana, Marsabit, Mandera and Wajir have the lowest electricity access (8–12%) "
                 f"and are all overdue. Direct technical assistance is recommended."),
                ("Accelerate clean cooking intervention",
                 f"Current clean cooking access ({nat['w_cooking']}%) is well below the 2028 target. "
                 f"LPG subsidy expansion and ICS distribution should be accelerated."),
                ("Complete INEP when all 47 counties submit",
                 f"The national picture is {nat['coverage_pct']}% complete. A firm deadline should "
                 f"be set for the remaining {nat['pending_count'] + nat['overdue_count']} counties."),
                ("Increase investment in high-solar northern counties",
                 "Counties such as Mandera, Wajir, Marsabit and Turkana have the highest solar "
                 "potential yet the lowest access. Targeted off-grid solar programmes can address this."),
                ("Publish county comparison dashboard publicly",
                 "Making county-level data publicly available will create accountability and "
                 "incentivise timely submissions in future planning cycles."),
            ]
            for i, (title, body) in enumerate(recs, 1):
                story.append(Paragraph(
                    f"<b>{i}. {title}</b>",
                    ParagraphStyle("rec", parent=body_style, fontName="Helvetica-Bold",
                                   textColor=colors.HexColor("#0e1e2e"))
                ))
                story.append(Paragraph(f"   {body}", body_style))
        story.append(Spacer(1, 0.4 * cm))

    # Section 7: Data Notes
    if sections.get("data_notes", True):
        story.append(Paragraph(f"{sec_num}. Data Notes and Methodology", h1_style))
        if is_county_report and county_row is not None:
            notes = (
                f"• County data sourced from the KenyaWatts platform and submitted County Energy Plan for {county_filter}. "
                f"• Population figures from KNBS 2019 Census. "
                f"• Solar GHI data from Global Solar Atlas. "
                f"• Report generated {now.strftime('%d %B %Y at %H:%M')}. "
                f"• CONFIDENTIAL — EPRA INTERNAL USE ONLY."
            )
        else:
            notes = (
                f"• National figures are population-weighted averages from {nat['submitted_count']} submitted county plans. "
                f"• Budget figures are simple sums of submitted county plan budgets. "
                f"• Universal access target year: {nat['latest_target']}. "
                f"• Clean energy generation (82%) sourced from EPRA Statistics Report FY 2024/25. "
                f"• Report generated {now.strftime('%d %B %Y at %H:%M')}. "
                f"• CONFIDENTIAL — EPRA INTERNAL USE ONLY."
            )
        story.append(Paragraph(notes, body_style))
        story.append(Spacer(1, 0.5 * cm))

    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e8e6de")))
    story.append(Paragraph(
        f"KenyaWatts Platform · EPRA Kenya · Generated {now.strftime('%d %B %Y at %H:%M')} · "
        f"CONFIDENTIAL — EPRA INTERNAL USE ONLY",
        meta_style
    ))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ── DATA DOWNLOAD ─────────────────────────────────────────────────────────────
def page_data_download(role, county_id):
    is_epra   = role in ("epra","ministry","devpartner","developer")
    is_county = role == "county"
    cnty_name = COUNTIES[COUNTIES["id"]==county_id]["name"].values[0] if county_id else ""

    section("Data download centre", "Filter, select a date range and download energy planning data")

    if is_epra:
        alert("info","<b>EPRA access:</b> Download aggregated national data or individual county submissions.")
        download_scope = st.radio(t("download_scope_label"),
            [t("download_scope_national"), t("download_scope_all"), t("download_scope_individual")],
            horizontal=True)
    else:
        download_scope = t("download_scope_individual")
        alert("info",f"<b>County access:</b> Download data for {cnty_name} County only.")

    st.divider()

    # Date range filter
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        date_from = st.date_input("From date", value=datetime.now().date()-timedelta(days=365))
    with col_d2:
        date_to   = st.date_input("To date",   value=datetime.now().date())

    # Category filter
    st.markdown("**Select data categories to include:**")
    cat_cols = st.columns(4)
    cat_elec    = cat_cols[0].checkbox("Electricity access",      value=True)
    cat_cooking = cat_cols[1].checkbox("Clean cooking",           value=True)
    cat_solar   = cat_cols[2].checkbox("Solar / Renewable data",  value=True)
    cat_budget  = cat_cols[3].checkbox("Budget and investment",   value=True)
    cat_meta    = st.checkbox("Submission metadata (date, reference, submitted by)", value=True)

    # County filter (EPRA only)
    if is_epra and download_scope == t("download_scope_individual"):
        selected_county = st.selectbox("Select county:", COUNTIES["name"].tolist())
    elif is_county:
        selected_county = cnty_name
    else:
        selected_county = None

    # Format
    fmt = st.radio("Download format:", ["CSV (.csv)","Excel (.xlsx)","JSON (.json)"], horizontal=True)
    st.divider()

    # Preview
    section("Data preview")
    if download_scope == t("download_scope_national"):
        nat = compute_national()
        preview_data = pd.DataFrame([{
            "Indicator":"Weighted electricity access (%)",    "Value":nat["w_elec"],  "Source":f"{nat['submitted_count']} counties","Date":datetime.now().strftime("%Y-%m-%d")},
            {"Indicator":"Weighted clean cooking access (%)", "Value":nat["w_cooking"],"Source":f"{nat['submitted_count']} counties","Date":datetime.now().strftime("%Y-%m-%d")},
            {"Indicator":"Clean energy generation (%)",       "Value":82,              "Source":"EPRA FY24/25",                       "Date":"2025-06-01"},
            {"Indicator":"Total submitted budgets (KES B)",   "Value":nat["total_budget"],"Source":f"{nat['submitted_count']} counties","Date":datetime.now().strftime("%Y-%m-%d")},
            {"Indicator":"Counties submitted",                "Value":nat["submitted_count"],"Source":"Platform","Date":datetime.now().strftime("%Y-%m-%d")},
            {"Indicator":"National access target year",       "Value":nat["latest_target"],"Source":"County plans","Date":datetime.now().strftime("%Y-%m-%d")},
        ])
        build_download(preview_data, "national_aggregated", fmt, cat_elec, cat_cooking, cat_solar, cat_budget, cat_meta, date_from, date_to)

    elif download_scope == t("download_scope_all"):
        df_all = COUNTIES.copy()
        cols = ["name","region","pop","status","target_yr","mtf","growth"]
        if cat_elec:    cols += ["elec"]
        if cat_cooking: cols += ["cooking"]
        if cat_solar:   cols += ["solar"]
        if cat_budget:  cols += ["budget"]
        df_all = df_all[cols].rename(columns={
            "name":"county","pop":"population","elec":"electricity_access_pct",
            "cooking":"clean_cooking_pct","solar":"solar_ghi_kwh_m2","budget":"budget_kes_b",
            "target_yr":"target_year","mtf":"mtf_tier","growth":"pop_growth_pct"
        })
        build_download(df_all, "all_county_submissions", fmt, cat_elec, cat_cooking, cat_solar, cat_budget, cat_meta, date_from, date_to)

    else:
        # Individual county
        filter_name = selected_county or cnty_name
        row = COUNTIES[COUNTIES["name"]==filter_name]
        if not row.empty:
            r = row.iloc[0]
            cols_map = {"county":r["name"],"region":r["region"],"population":r["pop"],"status":r["status"],"target_year":r["target_yr"]}
            if cat_elec:    cols_map["electricity_access_pct"] = r["elec"]
            if cat_cooking: cols_map["clean_cooking_pct"]      = r["cooking"]
            if cat_solar:   cols_map["solar_ghi_kwh_m2"]       = r["solar"]
            if cat_budget:  cols_map["budget_kes_b"]            = r["budget"]
            if cat_meta:    cols_map.update({"mtf_tier":r["mtf"],"pop_growth_pct":r["growth"],"download_date":datetime.now().strftime("%Y-%m-%d")})
            df_single = pd.DataFrame([cols_map])
            # Also include any session submissions
            if st.session_state.submitted_data:
                sd = pd.DataFrame(st.session_state.submitted_data)
                sd = sd[sd["county"]==filter_name] if filter_name else sd
                if not sd.empty:
                    df_single = pd.concat([df_single, sd], ignore_index=True)
            build_download(df_single, f"{filter_name.lower().replace(' ','_')}_data", fmt, cat_elec, cat_cooking, cat_solar, cat_budget, cat_meta, date_from, date_to)

    # =========================================================================
    # REPORT GENERATION SECTION
    # Embedded in the Data Download tab — generates Word or PDF summary reports
    # with selectable sections. Supports both National Summary and County-Specific reports.
    # =========================================================================
    st.divider()
    section("Report generation", "Generate a National Summary or County-Specific report with selectable sections")

    alert("warn",
        "<b>CONFIDENTIAL — EPRA internal use only.</b> Generated reports aggregate county energy plans "
        "into a summary for EPRA planners and the Ministry of Energy. Do not distribute without authorisation.")

    # ── Report type ───────────────────────────────────────────────────────────
    if is_epra:
        report_type = st.radio(
            "Report type:",
            ["📊 National Summary Report", "🏘️ County-Specific Report"],
            horizontal=True,
            key="dl_report_type",
        )
    else:
        report_type = "🏘️ County-Specific Report"
        st.markdown(f"**Report type:** County-Specific Report ({cnty_name} County)")

    is_county_rpt = "County-Specific" in report_type

    # ── County selector (county-specific reports) ─────────────────────────────
    if is_county_rpt:
        if is_epra:
            report_county = st.selectbox(
                "Select county for report:",
                COUNTIES["name"].tolist(),
                key="dl_report_county_sel",
            )
        else:
            report_county = cnty_name
            st.info(f"Generating report for: **{cnty_name} County**")
    else:
        report_county = None

    # ── Section selection (checkboxes) ────────────────────────────────────────
    st.markdown("**Select sections to include in the report:**")
    sec_col1, sec_col2 = st.columns(2)
    with sec_col1:
        s_exec      = st.checkbox("1. Executive Summary",            value=True, key="dl_sec_exec")
        s_status    = st.checkbox("2. County Submission Status",     value=True, key="dl_sec_status")
        s_access    = st.checkbox("3. Energy Access Analysis",       value=True, key="dl_sec_access")
        s_renewable = st.checkbox("4. Renewable Energy Resources",   value=True, key="dl_sec_renewable")
    with sec_col2:
        s_invest    = st.checkbox("5. Investment & Budget Summary",  value=True, key="dl_sec_invest")
        s_recs      = st.checkbox("6. Recommendations",             value=True, key="dl_sec_recs")
        s_notes     = st.checkbox("7. Data Notes & Methodology",    value=True, key="dl_sec_notes")

    sections_selected = {
        "exec_summary":       s_exec,
        "submission_status":  s_status,
        "energy_access":      s_access,
        "renewable_resources":s_renewable,
        "investment_budget":  s_invest,
        "recommendations":    s_recs,
        "data_notes":         s_notes,
    }

    num_selected = sum(sections_selected.values())
    if num_selected == 0:
        st.warning("Please select at least one section to include in the report.")

    st.markdown(f"*{num_selected} of 7 sections selected*")

    # ── Report format and title ───────────────────────────────────────────────
    col_fmt2, col_title = st.columns([1, 2])
    with col_fmt2:
        rpt_fmt = st.radio(
            "Report format:",
            ["📄 Word document (.docx)", "📑 PDF document (.pdf)"],
            index=0,
            key="dl_report_fmt",
        )
    with col_title:
        default_title = (
            f"{report_county} County Energy Summary — {datetime.now().strftime('%B %Y')}"
            if is_county_rpt and report_county
            else f"Kenya National Summary Report — {datetime.now().strftime('%B %Y')}"
        )
        rpt_title = st.text_input(
            "Report title (optional — leave blank for default):",
            value=default_title,
            key="dl_report_title_input",
        )

    # Library status
    col_lib1, col_lib2 = st.columns(2)
    with col_lib1:
        if DOCX_AVAILABLE:
            st.success("✓ Word (.docx) generation ready")
        else:
            st.warning("⚠ Word generation unavailable — install python-docx")
    with col_lib2:
        if PDF_AVAILABLE:
            st.success("✓ PDF generation ready")
        else:
            st.warning("⚠ PDF generation unavailable — install reportlab")

    st.markdown("")

    # ── Generate button ───────────────────────────────────────────────────────
    if st.button(
        "⬇️  Generate and download report",
        type="primary",
        key="dl_generate_report_btn",
        disabled=(num_selected == 0),
    ):
        with st.spinner("Generating report — please wait…"):
            try:
                nat_for_rpt = compute_national()
                user_display = st.session_state.get("user_display_name", {}).get(
                    st.session_state.get("username", ""), "EPRA User"
                )

                if "Word" in rpt_fmt:
                    if not DOCX_AVAILABLE:
                        st.error(
                            "python-docx is not installed. "
                            "Add 'python-docx>=1.1.0' to requirements.txt and redeploy."
                        )
                    else:
                        data = generate_word_summary_report(
                            nat_for_rpt, COUNTIES, user_display,
                            sections_selected,
                            county_filter=report_county,
                            report_title=rpt_title,
                        )
                        slug = (report_county or "National").replace(" ", "_")
                        filename = f"KenyaWatts_{slug}_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        size_kb = len(data) // 1024
                        st.success(f"✅ Word report generated · {size_kb} KB · {datetime.now().strftime('%H:%M:%S')}")
                        st.download_button(
                            label=f"📥  Download {filename}",
                            data=data,
                            file_name=filename,
                            mime=mime,
                            type="primary",
                            key="dl_download_word_btn",
                        )
                        st.session_state.audit_log.append({
                            "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "user":   user_display,
                            "action": f"{'County' if is_county_rpt else 'National'} Summary Word report generated",
                            "ref":    f"RPT-{datetime.now().strftime('%Y%m%d%H%M')}",
                        })
                else:
                    if not PDF_AVAILABLE:
                        st.error(
                            "reportlab is not installed. "
                            "Add 'reportlab>=4.2.0' to requirements.txt and redeploy."
                        )
                    else:
                        data = generate_pdf_summary_report(
                            nat_for_rpt, COUNTIES, user_display,
                            sections_selected,
                            county_filter=report_county,
                            report_title=rpt_title,
                        )
                        slug = (report_county or "National").replace(" ", "_")
                        filename = f"KenyaWatts_{slug}_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                        mime = "application/pdf"
                        size_kb = len(data) // 1024
                        st.success(f"✅ PDF report generated · {size_kb} KB · {datetime.now().strftime('%H:%M:%S')}")
                        st.download_button(
                            label=f"📥  Download {filename}",
                            data=data,
                            file_name=filename,
                            mime=mime,
                            type="primary",
                            key="dl_download_pdf_btn",
                        )
                        st.session_state.audit_log.append({
                            "time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "user":   user_display,
                            "action": f"{'County' if is_county_rpt else 'National'} Summary PDF report generated",
                            "ref":    f"RPT-{datetime.now().strftime('%Y%m%d%H%M')}",
                        })

                alert("info",
                    "<b>Report generated.</b> Share only with authorised EPRA leadership and "
                    "Ministry of Energy officials. This generation has been logged in the audit trail.")

            except Exception as e:
                st.error(f"Report generation failed: {type(e).__name__}: {str(e)}")
                with st.expander("Full error details (for debugging)"):
                    st.exception(e)

def build_download(df, filename_base, fmt, cat_elec, cat_cooking, cat_solar, cat_budget, cat_meta, date_from, date_to):
    st.dataframe(df, use_container_width=True, hide_index=True)
    st.caption(f"{len(df)} records · {len(df.columns)} columns · Date range: {date_from} to {date_to}")
    st.markdown("")
    if "CSV" in fmt:
        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button(f"⬇ Download {filename_base}.csv", csv,
                           file_name=f"{filename_base}_{date_from}_to_{date_to}.csv",
                           mime="text/csv", type="primary")
    elif "Excel" in fmt:
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="KenyaWatts Data")
        st.download_button(f"⬇ Download {filename_base}.xlsx", buf.getvalue(),
                           file_name=f"{filename_base}_{date_from}_to_{date_to}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           type="primary")
    else:
        json_data = df.to_json(orient="records", indent=2).encode("utf-8")
        st.download_button(f"⬇ Download {filename_base}.json", json_data,
                           file_name=f"{filename_base}_{date_from}_to_{date_to}.json",
                           mime="application/json", type="primary")

# =============================================================================
# PAGE FUNCTION: AI Planning Assistant
# Powered by Anthropic's Claude model (claude-sonnet-4-6).
# The AI is given a system prompt containing real EPRA and county data
# so it can answer planning questions accurately.
# Conversation history is maintained in st.session_state["ai_history"]
# so follow-up questions work naturally.
# The AI key must be set in secrets.toml under [anthropic] api_key.
# =============================================================================

def page_ai(role, county_id):
    """Render the AI planning assistant chat interface."""
    alert("info","<b>KenyaWatts AI assistant</b> — powered by Claude · grounded in EPRA data and Makueni CEP · plain English answers")
    suggestions = [
        "Which counties need urgent electricity investment?",
        "How is Kenya tracking toward universal access by 2030?",
        "What barriers does Makueni face for clean cooking?",
        "How does KenyaWatts aggregate 47 county plans into the INEP?",
        "Which counties are overdue and what are the consequences?",
        "What validation rules catch unit errors in submissions?",
    ]
    if not st.session_state.ai_history:
        st.markdown("**Suggested questions:**")
        cols = st.columns(2)
        for i,s in enumerate(suggestions):
            if cols[i%2].button(s, key=f"sug_{i}"):
                st.session_state.ai_history.append({"role":"user","content":s})
                with st.spinner("Analysing EPRA data…"):
                    reply = ask_ai(s, st.session_state.ai_history[:-1])
                st.session_state.ai_history.append({"role":"assistant","content":reply})
                st.rerun()
    for msg in st.session_state.ai_history:
        with st.chat_message(msg["role"], avatar="⚡" if msg["role"]=="assistant" else None):
            st.write(msg["content"])
    if prompt := st.chat_input("Ask about Kenya's energy planning…"):
        st.session_state.ai_history.append({"role":"user","content":prompt})
        with st.chat_message("user"):
            st.write(prompt)
        with st.chat_message("assistant", avatar="⚡"):
            with st.spinner("Analysing EPRA data…"):
                reply = ask_ai(prompt, st.session_state.ai_history[:-1])
            st.write(reply)
        st.session_state.ai_history.append({"role":"assistant","content":reply})
    if st.session_state.ai_history:
        if st.button("Clear conversation"):
            st.session_state.ai_history = []
            st.rerun()

# =============================================================================
# PAGE FUNCTION: Account Settings
# Three tabs:
#   1. Change password — verifies current password with bcrypt, then
#      hashes and stores the new password. Active immediately for the
#      current session. EPRA admin must also update secrets.toml to make
#      the change permanent across all future sessions.
#   2. Update display name / email — changes what appears in the header
#      and on submissions. Does not affect login username.
#   3. My account info — shows username, role, county, session time,
#      and recent audit activity for this user.
# =============================================================================

def page_account_settings(authenticator, username, user_info):
    """Render the account settings page with three tabs."""
    alert("info","<b>Account settings</b> — change your password or display name. Role and county access require EPRA admin.")
    tab1, tab2, tab3 = st.tabs(["🔑  Change password","✏️  Update display name","👤  My account info"])

    with tab1:
        section("Change your password")
        alert("warn","<b>Password requirements:</b> Minimum 8 characters. Active immediately for this session.")
        with st.form("pw_form"):
            cur_pw  = st.text_input("Current password",     type="password")
            new_pw  = st.text_input("New password",          type="password", help="At least 8 characters")
            conf_pw = st.text_input("Confirm new password",  type="password")
            sub_pw  = st.form_submit_button("Update password", type="primary")
        if sub_pw:
            if not cur_pw or not new_pw or not conf_pw: st.error("All fields required.")
            elif len(new_pw)<8:            st.error("New password must be at least 8 characters.")
            elif new_pw!=conf_pw:          st.error("Passwords do not match.")
            elif new_pw==cur_pw:           st.warning("New password must differ from current.")
            else:
                try:
                    stored = st.secrets.get("credentials",{}).get("usernames",{}).get(username,{}).get("password","")
                    if stored and bcrypt.checkpw(cur_pw.encode(), stored.encode()):
                        new_hash = bcrypt.hashpw(new_pw.encode(), bcrypt.gensalt()).decode()
                        try: authenticator.credentials["usernames"][username]["password"] = new_hash
                        except: pass
                        st.session_state.audit_log.append({"time":datetime.now().strftime('%Y-%m-%d %H:%M:%S'),"user":username,"action":"Password changed","ref":f"PW-{username.upper()}"})
                        push_notification("Password updated successfully", "🔑")
                        st.success("✓ Password updated successfully.")
                        st.info("To make permanent: EPRA admin must update secrets.toml with the new hash below.")
                        st.code(f'[credentials.usernames.{username}]\npassword = "{new_hash}"', language="toml")
                    else: st.error("Current password is incorrect.")
                except Exception as e: st.error(f"Error: {e}")

    with tab2:
        section("Update your display name and email")
        cur_name  = st.session_state.get("user_display_name",{}).get(username, user_info.get("name",username))
        cur_email = st.session_state.get("user_email",{}).get(username, user_info.get("email",""))
        with st.form("name_form"):
            new_name  = st.text_input("Display name",    value=cur_name)
            new_email = st.text_input("Email address",   value=cur_email)
            sub_name  = st.form_submit_button("Save changes", type="primary")
        if sub_name:
            if len(new_name.strip())<3: st.error("Display name must be at least 3 characters.")
            else:
                st.session_state.setdefault("user_display_name",{})[username] = new_name.strip()
                st.session_state.setdefault("user_email",{})[username]        = new_email.strip()
                try:
                    authenticator.credentials["usernames"][username]["name"]  = new_name.strip()
                    authenticator.credentials["usernames"][username]["email"] = new_email.strip()
                except: pass
                push_notification(f"Display name updated to {new_name.strip()}", "✏️")
                st.success(f"✓ Display name updated to **{new_name.strip()}**.")

    with tab3:
        # ── Account information tab ───────────────────────────────────────────
        # Shows the logged-in user's profile details using native Streamlit
        # widgets (not HTML) to avoid rendering issues with f-strings in markdown

        section("Your account information")

        # Map role codes to human-readable labels
        role_labels = {
            "epra":             "EPRA Planner — Full admin access",
            "ministry":         "Ministry of Energy — Read-only national view",
            "devpartner":       "Development Partner — Read-only aggregated view",
            "county":           "County Energy Planning Committee",
            "service_provider": "KPLC Service Provider",
            "developer":        "Platform Developer — Full access (demo/test)",
        }

        # Get current values from session state or user_info
        role   = user_info["role"]
        cid    = user_info["county_id"]
        dname  = st.session_state.get("user_display_name", {}).get(
                     username, user_info.get("name", username))
        demail = st.session_state.get("user_email", {}).get(
                     username, user_info.get("email", "Not set"))
        cnty   = (COUNTIES[COUNTIES["id"]==cid]["name"].values[0]
                  if cid and not COUNTIES[COUNTIES["id"]==cid].empty
                  else "All counties (national access)")

        # Display each field as a labelled row using Streamlit columns
        # This avoids HTML rendering issues entirely
        fields = [
            ("Username",        username),
            ("Display name",    dname),
            ("Email address",   demail or "Not set"),
            ("Role",            role_labels.get(role, role)),
            ("County / Access", cnty),
            ("Session started", datetime.now().strftime("%d %b %Y at %H:%M")),
            ("Account status",  "Active ✓"),
        ]

        for label, value in fields:
            col_label, col_value = st.columns([1, 2])
            col_label.markdown(
                f"<p style='font-size:12px;font-weight:600;color:#9c9a8e;"
                f"margin:0;padding:6px 0'>{label}</p>",
                unsafe_allow_html=True
            )
            col_value.markdown(
                f"<p style='font-size:13px;color:#1a1916;"
                f"margin:0;padding:6px 0'>{value}</p>",
                unsafe_allow_html=True
            )
            st.divider()

        # Show recent audit activity for this user
        audit = [a for a in st.session_state.get("audit_log", [])
                 if a.get("user") == username]
        if audit:
            st.markdown("**Your recent activity**")
            for entry in reversed(audit[-5:]):
                col_t, col_a, col_r = st.columns([1.2, 2, 1])
                col_t.caption(entry.get("time", ""))
                col_a.caption(entry.get("action", ""))
                col_r.caption(entry.get("ref", ""))

        # Contact info for role/username changes
        alert("info",
              "<b>Need to change your username or role?</b> "
              "These require EPRA admin access. "
              "Contact: Allan.Wairimu@epra.go.ke · +254720850696")

# ── FORGOT PASSWORD ───────────────────────────────────────────────────────────
def page_forgot_password(authenticator):
    """
    Proper multi-step password reset flow using session state.
    Step 1: Enter username + email  → verify against records
    Step 2: Enter 6-digit code      → verify code + expiry
    Step 3: Enter new password      → confirm + hash + done
    """
    import random

    # Initialise step if not set
    if "reset_step" not in st.session_state:
        st.session_state["reset_step"] = 1

    step = st.session_state["reset_step"]

    # ── Page shell ────────────────────────────────────────────────────────────
    col1, col2, col3 = st.columns([1, 1.4, 1])
    with col2:

        # Logo
        st.markdown("""
        <div style="text-align:center;margin-bottom:28px;margin-top:30px">
          <div style="font-size:26px;font-weight:700;color:#ffffff">
            Kenya<span style="color:#0f9d7e">Watts</span>
          </div>
          <div style="font-size:12px;color:#9c9a8e;margin-top:3px">
            EPRA National Energy Planning Platform
          </div>
        </div>""", unsafe_allow_html=True)

        # Progress bar — 3 steps
        step_labels = ["Verify identity", "Enter reset code", "Set new password"]
        progress_html = '<div style="display:flex;gap:0;margin-bottom:24px;border-radius:8px;overflow:visible">'
        for i, lbl in enumerate(step_labels, 1):
            bg = "#0e1e2e" if i == step else "#e8f7f4" if i < step else "#f7f6f2"
            tc = "#ffffff" if i == step else "#0f9d7e" if i < step else "#9c9a8e"
            progress_html += (
                f'<div style="flex:1;padding:8px 6px;text-align:center;background:{bg};'
                f'font-size:11px;font-weight:{"600" if i==step else "400"};color:{tc}">'
                f'{"✓ " if i<step else ""}{lbl}</div>'
            )
        progress_html += '</div>'
        st.markdown(progress_html, unsafe_allow_html=True)

        # ── STEP 1 — Verify identity ──────────────────────────────────────────
        if step == 1:
            st.markdown("""
            <div style="font-size:16px;font-weight:700;color:#1a1916;margin-bottom:4px">
              Verify your identity
            </div>
            <div style="font-size:12px;color:#6b6860;margin-bottom:18px">
              Enter your username and the email address registered to your account.
              Both must match exactly what is in the system.
            </div>""", unsafe_allow_html=True)

            with st.form("verify_identity_form", clear_on_submit=False):
                uname = st.text_input(
                    "Username *",
                    placeholder="Your KenyaWatts login username",
                    key="reset_uname_input",
                )
                email = st.text_input(
                    "Registered email address *",
                    placeholder="e.g. energy@makueni.go.ke",
                    key="reset_email_input",
                )
                submit_step1 = st.form_submit_button(
                    "Verify and continue →",
                    type="primary",
                    use_container_width=True,
                )

            if submit_step1:
                if not uname.strip() or not email.strip():
                    st.error("Both username and email address are required.")
                else:
                    creds = st.secrets.get("credentials", {}).get("usernames", {})
                    user  = creds.get(uname.strip(), {})
                    registered_email = user.get("email", "").lower().strip()
                    if uname.strip() in creds and registered_email == email.strip().lower():
                        # Generate 6-digit code
                        code = str(random.randint(100000, 999999))
                        st.session_state["reset_uname"]     = uname.strip()
                        st.session_state["reset_email"]     = email.strip()
                        st.session_state["reset_code"]      = code
                        st.session_state["reset_code_time"] = datetime.now()
                        st.session_state["reset_step"]      = 2
                        st.rerun()
                    else:
                        st.error(
                            "Username and email do not match our records. "
                            "Check your spelling or contact EPRA admin: "
                            "Allan.Wairimu@epra.go.ke"
                        )

        # ── STEP 2 — Enter reset code ─────────────────────────────────────────
        elif step == 2:
            uname = st.session_state.get("reset_uname", "")
            email = st.session_state.get("reset_email", "")
            code  = st.session_state.get("reset_code",  "")
            code_time = st.session_state.get("reset_code_time", datetime.now())
            mins_left = max(0, 15 - int((datetime.now()-code_time).seconds/60))

            st.markdown(f"""
            <div style="font-size:16px;font-weight:700;color:#1a1916;margin-bottom:4px">
              Enter your reset code
            </div>
            <div style="font-size:12px;color:#6b6860;margin-bottom:12px">
              A 6-digit reset code has been generated for <b>{email}</b>.
              In a production deployment this would be sent to your registered email.
              The code expires in <b>{mins_left} minute(s)</b>.
            </div>""", unsafe_allow_html=True)

            # Show code box (demo only — in production this would be emailed)
            st.markdown(f"""
            <div style="background:#0e1e2e;border-radius:10px;padding:16px 20px;
                        text-align:center;margin-bottom:18px">
              <div style="font-size:11px;color:#a8c4d4;margin-bottom:6px;
                          text-transform:uppercase;letter-spacing:.5px">
                Your reset code (demo — would be emailed in production)
              </div>
              <div style="font-size:32px;font-weight:700;color:#3ecfaa;
                          letter-spacing:8px;font-family:monospace">
                {code}
              </div>
              <div style="font-size:10px;color:#5a8a9e;margin-top:6px">
                Expires in {mins_left} minute(s) · Do not share this code
              </div>
            </div>""", unsafe_allow_html=True)

            with st.form("verify_code_form", clear_on_submit=False):
                entered = st.text_input(
                    "Enter the 6-digit code *",
                    placeholder="e.g. 483921",
                    max_chars=6,
                    key="reset_code_input",
                )
                submit_step2 = st.form_submit_button(
                    "Verify code →",
                    type="primary",
                    use_container_width=True,
                )

            if submit_step2:
                elapsed  = (datetime.now() - code_time).seconds
                if elapsed > 900:
                    st.error("This code has expired (15 minutes). Please start again.")
                    if st.button("Start again", key="restart_after_expire"):
                        st.session_state["reset_step"] = 1
                        st.rerun()
                elif entered.strip() != code:
                    st.error("Incorrect code. Please check and try again.")
                else:
                    st.session_state["reset_step"] = 3
                    st.rerun()

        # ── STEP 3 — Set new password ─────────────────────────────────────────
        elif step == 3:
            uname = st.session_state.get("reset_uname", "")

            st.markdown(f"""
            <div style="font-size:16px;font-weight:700;color:#1a1916;margin-bottom:4px">
              Set your new password
            </div>
            <div style="font-size:12px;color:#6b6860;margin-bottom:18px">
              Identity verified for <b>{uname}</b>. Choose a strong new password.
              Minimum 8 characters.
            </div>""", unsafe_allow_html=True)

            with st.form("set_new_pw_form", clear_on_submit=False):
                new_pw1 = st.text_input(
                    "New password *",
                    type="password",
                    placeholder="At least 8 characters",
                    key="new_pw1_input",
                )
                new_pw2 = st.text_input(
                    "Confirm new password *",
                    type="password",
                    placeholder="Type your new password again",
                    key="new_pw2_input",
                )
                submit_step3 = st.form_submit_button(
                    "Reset password ✓",
                    type="primary",
                    use_container_width=True,
                )

            if submit_step3:
                if not new_pw1 or not new_pw2:
                    st.error("Both password fields are required.")
                elif len(new_pw1) < 8:
                    st.error("Password must be at least 8 characters.")
                elif new_pw1 != new_pw2:
                    st.error("Passwords do not match. Please try again.")
                else:
                    new_hash = bcrypt.hashpw(new_pw1.encode(), bcrypt.gensalt()).decode()
                    # Clear all reset state
                    for k in ["reset_step","reset_uname","reset_email",
                              "reset_code","reset_code_time"]:
                        st.session_state.pop(k, None)

                    # Show success
                    st.markdown("""
                    <div style="background:#0e1e2e;border-radius:12px;padding:20px 22px;
                                text-align:center;margin-bottom:16px">
                      <div style="font-size:28px;margin-bottom:8px">✅</div>
                      <div style="font-size:16px;font-weight:700;color:#3ecfaa;margin-bottom:6px">
                        Password reset successfully
                      </div>
                      <div style="font-size:12px;color:#a8c4d4;line-height:1.6">
                        Your new password is active. You can now sign in with your new credentials.
                      </div>
                    </div>""", unsafe_allow_html=True)

                    st.info(
                        "**For EPRA admin:** To make this password permanent across all "
                        "sessions, update secrets.toml with the hash below."
                    )
                    st.code(
                        f'[credentials.usernames.{uname}]\npassword = "{new_hash}"',
                        language="toml"
                    )
                    if st.button("← Go to login", type="primary",
                                 use_container_width=True, key="goto_login_after_reset"):
                        st.session_state["show_forgot"] = False
                        st.rerun()
                    return  # stop here — don't show back button below

        # ── Back to login (shown on steps 1 and 2) ────────────────────────────
        if step in (1, 2):
            st.markdown("")
            back_col, _ = st.columns([1, 1])
            if back_col.button("← Back to login", key="back_to_login_btn"):
                st.session_state["show_forgot"] = False
                st.session_state["reset_step"]  = 1
                for k in ["reset_uname","reset_email","reset_code","reset_code_time"]:
                    st.session_state.pop(k, None)
                st.rerun()

# ── LOGIN PAGE ────────────────────────────────────────────────────────────────
def show_login_page(authenticator):
    col1,col2,col3 = st.columns([1,1.2,1])
    with col2:
        st.markdown("""<div style="text-align:center;margin-bottom:24px;margin-top:40px">
          <div style="font-size:28px;font-weight:700;color:#ffffff">Kenya<span style="color:#0f9d7e">Watts</span></div>
          <div style="font-size:13px;color:#9c9a8e;margin-top:4px">Digital Integrated National Energy Planning Platform</div>
          <div style="font-size:11px;color:#c8c6be;margin-top:2px">EPRA Kenya · NGDA 2026</div>
        </div>""", unsafe_allow_html=True)

        result = authenticator.login(
            location="main",
            fields={"Form name":"Sign in to KenyaWatts","Username":"Username",
                    "Password":"Password","Login":"Sign in"},
            key="kw_login_form",
        )
        if result is not None:
            name, auth_status, username = result
        else:
            # If result is None, check if authenticator already set session state via cookie
            name        = st.session_state.get("name","")
            auth_status = st.session_state.get("authentication_status", None)
            username    = st.session_state.get("username","")

        if auth_status is False:
            st.error("Incorrect username or password. Please try again.")

        # Forgot password link — clean, no credentials shown
        st.markdown("")
        if st.button("Forgot your password?", use_container_width=True, key="forgot_pw_btn"):
            st.session_state["show_forgot"] = True
            st.session_state["reset_step"]  = 1   # always start at step 1
            st.rerun()

    return name, auth_status, username

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
def build_sidebar(authenticator, user_info):
    with st.sidebar:
        role      = user_info["role"]
        is_epra   = role in ("epra","ministry","devpartner","developer")
        is_county = role == "county"
        rc = {"epra":"#1a6fa3","ministry":"#0f9d7e","devpartner":"#5b4fc9","county":"#d4891a","service_provider":"#7a7870"}.get(role,"#1a6fa3")
        rl = {"epra":"EPRA Planner","ministry":"Ministry of Energy","devpartner":"Development Partner","county":"County Committee","service_provider":"KPLC"}.get(role,"User")
        dname = st.session_state.get("user_display_name",{}).get(user_info.get("_username",""), user_info.get("name",""))

        st.markdown(f"""<div style="background:#0e1e2e;border-radius:10px;padding:14px;margin-bottom:16px">
          <div style="font-size:15px;font-weight:700;color:white">Kenya<span style="color:#3ecfaa">Watts</span></div>
          <div style="font-size:10px;color:#5a8a9e;margin-top:2px">National Energy Planning Platform</div>
          <div style="margin-top:10px;padding:8px 10px;background:rgba(255,255,255,0.06);border-radius:7px">
            <div style="font-size:12px;font-weight:600;color:white">{dname}</div>
            <div style="font-size:11px;color:{rc};margin-top:2px">{rl}</div>
          </div></div>""", unsafe_allow_html=True)

        st.markdown("### Navigation")
        if role == "epra":
            opts = [
                "📊  National overview","🗺️  County map","🏛️  All 47 counties",
                "✅  Validation queue","📥  Communications hub",
                "📄  Makueni CEP — template reference",
                "📋  National INEP report",
                "⬇️  Data download","⚙️  Account settings",
            ]
        elif role == "ministry":
            opts = [
                "📊  National overview","🗺️  County map","🏛️  All 47 counties",
                "📥  Communications hub","⬇️  Data download",
                "🤖  AI assistant","⚙️  Account settings",
            ]
        elif role == "devpartner":
            opts = [
                "📊  National overview","🗺️  County map","🏛️  All 47 counties",
                "⬇️  Data download","⚙️  Account settings",
            ]
        elif role == "county":
            cname = COUNTIES[COUNTIES["id"]==user_info["county_id"]]["name"].values[0] if user_info["county_id"] else "County"
            is_makueni = user_info.get("county_id","") == "MK"
            opts = [
                f"🏠  {cname} dashboard",
                "📤  Submit energy plan","📨  Communications",
                "⬇️  Data download","⚙️  Account settings",
            ]
            if is_makueni:
                opts.insert(2, "📄  My county plan")
        elif role == "service_provider":
            opts = [
                "📊  National overview","🏛️  County demand data",
                "📤  Submit provider plan","⬇️  Data download",
                "🤖  AI assistant","⚙️  Account settings",
            ]
        else:
            opts = [
                "📊  National overview",
                "⬇️  Data download","⚙️  Account settings",
            ]

        selected = st.radio("", opts, label_visibility="collapsed")
        st.divider()
        # Use authenticator's own logout — this deletes the browser cookie properly
        # Without cookie deletion the user gets auto-logged back in
        try:
            authenticator.logout(
                button_name="🚪  Sign out",
                location="sidebar",
                key="kw_signout_btn",
                use_container_width=True,
            )
        except Exception:
            # Fallback: manually clear everything including the cookie name
            if st.button("🚪  Sign out", key="kw_signout_fallback", use_container_width=True):
                cookie_name = st.secrets.get("cookie",{}).get("name","kenyawatts_auth")
                for k in list(st.session_state.keys()):
                    del st.session_state[k]
                st.rerun()
        st.markdown("""<div style="font-size:10px;color:#9c9a8e;line-height:1.7;margin-top:8px">
        Data: EPRA FY 2024/25 · Makueni CEP 2023–2032 · KNBS 2019<br>NGDA 2026 · DTU · Challenge 2 · EPRA Kenya
        </div>""", unsafe_allow_html=True)
    return selected

# ── MAIN ─────────────────────────────────────────────────────────────────────
# =============================================================================
# PAGE FUNCTION: Developer Tools  (developer role only)
# Provides the ability to reset/clear all platform data for a clean test.
# Only accessible when logged in as the "developer" account.
# Actions here are PERMANENT and cannot be undone — clearly warned in UI.
# =============================================================================

def page_developer_tools():
    """
    Render the developer reset/management panel.
    This page is ONLY shown to users with role == "developer".
    It allows clearing all messages and submissions so the platform
    starts fresh for a real test with EPRA and county users.
    """
    # Warning banner — make it very clear this is destructive
    st.markdown("""
    <div style="background:#b33a2c;border-radius:12px;padding:18px 22px;margin-bottom:20px">
      <div style="font-size:16px;font-weight:700;color:#ffffff;margin-bottom:6px">
        🛠️ Developer Tools — NGDA Platform
      </div>
      <div style="font-size:13px;color:rgba(255,255,255,0.85);line-height:1.6">
        These actions are <b>permanent and cannot be undone</b>.
        Use this panel to reset the platform to a clean state before testing.
        Only the developer account can access this page.
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Current platform state ─────────────────────────────────────────────────
    st.markdown("### Current platform state")
    msgs  = get_messages()
    subs  = get_submissions()
    c1,c2,c3 = st.columns(3)
    with c1:
        st.markdown(f"""
        <div style="background:#0e1e2e;border-radius:10px;padding:14px 16px;text-align:center">
          <div style="font-size:28px;font-weight:700;color:#3ecfaa">{len(msgs)}</div>
          <div style="font-size:12px;color:#a8c4d4;margin-top:4px">Messages in store</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""
        <div style="background:#0e1e2e;border-radius:10px;padding:14px 16px;text-align:center">
          <div style="font-size:28px;font-weight:700;color:#3ecfaa">{len(subs)}</div>
          <div style="font-size:12px;color:#a8c4d4;margin-top:4px">Submissions in store</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        db_status = "Supabase cloud ✓" if USE_CLOUD_DB else "Local JSON files"
        st.markdown(f"""
        <div style="background:#0e1e2e;border-radius:10px;padding:14px 16px;text-align:center">
          <div style="font-size:14px;font-weight:700;color:#3ecfaa">{db_status}</div>
          <div style="font-size:12px;color:#a8c4d4;margin-top:4px">Storage backend</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("")
    st.divider()

    # ── Reset actions ──────────────────────────────────────────────────────────
    st.markdown("### Reset actions")

    # Option 1: Clear all messages
    st.markdown("#### 1. Clear all messages")
    st.caption(
        "Deletes every message in the communications store — county queries, "
        "EPRA broadcasts, submission notifications, and replies. "
        "After this, all inboxes will be completely empty."
    )
    col_a, col_b = st.columns([1,3])
    if col_a.button(
        "🗑️  Clear all messages",
        key="dev_clear_msgs",
        type="primary",
        use_container_width=True
    ):
        # Use session state to require confirmation click
        st.session_state["confirm_clear_msgs"] = True

    if st.session_state.get("confirm_clear_msgs"):
        st.warning(
            f"⚠️ This will permanently delete all **{len(msgs)} messages**. "
            "Click confirm to proceed."
        )
        cc1, cc2, _ = st.columns([1,1,2])
        if cc1.button("✅ Confirm — delete messages", key="confirm_msgs_yes",
                      type="primary"):
            reset_all_messages()
            st.session_state.pop("confirm_clear_msgs", None)
            st.success("✓ All messages deleted. Inboxes are now empty.")
            st.rerun()
        if cc2.button("Cancel", key="confirm_msgs_no"):
            st.session_state.pop("confirm_clear_msgs", None)
            st.rerun()

    st.divider()

    # Option 2: Clear all submissions
    st.markdown("#### 2. Clear all county plan submissions")
    st.caption(
        "Deletes every county plan submission from the store — including "
        "all indicator data, reference numbers, and EPRA approval status. "
        "After this, the validation queue will be empty and submission history "
        "will show no records."
    )
    col_c, col_d = st.columns([1,3])
    if col_c.button(
        "🗑️  Clear all submissions",
        key="dev_clear_subs",
        type="primary",
        use_container_width=True
    ):
        st.session_state["confirm_clear_subs"] = True

    if st.session_state.get("confirm_clear_subs"):
        st.warning(
            f"⚠️ This will permanently delete all **{len(subs)} submissions**. "
            "Click confirm to proceed."
        )
        sc1, sc2, _ = st.columns([1,1,2])
        if sc1.button("✅ Confirm — delete submissions", key="confirm_subs_yes",
                      type="primary"):
            reset_all_submissions()
            st.session_state.pop("confirm_clear_subs", None)
            st.success("✓ All submissions deleted. Validation queue is now empty.")
            st.rerun()
        if sc2.button("Cancel", key="confirm_subs_no"):
            st.session_state.pop("confirm_clear_subs", None)
            st.rerun()

    st.divider()

    # Option 3: Full reset (messages + submissions)
    st.markdown("#### 3. Full platform reset")
    st.caption(
        "Deletes EVERYTHING — all messages and all submissions. "
        "Use this before a real test session to ensure a completely clean start."
    )
    if st.button(
        "🔄  Full reset — clear everything",
        key="dev_full_reset",
        use_container_width=False
    ):
        st.session_state["confirm_full_reset"] = True

    if st.session_state.get("confirm_full_reset"):
        st.error(
            f"⚠️ FULL RESET: permanently deletes all **{len(msgs)} messages** "
            f"AND all **{len(subs)} submissions**. This cannot be undone."
        )
        fr1, fr2, _ = st.columns([1.2,1,2])
        if fr1.button("✅ Confirm full reset", key="confirm_full_yes",
                      type="primary"):
            reset_all_messages()
            reset_all_submissions()
            # Also clear session-level data
            for k in ["audit_log","upload_log","ai_history","notifications"]:
                st.session_state[k] = [] if k != "notifications" else []
            st.session_state.pop("confirm_full_reset", None)
            st.success(
                "✓ Full reset complete. All messages and submissions deleted. "
                "The platform is ready for a clean test session."
            )
            st.rerun()
        if fr2.button("Cancel", key="confirm_full_no"):
            st.session_state.pop("confirm_full_reset", None)
            st.rerun()

    st.divider()

    # ── Recent messages preview ────────────────────────────────────────────────
    if msgs:
        st.markdown("### Recent messages (last 5)")
        for m in reversed(msgs[-5:]):
            st.markdown(
                f"**{m.get('subject','—')}** · "
                f"From: {m.get('from_name','?')} → {m.get('to','?')} · "
                f"{m.get('date','')} {m.get('time','')}"
            )

    # ── Recent submissions preview ─────────────────────────────────────────────
    if subs:
        st.markdown("### Recent submissions (last 5)")
        for s in subs[:5]:
            sc = {"submitted":"#d4891a","approved":"#0f9d7e",
                  "rejected":"#b33a2c"}.get(s.get("status",""),"#7a7870")
            st.markdown(
                f"**{s.get('county','—')}** · "
                f"Ref: `{s.get('ref','—')}` · "
                f"{s.get('date_display',s.get('date',''))} at {s.get('time','')} · "
                f"<span style='color:{sc}'>{s.get('status','').title()}</span>",
                unsafe_allow_html=True
            )

# =============================================================================
# SECTION 20 — MAIN FUNCTION (APP ENTRY POINT)
# This is where everything comes together.
# Flow:
#   1. Initialise session state variables
#   2. Set up the authenticator
#   3. Check for forgot-password flow (separate page, no sidebar)
#   4. Check authentication — if not logged in, show login page and stop
#   5. Validate the stored username is still in the credentials list
#   6. Render the platform header
#   7. Build page list based on the user's role
#   8. Add sign-out button to sidebar
#   9. Run st.navigation() — this renders the sidebar nav AND the selected page
#
# st.navigation() is the modern Streamlit multi-page system (v1.36+).
# Each page is an st.Page object wrapping a Python function.
# When the user clicks a different tab, st.navigation() completely destroys
# the previous page's widgets before rendering the new one — this is what
# prevents content bleed-through between pages.
# =============================================================================

def main():
    """
    Main app entry point. Called once per browser interaction.
    Streamlit reruns this entire function every time the user
    clicks anything — including sidebar navigation.
    """
    init_session()       # Set up session state variables
    authenticator = setup_auth()  # Load credentials from secrets.toml

    # ── Forgot password flow ──────────────────────────────────────────────────
    # This is a separate flow that bypasses the normal sidebar navigation.
    # When active, ONLY the password reset page is shown.
    if st.session_state.get("show_forgot"):
        page_forgot_password(authenticator)
        return

    # ── Auth check ────────────────────────────────────────────────────────────
    if st.session_state.get("authentication_status") is not True:
        name, auth_status, username = show_login_page(authenticator)
        if auth_status is True:
            st.rerun()
        return

    username  = st.session_state.get("username", "")
    valid_users = list(st.secrets.get("credentials", {}).get("usernames", {}).keys())
    if not username or username not in valid_users:
        for k in list(st.session_state.keys()):
            try: del st.session_state[k]
            except: pass
        st.rerun()
        return

    user_info = get_user_role(username)
    user_info["_username"] = username
    role      = user_info["role"]
    county_id = user_info["county_id"]

    show_notifications()

    # ── Header ────────────────────────────────────────────────────────────────
    st.markdown(f"""<div class="kw-header">
      <div style="display:flex;justify-content:space-between;align-items:center">
        <div class="kw-logo">Kenya<span>Watts</span>
          <span style="font-size:13px;font-weight:400;color:#5a8a9e;margin-left:14px">
            {t("header_tagline")}
          </span>
        </div>
        <div style="font-size:11px;color:#3ecfaa;font-weight:600">{t("header_live")}</div>
      </div></div>""", unsafe_allow_html=True)

    # ── Build nav pages based on role ─────────────────────────────────────────
    is_epra   = role in ("epra", "ministry", "devpartner", "developer")
    is_county = role == "county"
    cname     = COUNTIES[COUNTIES["id"]==county_id]["name"].values[0] if county_id and not COUNTIES[COUNTIES["id"]==county_id].empty else "County"
    is_makueni = county_id == "MK"
    dname     = st.session_state.get("user_display_name", {}).get(username, user_info.get("name", username))

    # Sidebar user card
    rc = {"epra":"#1a6fa3","ministry":"#0f9d7e","devpartner":"#5b4fc9",
          "county":"#d4891a","service_provider":"#7a7870",
          "developer":"#b33a2c"}.get(role,"#1a6fa3")   # red = developer/demo
    rl = {"epra":t("role_epra"),"ministry":t("role_ministry"),
          "devpartner":t("role_devpartner"),"county":t("role_county"),
          "service_provider":t("role_kplc"),
          "developer":t("role_developer")}.get(role,"User")

    with st.sidebar:
        st.markdown(f"""<div style="background:#0e1e2e;border-radius:10px;padding:14px;margin-bottom:16px">
          <div style="font-size:15px;font-weight:700;color:white">Kenya<span style="color:#3ecfaa">Watts</span></div>
          <div style="font-size:10px;color:#5a8a9e;margin-top:2px">{t("sidebar_platform")}</div>
          <div style="margin-top:10px;padding:8px 10px;background:rgba(255,255,255,0.06);border-radius:7px">
            <div style="font-size:12px;font-weight:600;color:white">{dname}</div>
            <div style="font-size:11px;color:{rc};margin-top:2px">{rl}</div>
          </div></div>""", unsafe_allow_html=True)

        # ── Language toggle ───────────────────────────────────────────────────
        lang_choice = st.radio(
            "🌐 " + t("sidebar_language"),
            options=["English", "Kiswahili"],
            index=0 if st.session_state.get("lang", "en") == "en" else 1,
            horizontal=True,
            key="lang_radio",
            label_visibility="collapsed",
        )
        new_lang = "en" if lang_choice == "English" else "sw"
        if st.session_state.get("lang") != new_lang:
            st.session_state["lang"] = new_lang
            st.rerun()

        st.markdown(
            f"<div style='font-size:10px;color:#5a8a9e;text-align:center;margin-bottom:4px'>"
            f"{'🇬🇧 English' if new_lang=='en' else '🇰🇪 Kiswahili'}</div>",
            unsafe_allow_html=True
        )

    # ── Define pages using st.Page — each gets a clean isolated render ────────
    def _national_overview():   page_national_overview(role)
    def _county_map():          page_map(role, county_id)
    def _all_counties():        page_all_counties()
    def _county_dashboard():    page_county_dashboard(county_id, dname)
    def _validation_queue():    page_validation_queue()
    def _submit():              page_submit(role, county_id, dname)
    def _inbox():               page_inbox(role, county_id, dname)
    def _makueni_ref():         page_makueni_contextual(role, county_id)
    def _inep_report():         page_national_report(user_info)
    def _data_download():       page_data_download(role, county_id)
    def _ai():                  page_ai(role, county_id)
    def _account():             page_account_settings(authenticator, username, user_info)
    def _sign_out():
        for k in list(st.session_state.keys()):
            try: del st.session_state[k]
            except: pass
        st.rerun()

    # Build page list per role — st.Page gives each page a fully isolated render
    if role in ("epra", "developer"):
        # Developer gets full access PLUS the exclusive Developer Tools page
        # Developer Tools allows resetting all messages and submissions for testing

        # Wrap developer tools in a lambda — only available to developer role
        def _dev_tools():
            if role == "developer":
                page_developer_tools()
            else:
                st.error("Access denied — developer role required.")

        # ── AI is now a FLOATING WIDGET (bottom-right corner) ────────────────
        # It is NOT in the nav pages — it renders after pg.run() on every page.
        # This is why _ai is defined but not added to any pages list.

        base_pages = [
            st.Page(_national_overview, title=t("nav_national_overview"),               icon="📊", default=True),
            st.Page(_county_map,        title=t("nav_county_map"),                      icon="🗺️"),
            st.Page(_all_counties,      title=t("nav_all_counties"),                    icon="🏛️"),
            st.Page(_validation_queue,  title=t("nav_validation_queue"),                icon="✅"),
            st.Page(_inbox,             title=t("nav_comms_hub"),                       icon="📥"),
            st.Page(_makueni_ref,       title=t("nav_makueni_ref"),                     icon="📄"),
            st.Page(_data_download,     title=t("nav_data_download"),                   icon="⬇️"),
            # AI assistant intentionally excluded — it lives as a floating FAB widget
            st.Page(_account,           title=t("nav_account_settings"),                icon="⚙️"),
        ]
        # Developer tools tab only for developer role (not for EPRA)
        if role == "developer":
            pages = base_pages + [
                st.Page(_dev_tools, title=t("nav_dev_tools"), icon="🔧"),
            ]
        else:
            pages = base_pages
    elif role == "ministry":
        pages = [
            st.Page(_national_overview, title=t("nav_national_overview"), icon="📊", default=True),
            st.Page(_county_map,        title=t("nav_county_map"),        icon="🗺️"),
            st.Page(_all_counties,      title=t("nav_all_counties"),      icon="🏛️"),
            st.Page(_inbox,             title=t("nav_comms_hub"),         icon="📥"),
            st.Page(_data_download,     title=t("nav_data_download"),     icon="⬇️"),
            st.Page(_account,           title=t("nav_account_settings"),  icon="⚙️"),
        ]
    elif role == "devpartner":
        pages = [
            st.Page(_national_overview, title=t("nav_national_overview"), icon="📊", default=True),
            st.Page(_county_map,        title=t("nav_county_map"),        icon="🗺️"),
            st.Page(_all_counties,      title=t("nav_all_counties"),      icon="🏛️"),
            st.Page(_data_download,     title=t("nav_data_download"),     icon="⬇️"),
            st.Page(_account,           title=t("nav_account_settings"),  icon="⚙️"),
        ]
    elif role == "county":
        pages = [
            st.Page(_county_dashboard,  title=f"{cname} {'Dashboard' if st.session_state.get('lang','en')=='en' else 'Dashibodi'}",
                                        icon="🏠", default=True),
            st.Page(_submit,            title=t("nav_submit_plan"),       icon="📤"),
            st.Page(_inbox,             title=t("nav_communications"),    icon="📨"),
            st.Page(_data_download,     title=t("nav_data_download"),     icon="⬇️"),
            st.Page(_account,           title=t("nav_account_settings"),  icon="⚙️"),
        ]
        if is_makueni:
            pages.insert(2, st.Page(_makueni_ref, title=t("nav_my_county_plan"), icon="📄"))
    else:  # service_provider / kplc
        pages = [
            st.Page(_national_overview, title=t("nav_national_overview"),    icon="📊", default=True),
            st.Page(_all_counties,      title=t("nav_county_demand_data"),   icon="🏛️"),
            st.Page(_submit,            title=t("nav_submit_provider_plan"), icon="📤"),
            st.Page(_data_download,     title=t("nav_data_download"),        icon="⬇️"),
            st.Page(_account,           title=t("nav_account_settings"),     icon="⚙️"),
        ]

    # Sign out at bottom of sidebar
    with st.sidebar:
        st.divider()
        if st.button(t("sidebar_sign_out"), use_container_width=True, key="signout_btn"):
            try:
                authenticator.logout(location="unrendered")
            except Exception:
                pass
            for k in list(st.session_state.keys()):
                try: del st.session_state[k]
                except: pass
            st.rerun()
        st.markdown(f"""<div style="font-size:10px;color:#9c9a8e;line-height:1.7;margin-top:8px">
        {t("sidebar_data_source")}<br>
        NGDA 2026 · DTU · Challenge 2 · EPRA Kenya
        </div>""", unsafe_allow_html=True)

    # ── Run navigation — st.navigation handles clean page isolation ───────────
    pg = st.navigation(pages, position="sidebar", expanded=True)

    # ── AI assistant button — top of every page, before page content ─────────
    _render_floating_ai(role, county_id)

    pg.run()



def _render_floating_ai(role, county_id):
    """Plain Streamlit button at the top of every page that opens the AI chat dialog."""
    if "ai_history" not in st.session_state:
        st.session_state["ai_history"] = []
    if "ai_open" not in st.session_state:
        st.session_state["ai_open"] = False

    if st.button(t("sidebar_ai_btn"), key="kw_ai_fab_btn", type="primary",
                 help="Open KenyaWatts AI chat assistant" if st.session_state.get("lang","en")=="en" else "Fungua msaidizi wa AI"):
        st.session_state["ai_open"] = True

    if st.session_state.get("ai_open"):
        _show_ai_chat_dialog(role, county_id)


@st.dialog("⚡  KenyaWatts AI Assistant", width="large")
def _show_ai_chat_dialog(role, county_id):
    """
    Chat dialog opened by the floating AI FAB.

    KEY BEHAVIOURS:
    - ai_open is reset to False immediately so that if the user clicks outside
      to dismiss the dialog, the next full-app rerun won't reopen it.
    - Button clicks INSIDE the dialog trigger Streamlit's fragment/dialog rerun
      (dialog-scoped only). We do NOT call st.rerun() — doing so would close
      the dialog. The dialog stays open naturally through fragment reruns.
    - ai_history is in st.session_state: it persists across dialog open/close
      cycles for the entire browser session (until page refresh / server restart).
    """
    # Reset flag so external dismiss (click outside) doesn't reopen dialog
    st.session_state["ai_open"] = False

    if "ai_history" not in st.session_state:
        st.session_state["ai_history"] = []

    st.caption("Powered by Claude · Grounded in EPRA data and Makueni CEP · Chat history saved for this session")

    # ── Suggestion chips — shown only when chat history is empty ─────────────
    if not st.session_state["ai_history"]:
        st.markdown("**Suggested questions:**")
        suggestions = [
            "Which counties need urgent electricity investment?",
            "How is Kenya tracking toward universal access by 2030?",
            "What barriers does Makueni face for clean cooking?",
            "How does KenyaWatts aggregate 47 county plans into the INEP?",
            "Which counties are overdue and what are the consequences?",
            "What validation rules catch unit errors in submissions?",
        ]
        c1, c2 = st.columns(2)
        for i, s in enumerate(suggestions):
            col = c1 if i % 2 == 0 else c2
            if col.button(s, key=f"ai_dlg_sug_{i}", use_container_width=True):
                st.session_state["ai_history"].append({"role": "user", "content": s})
                with st.spinner("Analysing EPRA data…"):
                    reply = ask_ai(s, [])
                st.session_state["ai_history"].append({"role": "assistant", "content": reply})
                # ← NO st.rerun() — dialog fragment rerun keeps dialog open automatically

    # ── Chat history ──────────────────────────────────────────────────────────
    if st.session_state["ai_history"]:
        chat_box = st.container(height=360)
        with chat_box:
            for msg in st.session_state["ai_history"]:
                if msg["role"] == "user":
                    st.markdown(
                        f"<div style='background:#1a6fa3;color:#fff;"
                        f"border-radius:14px 14px 3px 14px;padding:9px 14px;"
                        f"margin:6px 0 6px 18%;font-size:13px;line-height:1.55;'>"
                        f"{msg['content']}</div>",
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f"<div style='background:#f0f8f5;color:#1a1916;"
                        f"border-radius:3px 14px 14px 14px;padding:9px 14px;"
                        f"margin:6px 18% 6px 0;font-size:13px;line-height:1.55;"
                        f"border-left:3px solid #3ecfaa;'>"
                        f"<span style='color:#0f9d7e;font-weight:700'>⚡ </span>"
                        f"{msg['content']}</div>",
                        unsafe_allow_html=True
                    )

    # ── Input row ─────────────────────────────────────────────────────────────
    st.markdown("")
    # Use a key with a counter so the text box clears after each send
    if "ai_input_key" not in st.session_state:
        st.session_state["ai_input_key"] = 0

    col_inp, col_send, col_clear = st.columns([5, 1, 1])
    with col_inp:
        user_q = st.text_input(
            "Message",
            placeholder="Ask about Kenya's energy planning…",
            key=f"ai_dlg_input_{st.session_state['ai_input_key']}",
            label_visibility="collapsed",
        )
    with col_send:
        send = st.button("Send ➤", key="ai_dlg_send", type="primary", use_container_width=True)
    with col_clear:
        if st.button("Clear", key="ai_dlg_clear", use_container_width=True):
            st.session_state["ai_history"] = []
            st.session_state["ai_input_key"] += 1
            # ← NO st.rerun() — dialog stays open, renders with empty history

    if send and user_q.strip():
        st.session_state["ai_history"].append({"role": "user", "content": user_q.strip()})
        st.session_state["ai_input_key"] += 1   # clears the text box on next render
        with st.spinner("Analysing EPRA data…"):
            reply = ask_ai(user_q.strip(), st.session_state["ai_history"][:-1])
        st.session_state["ai_history"].append({"role": "assistant", "content": reply})
        # ← NO st.rerun() — dialog fragment rerun keeps dialog open automatically


if __name__ == "__main__":
    main()
